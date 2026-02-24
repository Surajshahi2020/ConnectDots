# collect/views.py
from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.paginator import Paginator
from django.db import transaction
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from django.db import IntegrityError
from urllib.parse import quote
from django.db.models import Q, Avg  
from django.utils import timezone
from datetime import datetime, timedelta
from urllib.parse import urlencode
import pytz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.utils.html import escapejs
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
import io

from .models import ThreatAlert, CurrentInformation, NewsSource, DangerousKeyword, User, AutoNewsArticle, MapMarker, SocialMediaURL, SharedFile, Website
from collections import Counter
from django.utils import timezone
from datetime import datetime, timedelta
from django.db.models import Count
import json
from django.db.models.functions import TruncDate
from django.http import JsonResponse, HttpResponseRedirect, HttpResponse
from django.core.files.storage import FileSystemStorage
import os
import re
from pathlib import Path
from django.conf import settings
from django.core.cache import cache
from html import escape
from django.utils.safestring import mark_safe
from collections import defaultdict
from django.contrib.auth.hashers import check_password
from django.contrib.auth.hashers import make_password
from collect.decorators import session_auth_required
from django.http import JsonResponse
from collect.scrapers.kantipurdaily import kantipur_to_json
from collect.scrapers.kathmandu_post import kathmandu_post_extractor
from collect.scrapers.techpana import techpana_to_json 
from collect.scrapers.keyboard_techpana import keyboard_techpana_to_json 
from collect.scrapers.keyboard_kantipur import keyboard_kantipur_to_json 
from collect.scrapers.keyboard_kathmandupost import keyboard_kathmandu_post_to_json 
from collect.scrapers.keyboard_onlinekhabar import keyboard_onlinekhabar_to_json
from collect.scrapers.keyboard_paschim import keyboard_paschimnepal_to_json
from collect.scrapers.tvnepal import keyboard_onlinetvnepal_to_json
from collect.scrapers.osnepal import keyboard_osnepal_to_json
from collect.scrapers.eAdarsha import keyboard_eadarsha_to_json
from collect.scrapers.keyboard_arthasarokar import keyboard_arthasarokar_to_json
from collect.scrapers.keyboard_newsofnepal import keyboard_newsofnepal_to_json
from collect.scrapers.keyboard_rajdhanidaily import keyboard_rajdhanidaily_to_json


from django.urls import reverse  # Add this import
from django.db.models import Q
from datetime import datetime, timedelta
from utils.sentiment import predict_sentiment
from django.utils.dateparse import parse_date
from utils.permission import check_access
import requests



def loginLogic(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        
        # Debug print
        print(f"Login attempt: email={email}, password={'*' * len(password) if password else ''}")
        
        try:
            # Step 1: Check if user exists by email
            user = User.objects.get(email=email)
            
            # Step 2: Verify password
            if check_password(password, user.password):
                if user.is_active:
                    # ✅ LOG USER IN (creates session)
                    login(request, user)
                    
                    # ✅ STORE USER INFO IN SESSION (optional but useful)
                    request.session['auth'] = True
                    request.session['user_id'] = str(user.id)
                    request.session['user_role'] = user.role
                    request.session['user_username'] = user.username
                    request.session['user_unit'] = user.unit
                    request.session['user_email'] = user.email
                    request.session['user_rank'] = user.rank

                    return redirect('/adding_new/')  # ✅ Proper redirect
                else:
                    # Account disabled
                    return render(request, 'login.html', {
                        'alert_type': 'error',
                        'alert_message': 'Your account is disabled. Contact administrator.'
                    })
            else:
                # Invalid password
                return render(request, 'login.html', {
                    'alert_type': 'error',
                    'alert_message': 'Invalid email or password.'
                })
                
        except User.DoesNotExist:
            # User not found
            return render(request, 'login.html', {
                'alert_type': 'error',
                'alert_message': 'Invalid email or password.'
            })
        except Exception as e:
            # Unexpected error
            print(f"Login error: {e}")
            return render(request, 'login.html', {
                'alert_type': 'error',
                'alert_message': 'An unexpected error occurred during login.'
            })
    
    # GET request - show login form
    return render(request, 'login.html')

@session_auth_required
def dashboard(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    user = request.user
    role_map = {2: 'Admin', 1: 'Analyst', 0: 'Viewer'}
    role_display = role_map.get(getattr(user, 'role', 0), 'Unknown')

    query = request.GET.get('q', '').strip()
    
    # Start with appropriate queryset based on user's role and department
    if user.is_superuser:
        # Superusers can see ALL threats from ALL departments
        threats = ThreatAlert.objects.all()
    elif user.role == 2:  # Admin role
        # Admins can see threats from their own department only
        threats = ThreatAlert.objects.filter(
            models.Q(created_by__unit=user.unit) |  # Threats from same unit
            models.Q(created_by=user) |  # Threats created by themselves
            models.Q(created_by__isnull=True)  # System-generated threats
        ).distinct()
    else:  # Analyst (role=1) and Viewer (role=0)
        # Analysts and Viewers can only see threats from their department
        # that they have permission to view
        threats = ThreatAlert.objects.filter(
            models.Q(created_by__unit=user.unit)  # Only from same unit
        )
    
    threats = threats.order_by('-timestamp')

    if query:
        threats = threats.filter(
            Q(title__icontains=query) |
            Q(content__icontains=query) |
            Q(source__icontains=query) |
            Q(category__name__icontains=query) |
            Q(province__icontains=query)
        )

    # Get counts for dashboard statistics
    total_threats = threats.count()
    high_severity = threats.filter(severity='high').count()
    medium_severity = threats.filter(severity='medium').count()
    low_severity = threats.filter(severity='low').count()
    critical_severity = threats.filter(severity='critical').count()
    
    # Get threats by department/unit (if applicable)
    if hasattr(user, 'unit'):
        unit_threats = threats.filter(created_by__unit=user.unit).count()
    else:
        unit_threats = total_threats

    # Get recent threats (last 7 days)
    from datetime import datetime, timedelta
    one_week_ago = datetime.now() - timedelta(days=7)
    recent_threats = threats.filter(timestamp__gte=one_week_ago).count()

    # Get top categories in user's department
    from django.db.models import Count
    top_categories = threats.values(
        'category__name'
    ).annotate(
        count=Count('id')
    ).order_by('-count')[:5]

    # Pagination
    paginator = Paginator(threats, 5)
    page_number = request.GET.get('page')
    threats_page = paginator.get_page(page_number)

    return render(request, 'dashboard.html', {
        'threats': threats_page,
        'total_threats': total_threats,
        'high_severity': high_severity,
        'medium_severity': medium_severity,
        'low_severity': low_severity,
        'critical_severity': critical_severity,
        'recent_threats': recent_threats,
        'unit_threats': unit_threats,
        'role_display': role_display,
        'user_unit': user.unit if hasattr(user, 'unit') else 'N/A',
        'user_department': user.unit if hasattr(user, 'unit') else 'System',
        'search_query': query,
        'top_categories': top_categories,
        'is_superuser': user.is_superuser,
        'is_admin': getattr(user, 'role', 0) == 2,
    })

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db import models
from .models import ThreatAlert, ThreatCategory
import os

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from .models import ThreatCategory, User  # Adjust import path as needed

@login_required
def newsfeeding(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    # Get the current user
    current_user = request.user
    
    # Check if user account is active
    if not current_user.is_active:
        return render(request, 'news_add.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'user_categories': [],
            'selected_category': ''
        })
    
    # Get visible categories based on user's unit
    visible_categories = ThreatCategory.objects.filter(
        is_active=True
    ).order_by('name')
    
    # For non-superusers, filter categories created by users in the same unit
    if not current_user.is_superuser:
        visible_categories = visible_categories.filter(
            models.Q(created_by__unit=current_user.unit) |
            models.Q(created_by=current_user) |
            models.Q(created_by__isnull=True)  # Include system-created categories
        ).distinct()
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        source = request.POST.get('source', 'unknown').strip()
        category_id = request.POST.get('category', '').strip()
        description = request.POST.get('description', '').strip()
        url = request.POST.get('url', '').strip()
        severity = request.POST.get('severity', 'medium').lower()
        image = request.FILES.get('image')
        video = request.FILES.get('video')
        province = request.POST.get('province', '').strip()

        # Validation
        errors = []
        
        if not title:
            errors.append("Title is required.")
        if not description:
            errors.append("Description is required.")
        if not category_id:
            errors.append("Category is required.")
        else:
            # Validate that category exists and user has permission
            try:
                category = ThreatCategory.objects.get(id=category_id)
                if not category.is_active:
                    errors.append(f"Category '{category.name}' is not active.")
                
                # Check if user has permission to use this category
                if not request.user.is_superuser:
                    user_can_use = False
                    
                    # User can use categories from:
                    # 1. Their own creation
                    if category.created_by == request.user:
                        user_can_use = True
                    # 2. Users from the same unit
                    elif category.created_by and category.created_by.unit == request.user.unit:
                        user_can_use = True
                    # 3. System-created categories (created_by is None)
                    elif category.created_by is None:
                        user_can_use = True
                    
                    if not user_can_use:
                        errors.append(f"You don't have permission to use category '{category.name}'.")
                        
            except ThreatCategory.DoesNotExist:
                errors.append("Invalid category selected.")

        # Validate severity
        valid_severities = ['low', 'medium', 'high', 'critical']
        severity = severity if severity in valid_severities else 'medium'

        # Validate URL - if empty, set to None
        if not url:
            url = None
        elif not url.startswith(('http://', 'https://')):
            url = 'https://' + url

        # File validation
        if image:
            if image.size > 10 * 1024 * 1024:  # 10MB
                errors.append("Image size must be less than 10MB")
            
            valid_image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp']
            ext = os.path.splitext(image.name)[1].lower()
            if ext not in valid_image_extensions:
                errors.append("Invalid image format. Supported: JPG, JPEG, PNG, GIF, WEBP")

        if video:
            if video.size > 100 * 1024 * 1024:  # 100MB
                errors.append("Video size must be less than 100MB")
            
            valid_video_extensions = ['.mp4', '.mov', '.avi', '.webm', '.mkv']
            ext = os.path.splitext(video.name)[1].lower()
            if ext not in valid_video_extensions:
                errors.append("Invalid video format. Supported: MP4, MOV, AVI, WebM, MKV")

        if errors:
            # Use the same alert structure for form errors
            return render(request, 'news_add.html', {
                'alert_type': 'error',
                'alert_message': ' | '.join(errors),
                'form_data': request.POST,
                'user_categories': visible_categories,
                'selected_category': category_id
            })

        try:
            # Get the category object
            category = ThreatCategory.objects.get(id=category_id)
            
            # Create threat alert
            threat_alert = ThreatAlert.objects.create(
                title=title,
                content=description,
                source=source,
                category=category,
                url=url,
                severity=severity,
                image=image,
                video=video, 
                province=province,
                created_by=request.user
            )

            # Use success message format
            return render(request, 'news_add.html', {
                'alert_type': 'success',
                'alert_message': f'✅ Threat report #{threat_alert.id} saved successfully!',
                'user_categories': visible_categories,
                'selected_category': '',
                'form_data': {}  # Clear form data on success
            })
            
        except ThreatCategory.DoesNotExist:
            return render(request, 'news_add.html', {
                'alert_type': 'error',
                'alert_message': 'Selected category does not exist.',
                'form_data': request.POST,
                'user_categories': visible_categories,
                'selected_category': category_id
            })
        except Exception as e:
            import traceback
            print(f"Error saving threat alert: {str(e)}")
            print(traceback.format_exc())
            return render(request, 'news_add.html', {
                'alert_type': 'error',
                'alert_message': f'⚠️ Failed to save: {str(e)}',
                'form_data': request.POST,
                'user_categories': visible_categories,
                'selected_category': category_id
            })

    # GET request - show form with visible categories
    return render(request, 'news_add.html', {
        'user_categories': visible_categories,
        'selected_category': ''
    })

def newsSearching(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    selected = request.GET.getlist('category')
    
    # Get the current user's unit from session
    user_unit = request.session.get('user_unit', None)
    
    # Start with all threats
    threats = ThreatAlert.objects.all()
    
    # Filter threats based on the unit of the user who created the category
    if user_unit:
        threats = threats.filter(category__created_by__unit=user_unit)
    
    # Then filter by selected categories
    if selected:
        threats = threats.filter(category__id__in=selected)
    
    threats = threats.order_by('-timestamp')
    
    # Get categories available for the current user's unit
    if user_unit:
        available_categories = ThreatCategory.objects.filter(created_by__unit=user_unit)
    else:
        available_categories = ThreatCategory.objects.all()
    
    # 🔢 Chart data - Count threats per category
    # First, get all category IDs from filtered threats
    category_ids = threats.values_list('category_id', flat=True)
    
    # Count occurrences of each category ID
    category_counts = Counter(category_ids)
    
    # Prepare chart data
    chart_data = []
    for category in available_categories:
        if not selected or str(category.id) in selected:
            chart_data.append({
                'label': category.name,  # Assuming ThreatCategory has a 'name' field
                'count': category_counts.get(category.id, 0),
                'value': category.id
            })
    
    # Prepare categories for template - list of (id, name) tuples
    categories_for_template = [(str(category.id), category.name) for category in available_categories]
    
    # 📄 Paginate (15 items per page)
    paginator = Paginator(threats, 15)
    page_number = request.GET.get('page')
    threats_page = paginator.get_page(page_number)

    return render(request, 'searchNews.html', {
        'threats': threats_page,
        'all_categories': categories_for_template,
        'selected_categories': selected,
        'chart_data': chart_data,
        'user_unit': user_unit,
    })

# At the VERY TOP of your views.py file, use these imports:
from datetime import datetime, timedelta  # Changed this line
from django.utils import timezone
from django.db.models import Count, Q
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib.auth import get_user_model
from .models import ThreatAlert  # Add this import for the ThreatAlert model
from utils.permission import check_access

@login_required
def newsVisualization(request):
    if not check_access(request):
        return redirect('logout')
    
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return render(request, 'newsVisualization.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'total_threats': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'critical_severity': 0,
            'chart_labels': [],
            'chart_data': [],
            'severity_labels': [],
            'severity_data': [],
            'timeline_labels': [],
            'timeline_data': [],
            'trend_labels': [],
            'trend_data': {'critical': [], 'high': [], 'medium': [], 'low': []},
            'sources_labels': [],
            'sources_data': [],
            'unit_stats': {},
            'user_unit': user.unit if hasattr(user, 'unit') else None,
            'is_superuser': user.is_superuser,
            'is_admin': getattr(user, 'role', 0) == 2,
            'department_data': [],
        })
    
    # Get threats based on user's unit
    if user.is_superuser:
        # Superusers can see all threats
        threats = ThreatAlert.objects.all()
    elif hasattr(user, 'role') and user.role == 2:  # Admin
        # Admins can see threats from their unit + their own + system threats
        threats = ThreatAlert.objects.filter(
            Q(created_by__unit=user.unit) |
            Q(created_by=user) |
            Q(created_by__isnull=True)
        )
    else:
        # Regular users can only see threats from their unit
        threats = ThreatAlert.objects.filter(
            created_by__unit=user.unit
        )
    
    threats = threats.order_by('-timestamp')

    # Total counts
    total_threats = threats.count()
    high_severity = threats.filter(severity='high').count()
    medium_severity = threats.filter(severity='medium').count()
    low_severity = threats.filter(severity='low').count()
    critical_severity = threats.filter(severity='critical').count()

    # Unit-specific statistics
    unit_stats = {}
    if hasattr(user, 'unit') and user.unit:
        unit_threats = threats.filter(created_by__unit=user.unit)
        unit_stats = {
            'total': unit_threats.count(),
            'high': unit_threats.filter(severity='high').count(),
            'medium': unit_threats.filter(severity='medium').count(),
            'low': unit_threats.filter(severity='low').count(),
            'critical': unit_threats.filter(severity='critical').count(),
        }

    # Category distribution for chart
    category_data = threats.values('category__name').annotate(
        count=Count('id')
    ).order_by('-count')[:10]  # Top 10 categories

    # Prepare chart data
    category_labels = []
    category_counts = []
    for item in category_data:
        category_name = item['category__name'] if item['category__name'] else 'Uncategorized'
        category_labels.append(category_name)
        category_counts.append(item['count'])

    # Severity distribution for chart
    severity_labels = []
    severity_counts = []
    severity_order = ['critical', 'high', 'medium', 'low']
    
    # Create ordered list
    for severity in severity_order:
        count = threats.filter(severity=severity).count()
        if count > 0:
            severity_labels.append(severity.title())
            severity_counts.append(count)

    # Timeline data (last 30 days)
    thirty_days_ago = timezone.now() - timedelta(days=30)  # Changed datetime.timedelta to timedelta
    
    # Simple manual approach for timeline data
    date_counts = {}
    for i in range(30):
        date = (timezone.now() - timedelta(days=i)).date()  # Changed datetime.timedelta to timedelta
        date_counts[date] = 0
    
    # Count threats for each date
    for threat in threats.filter(timestamp__gte=thirty_days_ago):
        threat_date = threat.timestamp.date()
        if threat_date in date_counts:
            date_counts[threat_date] += 1
    
    # Convert to sorted lists
    sorted_dates = sorted(date_counts.keys())
    timeline_labels = [date.strftime('%m/%d') for date in sorted_dates]
    timeline_counts = [date_counts[date] for date in sorted_dates]

    # Threat sources from the source field
    sources_data = threats.values('source').annotate(
        count=Count('id')
    ).order_by('-count')[:10]  # Top 10 sources

    sources_labels = []
    sources_counts = []
    for item in sources_data:
        source_name = item['source'] if item['source'] else 'Unknown'
        sources_labels.append(source_name)
        sources_counts.append(item['count'])

    # Trend data (last 7 days by severity)
    # Get dates for last 7 days
    dates = [(timezone.now() - timedelta(days=i)).date() for i in range(6, -1, -1)]  # Changed datetime.timedelta to timedelta
    trend_labels = [date.strftime('%m/%d') for date in dates]
    
    # Initialize trend data structure
    trend_data = {
        'critical': [0] * 7,
        'high': [0] * 7,
        'medium': [0] * 7,
        'low': [0] * 7
    }
    
    # Get threat counts for each severity for last 7 days
    for i, date in enumerate(dates):
        # Create timezone-aware datetime objects - FIXED VERSION
        day_start = timezone.make_aware(
            datetime.combine(date, datetime.min.time()),  # Changed datetime.datetime to datetime
            timezone.get_current_timezone()
        )
        day_end = timezone.make_aware(
            datetime.combine(date, datetime.max.time()),  # Changed datetime.datetime to datetime
            timezone.get_current_timezone()
        )
        
        day_threats = threats.filter(timestamp__range=(day_start, day_end))
        
        trend_data['critical'][i] = day_threats.filter(severity='critical').count()
        trend_data['high'][i] = day_threats.filter(severity='high').count()
        trend_data['medium'][i] = day_threats.filter(severity='medium').count()
        trend_data['low'][i] = day_threats.filter(severity='low').count()

    # Department/Unit comparison (for superusers and admins only)
    department_data = []
    if user.is_superuser or (hasattr(user, 'role') and user.role == 2):
        User = get_user_model()
        
        department_data = User.objects.values(
            'unit'
        ).annotate(
            threat_count=Count('created_alerts')
        ).exclude(unit__isnull=True).exclude(unit='').order_by('-threat_count')[:5]

    # Recent threats for table
    recent_threats = threats[:10]

    return render(request, 'newsVisualization.html', {
        'threats': recent_threats,
        'total_threats': total_threats,
        'high_severity': high_severity,
        'medium_severity': medium_severity,
        'low_severity': low_severity,
        'critical_severity': critical_severity,
        'chart_labels': category_labels,
        'chart_data': category_counts,
        'severity_labels': severity_labels,
        'severity_data': severity_counts,
        'timeline_labels': timeline_labels,
        'timeline_data': timeline_counts,
        'trend_labels': trend_labels,
        'trend_data': trend_data,
        'sources_labels': sources_labels,
        'sources_data': sources_counts,
        'unit_stats': unit_stats,
        'user_unit': user.unit if hasattr(user, 'unit') else None,
        'is_superuser': user.is_superuser,
        'is_admin': getattr(user, 'role', 0) == 2,
        'department_data': department_data,
    })

@login_required
def newsTrending(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return render(request, 'newsTrending.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'threats': [],
            'critical_threats': [],
            'total_critical_with_videos': 0,
            'total_all_videos': 0,
            'page_obj': None,
            'user_unit': None,
            'unit_stats': {},
        })
    
    # Start with base queryset based on user's unit
    if user.is_superuser:
        base_threats = ThreatAlert.objects.all()
    elif hasattr(user, 'role') and user.role == 2:  # Admin
        base_threats = ThreatAlert.objects.filter(
            Q(created_by__unit=user.unit) |
            Q(created_by=user) |
            Q(created_by__isnull=True)
        )
    else:
        base_threats = ThreatAlert.objects.filter(
            created_by__unit=user.unit
        )
    
    # Get critical threats with videos from user's unit
    critical_threats_with_videos = base_threats.filter(
        severity='critical'
    ).exclude(
        Q(video__isnull=True) | Q(video='')
    ).order_by('-timestamp')
    
    # Get all threats with videos for statistics (from user's unit)
    all_threats_with_videos = base_threats.exclude(
        Q(video__isnull=True) | Q(video='')
    )
    
    # Get unit-specific statistics
    unit_stats = {}
    if hasattr(user, 'unit') and user.unit:
        unit_critical_videos = base_threats.filter(
            severity='critical',
            created_by__unit=user.unit
        ).exclude(
            Q(video__isnull=True) | Q(video='')
        ).count()
        
        unit_all_videos = base_threats.filter(
            created_by__unit=user.unit
        ).exclude(
            Q(video__isnull=True) | Q(video='')
        ).count()
        
        unit_stats = {
            'critical_videos': unit_critical_videos,
            'all_videos': unit_all_videos,
        }
    
    # Get category statistics for videos
    video_categories = all_threats_with_videos.values(
        'category__name'
    ).annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # Get severity distribution for videos
    video_severities = all_threats_with_videos.values(
        'severity'
    ).annotate(
        count=Count('id')
    ).order_by('severity')
    
    # Format severity names for display
    severity_names = {
        'critical': 'Critical',
        'high': 'High',
        'medium': 'Medium',
        'low': 'Low'
    }
    
    video_severity_data = []
    for item in video_severities:
        video_severity_data.append({
            'severity': severity_names.get(item['severity'], item['severity'].title()),
            'count': item['count']
        })
    
    # Get recent videos (last 7 days)
    from datetime import timedelta
    one_week_ago = timezone.now() - timedelta(days=7)
    recent_videos = all_threats_with_videos.filter(
        timestamp__gte=one_week_ago
    ).count()
    
    # Add pagination for critical threats with videos
    paginator = Paginator(critical_threats_with_videos, 15)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'newsTrending.html', {
        'threats': page_obj,
        'critical_threats': critical_threats_with_videos,
        'total_critical_with_videos': critical_threats_with_videos.count(),
        'total_all_videos': all_threats_with_videos.count(),
        'page_obj': page_obj,
        'user_unit': user.unit if hasattr(user, 'unit') else None,
        'unit_stats': unit_stats,
        'video_categories': video_categories,
        'video_severity_data': video_severity_data,
        'recent_videos': recent_videos,
        'is_superuser': user.is_superuser,
        'is_admin': getattr(user, 'role', 0) == 2,
    })


def newsReport(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    return render(request, 'news_report.html', {
    })

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import redirect
from django.utils import timezone
from django.db.models import Q, Count
import pytz
from io import BytesIO
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

@login_required
def generate_word_report(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Generate and download Word document report with severity filter
    """
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return HttpResponse("Your account is disabled. Contact administrator.", status=403)
    
    if request.method != 'POST':
        return redirect('newsCurrent')
    
    try:
        # Get form data
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        province = request.POST.get('province', '')
        
        # Get selected severities
        selected_severities = request.POST.getlist('severity')
        
        # Determine if "all" is selected
        if 'all' in selected_severities:
            include_all_severities = True
            selected_severities = [sev for sev in selected_severities if sev != 'all']
        else:
            include_all_severities = False
        
        # Validate dates
        if not start_date or not end_date:
            return HttpResponse("Start date and end date are required.", status=400)
        
        # Convert dates to datetime using Kathmandu timezone
        try:
            from datetime import datetime
            start_naive = datetime.strptime(start_date, '%Y-%m-%d')
            end_naive = datetime.strptime(end_date, '%Y-%m-%d')
            
            kathmandu_tz = pytz.timezone('Asia/Kathmandu')
            
            start_datetime = kathmandu_tz.localize(start_naive)
            end_datetime = kathmandu_tz.localize(end_naive.replace(hour=23, minute=59, second=59))
            
        except ValueError:
            return HttpResponse("Invalid date format", status=400)
        
        # Get threats based on user's unit
        if user.is_superuser:
            alerts = ThreatAlert.objects.all()
        elif hasattr(user, 'role') and user.role == 2:  # Admin
            alerts = ThreatAlert.objects.filter(
                Q(created_by__unit=user.unit) |
                Q(created_by=user) |
                Q(created_by__isnull=True)
            )
        else:
            alerts = ThreatAlert.objects.filter(
                created_by__unit=user.unit
            )
        
        # Apply date filter
        alerts = alerts.filter(
            created_at__gte=start_datetime,
            created_at__lte=end_datetime
        )
        
        # Apply province filter if selected
        if province:
            alerts = alerts.filter(province=province)
        
        # Apply severity filter
        if include_all_severities:
            pass  # Don't filter by severity - include all
        elif selected_severities:
            alerts = alerts.filter(severity__in=selected_severities)
        else:
            alerts = alerts.none()  # No severity selected, return empty
        
        # Get statistics
        total_alerts = alerts.count()
        
        # Count by severity
        severity_order = ['low', 'medium', 'high', 'critical']
        severity_counts = alerts.values('severity').annotate(
            count=Count('id')
        )
        severity_counts_list = list(severity_counts)
        severity_counts_list.sort(key=lambda x: severity_order.index(x['severity']) if x['severity'] in severity_order else 999)
        
        # Count by category
        category_counts = alerts.values('category__name').annotate(
            count=Count('id')
        ).order_by('-count')[:10]
        
        # Get province name for display
        province_name = 'All Provinces'
        if province:
            province_dict = dict(ThreatAlert.PROVINCE_CHOICES)
            province_name = province_dict.get(province, province.title())
        
        # Get severity display text
        if include_all_severities:
            severity_display = 'All Levels'
        elif selected_severities:
            severity_mapping = {
                'low': 'Low',
                'medium': 'Medium', 
                'high': 'High',
                'critical': 'Critical'
            }
            sorted_severities = sorted(
                selected_severities, 
                key=lambda x: severity_order.index(x) if x in severity_order else 999
            )
            severity_display = ', '.join([severity_mapping.get(sev, sev.title()) for sev in sorted_severities])
        else:
            severity_display = 'None Selected'
        
        # Get current time in Kathmandu
        kathmandu_tz = pytz.timezone('Asia/Kathmandu')
        now_kathmandu = timezone.now().astimezone(kathmandu_tz)
        generated_date = now_kathmandu.strftime('%Y-%m-%d')
        generated_time = now_kathmandu.strftime('%I:%M %p')
        
        # ================= CREATE WORD DOCUMENT =================
        document = Document()
        
        # Set document styles
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        
        # Add title
        title = document.add_heading('Threat Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add subtitle
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = f'Generated by: {user.username}'
        if hasattr(user, 'unit') and user.unit:
            subtitle_text += f' ({user.unit})'
        subtitle_text += f' | {generated_date} at {generated_time} NPT'
        
        run = subtitle.add_run(subtitle_text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Add report parameters
        document.add_heading('Report Parameters', 1)
        
        params_table = document.add_table(rows=6, cols=2)
        params_table.style = 'Colorful Grid Accent 1'
        
        for row in params_table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(4.2)
        
        params_data = [
            ('Date Range:', f'{start_date} to {end_date}'),
            ('Time Zone:', 'Asia/Kathmandu (NPT)'),
            ('Province:', province_name),
            ('Severity:', severity_display),
            ('Total Alerts:', str(total_alerts)),
            ('User Unit:', user.unit if hasattr(user, 'unit') and user.unit else 'All Units'),
        ]
        
        for i, (label, value) in enumerate(params_data):
            params_table.rows[i].cells[0].text = label
            params_table.rows[i].cells[1].text = value
            label_cell = params_table.rows[i].cells[0]
            label_cell.paragraphs[0].runs[0].bold = True
            label_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            # Add background to label cells
            tcPr = label_cell._element.tcPr
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '4F81BD')
            tcPr.append(shading)
        
        # Add summary statistics
        document.add_heading('Summary Statistics', 1)
        
        # Severity breakdown
        document.add_heading('Alerts by Severity Level', 2)
        if severity_counts_list:
            severity_table = document.add_table(rows=len(severity_counts_list) + 1, cols=3)
            severity_table.style = 'Colorful Grid Accent 2'
            
            # Header row
            header_cells = severity_table.rows[0].cells
            header_cells[0].text = 'Severity Level'
            header_cells[1].text = 'Count'
            header_cells[2].text = 'Percentage'
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = cell._element.tcPr
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '1F497D')
                tcPr.append(shading)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            # Data rows
            severity_display_names = {
                'low': 'Low',
                'medium': 'Medium',
                'high': 'High', 
                'critical': 'Critical'
            }
            
            severity_row_colors = {
                'critical': 'C0504D',
                'high': 'F79646',
                'medium': 'FFC000',
                'low': '9BBB59'
            }
            
            for i, item in enumerate(severity_counts_list, 1):
                row_cells = severity_table.rows[i].cells
                display_name = severity_display_names.get(item['severity'], item['severity'].title())
                row_cells[0].text = display_name
                row_cells[1].text = str(item['count'])
                
                if total_alerts > 0:
                    percentage = (item['count'] / total_alerts) * 100
                    row_cells[2].text = f"{percentage:.1f}%"
                else:
                    row_cells[2].text = "0%"
                
                # Add colored background
                row_color = severity_row_colors.get(item['severity'], 'FFFFFF')
                for cell in row_cells:
                    tcPr = cell._element.tcPr
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), row_color)
                    tcPr.append(shading)
                
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            document.add_paragraph('No severity data available.').italic = True
        
        # Category breakdown
        document.add_heading('Top Categories', 2)
        if category_counts:
            category_table = document.add_table(rows=len(category_counts) + 1, cols=2)
            category_table.style = 'Light Grid Accent 1'
            
            header_cells = category_table.rows[0].cells
            header_cells[0].text = 'Category'
            header_cells[1].text = 'Count'
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, item in enumerate(category_counts, 1):
                row_cells = category_table.rows[i].cells
                row_cells[0].text = item['category__name'] if item['category__name'] else 'Uncategorized'
                row_cells[1].text = str(item['count'])
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            document.add_paragraph('No category data available.').italic = True
        
        # Recent alerts table
        document.add_heading('Recent Alerts (Last 20)', 1)
        recent_alerts = alerts.order_by('-created_at')[:20]
        
        if recent_alerts:
            # Create table
            has_url_field = hasattr(ThreatAlert, 'url')
            
            if has_url_field:
                alerts_table = document.add_table(rows=len(recent_alerts) + 1, cols=8)
                headers = ['Title', 'Severity', 'Category', 'Province', 'Creator Unit', 'Date (NPT)', 'Source URL']
                col_count = 7
            else:
                alerts_table = document.add_table(rows=len(recent_alerts) + 1, cols=7)
                headers = ['Title', 'Severity', 'Category', 'Province', 'Creator Unit', 'Date (NPT)']
                col_count = 6
            
            alerts_table.style = 'Medium Grid 1 Accent 1'
            
            # Header row
            header_cells = alerts_table.rows[0].cells
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                
                tcPr = header_cells[i]._element.tcPr
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '2E75B6')
                tcPr.append(shading)
            
            # Data rows
            severity_names = {'low': 'Low', 'medium': 'Medium', 'high': 'High', 'critical': 'Critical'}
            severity_colors = {
                'critical': ('C00000', RGBColor(255, 255, 255)),
                'high': ('FF6600', RGBColor(255, 255, 255)),
                'medium': ('FFC000', RGBColor(0, 0, 0)),
                'low': ('00B050', RGBColor(255, 255, 255))
            }
            
            province_names = dict(ThreatAlert.PROVINCE_CHOICES)
            
            for i, alert in enumerate(recent_alerts, 1):
                row_cells = alerts_table.rows[i].cells
                
                # Title
                title_text = alert.title[:40] + ('...' if len(alert.title) > 40 else '')
                row_cells[0].text = title_text
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                
                # Severity
                severity_text = severity_names.get(alert.severity, alert.severity.title() if alert.severity else 'N/A')
                row_cells[1].text = severity_text
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].paragraphs[0].runs[0].bold = True
                
                if alert.severity in severity_colors:
                    bg_color, text_color = severity_colors[alert.severity]
                    row_cells[1].paragraphs[0].runs[0].font.color.rgb = text_color
                    tcPr = row_cells[1]._element.tcPr
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), bg_color)
                    tcPr.append(shading)
                
                # Category
                category_name = alert.category.name if alert.category else 'Uncategorized'
                row_cells[2].text = category_name
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Province
                province_text = province_names.get(alert.province, alert.province.title() if alert.province else 'N/A')
                row_cells[3].text = province_text
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Creator Unit
                if alert.created_by and hasattr(alert.created_by, 'unit') and alert.created_by.unit:
                    creator_unit = alert.created_by.unit
                    if creator_unit == user.unit:
                        row_cells[4].text = f"★ {creator_unit}"
                    else:
                        row_cells[4].text = creator_unit
                else:
                    row_cells[4].text = 'System'
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Date
                if alert.created_at:
                    alert_date_kathmandu = alert.created_at.astimezone(kathmandu_tz)
                    date_text = alert_date_kathmandu.strftime('%Y-%m-%d\n%H:%M')
                else:
                    date_text = 'N/A'
                row_cells[5].text = date_text
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[5].paragraphs[0].runs[0].font.size = Pt(8)
                
                # Source URL (if available)
                if has_url_field and col_count > 6:
                    if hasattr(alert, 'url') and alert.url:
                        url_text = str(alert.url)[:30] + ('...' if len(str(alert.url)) > 30 else '')
                        row_cells[6].text = url_text
                        row_cells[6].paragraphs[0].runs[0].font.size = Pt(8)
                        row_cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)
                    else:
                        row_cells[6].text = 'No URL'
                        row_cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(153, 153, 153)
                        row_cells[6].paragraphs[0].runs[0].italic = True
        else:
            document.add_paragraph('No alerts found for the selected criteria.').italic = True
        
        # Add footer
        document.add_heading('Report Information', 1)
        notes = [
            f'• Report generated: {generated_date} at {generated_time} Nepal Time',
            f'• Date Range: {start_date} to {end_date}',
            f'• Province Filter: {province_name}',
            f'• Severity Levels: {severity_display}',
            f'• Total Alerts: {total_alerts}',
            f'• Generated by: {user.username}',
        ]
        
        if hasattr(user, 'unit') and user.unit:
            notes.append(f'• User Unit: {user.unit}')
        
        for note in notes:
            p = document.add_paragraph(style='ListBullet')
            run = p.add_run(note.replace('• ', ''))
            run.font.size = Pt(10)
        
        # ================= SAVE DOCUMENT TO BYTESIO =================
        # *** THIS IS WHERE file_stream IS DEFINED ***
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        # Create filename
        filename_date = now_kathmandu.strftime('%Y%m%d_%H%M')
        filename = f'threat_report_{start_date}_to_{end_date}_{filename_date}'
        if hasattr(user, 'unit') and user.unit:
            clean_unit = user.unit.replace(" ", "_").replace("/", "_").replace("\\", "_")
            filename += f'_{clean_unit}'
        if province:
            filename += f'_{province}'
        if not include_all_severities and selected_severities:
            sorted_sev = sorted(selected_severities, key=lambda x: severity_order.index(x) if x in severity_order else 999)
            severity_str = '_'.join(sorted_sev)
            filename += f'_{severity_str}'
        filename += '.docx'
        
        # ================= RETURN RESPONSE =================
        response = HttpResponse(
            file_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        import traceback
        print(f"Error generating report: {str(e)}")
        print(traceback.format_exc())
        return HttpResponse(f"Error generating report: {str(e)}", status=500)

# Optional: Add this view for displaying the form
def report_form(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Display the report generation form"""
    return render(request, 'newsCurrent.html')

from django.db.models import Q
from django.core.paginator import Paginator
from urllib.parse import urlencode
from datetime import datetime

@login_required
def newsCurrent(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return render(request, 'newsCurrent.html', {
            'page_obj': None,
            'total_count': 0,
            'search': '',
            'province': '',
            'status': '',
            'start_date': '',
            'end_date': '',
            'query_string': '',
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
        })
    
    # Start with filtered records based on user's unit
    if user.is_superuser:
        # Superusers can see all records
        queryset = CurrentInformation.objects.all()
    elif hasattr(user, 'role') and user.role == 2:  # Admin
        # Admins can see records from their unit + their own + system records
        queryset = CurrentInformation.objects.filter(
            Q(created_by__unit=user.unit) |
            Q(created_by=user) |
            Q(created_by__isnull=True)
        )
    else:
        # Regular users can only see records from their unit
        queryset = CurrentInformation.objects.filter(
            created_by__unit=user.unit
        )
    
    # Get filter parameters
    search = request.GET.get('search', '').strip()
    province = request.GET.get('province', '')
    status = request.GET.get('status', '')
    start_date_str = request.GET.get('start_date', '')
    end_date_str = request.GET.get('end_date', '')
    
    # Apply search filter
    if search:
        queryset = queryset.filter(
            Q(leader__icontains=search) |
            Q(location__icontains=search) |
            Q(vehicle__icontains=search) |
            Q(description__icontains=search) |
            Q(number__icontains=search) |
            Q(province__icontains=search) |
            Q(created_by__username__icontains=search) |  # Search by creator username
            Q(created_by__unit__icontains=search)  # Search by unit name
        )
    
    # Apply province filter
    if province:
        queryset = queryset.filter(province=province)
    
    # Apply status filter
    if status:
        queryset = queryset.filter(status=status)
    
    # Apply date range filter
    try:
        # Start date filter
        if start_date_str:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            queryset = queryset.filter(created_at__date__gte=start_date)
        
        # End date filter
        if end_date_str:
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            queryset = queryset.filter(created_at__date__lte=end_date)
    except ValueError:
        # If date format is invalid, ignore the filter
        pass
    
    # Order by creation date (newest first)
    queryset = queryset.order_by('-created_at')
    
    # Get total count
    total_count = queryset.count()
    
    # Get unit statistics
    unit_stats = {}
    if hasattr(user, 'unit'):
        unit_queryset = queryset.filter(created_by__unit=user.unit)
        unit_stats = {
            'total': unit_queryset.count(),
            'pending': unit_queryset.filter(status='pending').count(),
            'completed': unit_queryset.filter(status='completed').count(),
            'cancelled': unit_queryset.filter(status='cancelled').count(),
        }
    
    # Pagination
    paginator = Paginator(queryset, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Create a query string without the page parameter
    query_params = {}
    for key, value in request.GET.items():
        if key != 'page':
            query_params[key] = value
    
    # Generate URL query string
    query_string = urlencode(query_params)
    
    # Prepare context
    context = {
        'page_obj': page_obj,
        'total_count': total_count,
        'search': search,
        'province': province,
        'status': status,
        'start_date': start_date_str,
        'end_date': end_date_str,
        'query_string': query_string,
        'user_unit': user.unit if hasattr(user, 'unit') else None,
        'unit_stats': unit_stats,
        'is_superuser': user.is_superuser,
        'is_admin': getattr(user, 'role', 0) == 2,
    }
    return render(request, 'newsCurrent.html', context)

@login_required
def newsSpy(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Combined view: list + manual form submission with alert context"""
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        current_info_list = CurrentInformation.objects.none()
        paginator = Paginator(current_info_list, 15)
        page_obj = paginator.get_page(request.GET.get('page'))
        return render(request, 'newsSpy.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'page_obj': page_obj,
        })
    
    if request.method == 'POST':
        # Extract and clean data
        timing = request.POST.get('timing', '').strip()
        location = request.POST.get('location', '').strip()
        leader = request.POST.get('leader', '').strip()
        number = request.POST.get('number', '').strip() or None
        vehicle = request.POST.get('vehicle', '').strip() or None
        description = request.POST.get('description', '').strip() or None
        status = request.POST.get('status', 'pending').strip()
        province = request.POST.get('province', '').strip() or None

        # 🔒 Validation: Required fields
        if not timing or not location or not leader:
            # Get appropriate list based on user role/unit
            if user.is_superuser:
                current_info_list = CurrentInformation.objects.all()
            else:
                current_info_list = CurrentInformation.objects.filter(
                    models.Q(created_by__unit=user.unit) |
                    models.Q(created_by=user) |
                    models.Q(created_by__isnull=True)
                )
            current_info_list = current_info_list.order_by('-created_at')
            paginator = Paginator(current_info_list, 15)
            page_obj = paginator.get_page(request.GET.get('page'))
            return render(request, 'newsSpy.html', {
                'alert_type': 'error',
                'alert_message': 'Timestamp, Location, and Field Leader are required.',
                'form_data': {
                    'timing': timing,
                    'location': location,
                    'leader': leader,
                    'number': number,
                    'vehicle': vehicle,
                    'description': description,
                    'province': province,
                    'status': status
                },
                'page_obj': page_obj,
            })

        # 🔒 Validate status
        valid_statuses = ['pending', 'completed', 'cancelled']
        if status not in valid_statuses:
            status = 'pending'

        try:
            # ✅ Save to DB
            CurrentInformation.objects.create(
                timing=timing,
                location=location,
                leader=leader,
                number=number,
                vehicle=vehicle,
                description=description,
                province=province,
                status=status,
                created_by=request.user
            )
            
            # On success: Get appropriate list based on user role/unit
            if user.is_superuser:
                current_info_list = CurrentInformation.objects.all()
            else:
                current_info_list = CurrentInformation.objects.filter(
                    models.Q(created_by__unit=user.unit) |
                    models.Q(created_by=user) |
                    models.Q(created_by__isnull=True)
                )
            current_info_list = current_info_list.order_by('-created_at')
            paginator = Paginator(current_info_list, 15)
            page_obj = paginator.get_page(1)  # Go to page 1 to show new entry
            return render(request, 'newsSpy.html', {
                'alert_type': 'success',
                'alert_message': '✅ Intel report submitted successfully!',
                'page_obj': page_obj,
            })

        except Exception as e:
            # Get appropriate list based on user role/unit
            if user.is_superuser:
                current_info_list = CurrentInformation.objects.all()
            else:
                current_info_list = CurrentInformation.objects.filter(
                    models.Q(created_by__unit=user.unit) |
                    models.Q(created_by=user) |
                    models.Q(created_by__isnull=True)
                )
            current_info_list = current_info_list.order_by('-created_at')
            paginator = Paginator(current_info_list, 7)
            page_obj = paginator.get_page(request.GET.get('page'))
            return render(request, 'newsSpy.html', {
                'alert_type': 'error',
                'alert_message': f'⚠️ Failed to save: {str(e)}',
                'form_data': {
                    'timing': timing,
                    'location': location,
                    'leader': leader,
                    'number': number,
                    'vehicle': vehicle,
                    'description': description,
                    'status': status,
                },
                'page_obj': page_obj,
            })

    # GET request - Get appropriate list based on user role/unit
    if user.is_superuser:
        current_info_list = CurrentInformation.objects.all()
    else:
        current_info_list = CurrentInformation.objects.filter(
            models.Q(created_by__unit=user.unit) |
            models.Q(created_by=user) |
            models.Q(created_by__isnull=True)
        )
    
    current_info_list = current_info_list.order_by('-created_at')
    paginator = Paginator(current_info_list, 15)
    page_obj = paginator.get_page(request.GET.get('page'))
    
    return render(request, 'newsSpy.html', {
        'page_obj': page_obj,
        'user_unit': user.unit if hasattr(user, 'unit') else None,
    })

def loginPage(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    return render(request, 'login.html', {
    })

def logOut(request):
    request.session.flush()
    return redirect('/login/')

def signinPage(request):
    return render(request, 'signin.html', {
    })

def signinAddView(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    if request.method == 'POST':
        # Get all form values
        username = request.POST.get('username')
        email = request.POST.get('email')
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')
        phone = request.POST.get('phone')
        unit_id = request.POST.get('unit_id')  # This is the unit code
        role = request.POST.get('role')
        rank_code = request.POST.get('rank_code')  # This is the rank code
        
        # Validation
        if not username or not email or not password1 or not password2:
            alert_type = 'error'
            alert_message = 'All required fields must be filled!'
            return render(request, 'signin.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
            
            
        if password1 != password2:
            alert_type = 'error'
            alert_message = 'Passwords do not match!'
            return render(request, 'signin.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
        
        if User.objects.filter(username=username).exists():
            alert_type = 'error'
            alert_message = 'Username already exists!'
            return render(request, 'signin.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
            
            
        if User.objects.filter(email=email).exists():
            alert_type = 'error'
            alert_message = 'Email already registered!'
            return render(request, 'signin.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
        
        # Create user
        try:
            # Get the actual unit name from your unit mapping
            unit_mapping = {
                '74': 'अति विशिष्त व्यक्ति सुरक्षा निर्देशनालय',
                '27': 'अनुसन्धान र विस्तार निर्देशनालय',
                '58': 'असवाव निर्देशनालय',
                '113': 'अस्थायी निबेश',
                '106': 'आपुर्ति तथा परिवहन गण',
                '86': 'आपुर्ति तथा परिवहन निर्देशनालय',
                '115': 'आर्टिलरी तालिम शिक्षालय',
                '18': 'आवास तथा विमा शाखा, कल्याणकारी योजना निर्देशनालय',
                '49': 'ईन्जिनियरिङ्ग विभाग',
                '127': 'उत्तर पश्चिम पृतना',
                '124': 'उपत्यका पृतना',
                '91': 'एम.र्इ.एस. (के.शा.)',
                '90': 'एम्युनिशन कार्य महाशाखा',
                '68': 'कलेज अफ मेडिकल पोलिटेक्निक',
                '12': 'कल्याणकारी योजना निर्देशनालय',
                '23': 'कल्याणकारी स्वास्थ्य सेवा व्यवस्थापन महाशाखा',
                '21': 'का. मु. प्रधान सेनापतिको कार्यालय',
                '136': 'काठमाडौँ-तराई मधेश द्रुतमार्ग सडक आयोजना कार्यालय',
                '46': 'कार्यरथीको कार्यालय',
                '20': 'कोष तथा लेखा नियन्त्रक महाशाखा, कल्याणकारी योजना निर्देशनालय',
                '81': 'गणेशदल गण',
                '8': 'गुणस्तर नियन्त्रण निर्देशनालय',
                '94': 'गुणस्तर नियन्त्रण निर्देशनालय',
                '67': 'चिकित्सा शास्त्र महाविद्यालय',
                '98': 'छापाखाना शाखा, श्रब्यदृश्य माहाशाखा, सै.ज.नि.',
                '76': 'जगदल गण',
                '60': 'जंगी असवाव खाना',
                '73': 'जनसम्पर्क तथा सूचना निर्देशनालय',
                '64': 'टु.मु. कोतखाना डिपो',
                '128': 'त्रिभुवन आर्मी अफिसर्स क्लब',
                '144': 'नं. १ युद्धकवच गुल्म',
                '134': 'नं. १ र्इन्टेलिजेन्स गण',
                '78': 'नं. १३ बाहिनी अड्रडा',
                '77': 'नं. १४ बाहिनी अड्रडा',
                '72': 'नं. १५ बाहिनी अड्डा',
                '79': 'नं. १६ बाहिनी अड्रडा',
                '71': 'नं. १७ बाहिनी अड्रडा',
                '143': 'नयाँ गोरख गण',
                '66': 'नर्सिङ्ग महाविद्यालय',
                '7': 'निरीक्षण निर्देशनालय',
                '51': 'निरीक्षण निर्देशनालय',
                '50': 'निरीक्षणाधिकृतको कार्यालय',
                '87': 'निवेश, कार्यरथी विभाग',
                '36': 'नीति तथा योजना निर्देशनालय',
                '43': 'नीति तथा योजना महाशाखा',
                '105': 'नेपाल क्याभलरी',
                '114': 'नेपाली सेना वार कलेज',
                '59': 'नेपाली सेना स्वास्थ्य विज्ञान संस्थान',
                '118': 'नेपाली सैनिक प्रतिष्ठान',
                '111': 'पशु बिकास तथा पशु चिकित्शा निर्देशनालय',
                '61': 'पश्चिम एयर बेश',
                '125': 'पश्चिम पृतना',
                '69': 'पि.एस्.ओज. कार्यालय',
                '120': 'पूर्वी पृतना',
                '9': 'प्रधान सेनापतिको कार्यालय',
                '47': 'प्रबन्धरथीको कार्यालय',
                '56': 'प्राप्ती उपशाखा, सैनिक सामग्री प्राप्ती निर्देशनालय',
                '44': 'फौज योजना महाशाखा',
                '53': 'बजेट निर्देशनालय',
                '28': 'बलाधिकृतको कार्यालय',
                '1': 'बलाध्यक्षको कार्यालय',
                '135': 'बहु व्यवसाय उद्योग पहिरन सामाग्री शाखा',
                '83': 'भर्ना छनौट निर्देशनालय',
                '109': 'भैरव बहान गुल्म',
                '126': 'मध्य पश्चिम पृतना',
                '122': 'मध्य पूर्वी पृतना',
                '123': 'मध्य पृतना',
                '99': 'मनोबैज्ञानिक कार्य महाशाखा',
                '3': 'मानव अधिकार निर्देशनालय',
                '55': 'मिन्हा मोजारा',
                '97': 'युद्ध कवच गण',
                '40': 'युद्धकार्य निर्देशनालय',
                '29': 'युद्धकार्य महानिर्देशनालय',
                '19': 'योजना तथा अनुसन्धान महाशाखा, कल्याणकारी योजना निर्देशनालय',
                '138': 'रक्षा मन्त्रालय',
                '41': 'रणनीतिक तथा दीर्घकालीन योजना निर्देशनालय',
                '75': 'राजदल गण',
                '14': 'राष्ट्रिय निकुञ्ज तथा बन्यजन्तु आरक्ष निर्देशनालय',
                '141': 'राष्ट्रिय सुरक्षा परिषदको सचिवालय',
                '62': 'राष्ट्रिय सेवा दल',
                '104': 'रेडियो तथा प्रकाशन शाखा',
                '132': 'र्इन्जिनियर तालिम शिक्षालय',
                '129': 'वन तथा पर्यावरण सुरक्षा निर्देशनालय',
                '101': 'विकास निर्माण कार्यदल गण १',
                '102': 'विकास निर्माण कार्यदल गण २',
                '54': 'विकास निर्माण निर्देशनालय',
                '65': 'विदेश उपशाखा, सैनिक सामग्री प्राप्ती निर्देशनालय',
                '38': 'विदेश तालिम शाखा, सैनिक तालिम निर्देशनालय, सै.ता. तथा ड.म‍.नि.',
                '96': 'विद्युत तथा यान्त्रिक शिक्षालय',
                '95': 'विद्युत तथा यान्त्रिक सेवा केन्द्र',
                '11': 'विद्युत तथा यान्त्रिक सेवा निर्देशनालय',
                '35': 'विपद व्यवस्थापन निर्देशनालय',
                '137': 'विशेष फौज बाहिनी',
                '25': 'वीरेन्द्र अस्पताल',
                '84': 'वेतन बृति तथा समारोह निर्देशनालय',
                '33': 'व्यवस्था, नीति तथा योजना महानिर्देशनालय',
                '34': 'शान्ति सेना संचालन निर्देशनालय',
                '110': 'शार्दुलजंग गुल्म',
                '6': 'शिकायत जाँच निर्देशनालय',
                '93': 'शिकायत जाँच निर्देशनालय',
                '17': 'शैक्षिक शाखा, कल्याणकारी योजना निर्देशनालय',
                '103': 'श्रब्यदृश्य शाखा',
                '48': 'संभाररथीको कार्यालय',
                '80': 'समन्वय शाखा, बलाधिकृतको कार्यालय',
                '10': 'समन्वय शाखा, बलाध्यक्षको कार्यालय',
                '100': 'समारोह श्रब्यदृश्य महाशाखा',
                '130': 'सर्भे महाशाखा',
                '4': 'सर्वोत्कृष्ट अभ्यास महाशाखा',
                '92': 'सर्वोत्कृष्ट अभ्यास महाशाखा',
                '2': 'सहायक बलाध्यक्षको कार्यालय',
                '142': 'साइबर सुरक्षा निर्देशनालय',
                '82': 'सिग्नल तालिम शिक्षालय',
                '121': 'सुदुर पश्चिम पृतना',
                '131': 'सुन्दरीजल आर्सनल कार्यालय',
                '140': 'सुरक्षा तथा समारोह व्यवस्था सचिवालय, उपराष्ट्रपतिको कार्यालय',
                '139': 'सुरक्षा तथा समारोह व्यवस्था सचिवालय, राष्ट्रपति भवन',
                '22': 'सूचना तथा प्रविधि महाशाखा',
                '5': 'सेना प्राड विवाक',
                '89': 'सैनिक अभिलेखालय',
                '70': 'सैनिक आर्थिक प्रशासन विभाग',
                '133': 'सैनिक आवसिय माध्यमिक बिद्यालय, भक्तपुर',
                '31': 'सैनिक इन्टेलिजेन्स महानिर्देशनालय',
                '39': 'सैनिक ईण्टेलिजेन्स बाहिनी',
                '117': 'सैनिक कमाण्ड तथा स्टाफ कलेज',
                '116': 'सैनिक केन्द्रिय पुस्तकालय',
                '63': 'सैनिक गठ्ठाघर डिपो',
                '30': 'सैनिक तालिम तथा डक्ट्रिन महानिर्देशनालय',
                '24': 'सैनिक पुनर्स्थापना केन्द्र',
                '88': 'सैनिक प्रहरी गण',
                '112': 'सैनिक बन्दोबस्ती शिक्षालय',
                '107': 'सैनिक ब्याण्ड',
                '42': 'सैनिक शारीरिक तालिम तथा खेलकुद केन्द्र',
                '52': 'सैनिक संगठन निर्देशनालय',
                '108': 'सैनिक संग्रहालय',
                '26': 'सैनिक सचिव  विभाग',
                '15': 'सैनिक सामग्री उत्पादन तथा यान्त्रिक सेवा महानिर्देशनालय',
                '16': 'सैनिक सामग्री उत्पादन निर्देशनालय',
                '13': 'सैनिक स्वास्थ्य महानिर्देशनालय',
                '32': 'सैनिक हवाई महानिर्देशनालय',
                '57': 'सैनिक हवाई महानिर्देशनालय,मर्मत तथा सम्भार',
                '85': 'स्थपति निर्देशनालय',
                '37': 'स्वदेश तालिम शाखा, सैनिक तालिम निर्देशनालय, सै.ता. तथा ड.म‍.नि.',
                '119': 'स्वास्थ्य सेवा व्यवस्थापन महाशाखा',
                '45': 'हतियार तथा उपकरण महाशाखा'
            }
            
            # Get the actual unit name from the code
            unit_name = unit_mapping.get(unit_id, '')
            
            # Get rank name from code
            rank_mapping = {
                '3': 'महारथी',
                '5': 'रथी',
                '7': 'उपरथी',
                '9': 'सहायक रथी',
                '10': 'आ.सहायक रथी',
                '11': 'महा सेनानी',
                '12': 'आ.महा सेनानी',
                '13': 'प्रमुख सेनानी',
                '14': 'आ. प्रमुख सेनानी',
                '15': 'सेनानी',
                '16': 'आ.सेनानी',
                '17': 'सह. सेनानी',
                '18': 'आ.सह सेनानी',
                '19': 'उप सेनानी',
                '21': 'सहायक सेनानी',
                '23': 'अधिकृत क्याडेट शुरु',
                '24': 'पदिक क्याडेट',
                '25': 'पदिक कर्मचारी क्याडेट',
                '26': 'इन्सर्भिस क्याडेट',
                '27': 'मानार्थ सह सेनानी',
                '29': 'मानार्थ उप सेनानी',
                '31': 'प्रमुख सुवेदार',
                '33': 'सिनियर सुवेदार',
                '35': 'सुबेदार',
                '37': 'जमदार',
                '39': 'गण कार्य हुद्बा',
                '41': 'गण प्रबन्ध हुद्बा',
                '43': 'गुल्म कार्य हुद्बा',
                '45': 'गुल्म प्रबन्ध हुद्बा',
                '47': 'हुद्बा',
                '49': 'अमल्दार',
                '51': 'प्युठ',
                '53': 'सिपाही',
                '54': 'एन.सि.ई.  पाँचौ स्तर',
                '55': 'सैन्य',
                '56': 'ए.सैन्य',
                '58': 'एन.सि.ई. चौथो स्तर',
                '59': 'कोते',
                '60': 'एन.सि.ई. तेस्रो स्तर',
                '62': 'एन.सि.ई. दोस्रो स्तर',
                '63': 'एन.सि.ई. प्रथम स्तर',
                '77': 'हुद्बा क्याडेट',
                '100': 'वरिष्ठ चार्टर्ड एकाउन्टेन्ट',
                '101': 'प्राड सहायक रथी',
                '102': 'शाखा अधिकृत',
                '103': 'सेनानी (अ.प्रा.)',
                '104': 'प्रमुख सेनानी (अ.प्रा.)',
                '105': 'नायव सुब्बा'
            }
            
            rank_name = rank_mapping.get(rank_code, '')
            
            # Create user with hashed password
            user = User.objects.create(
                username=username,
                email=email,
                password=make_password(password1),
                phone=phone,
                unit=unit_name,
                role=role,
                rank=rank_name,
                is_active=False
            )
            
            alert_type = 'success'
            alert_message = f'Account created successfully for {username}!'
            return render(request, 'login.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
            
        except Exception as e:
            alert_type = 'error'
            alert_message = f'Error creating account: {str(e)}'
            return render(request, 'signin.html', {
                'alert_type': alert_type,
                'alert_message': alert_message
            })
    
    
    # GET request - show registration form
    return render(request, 'signin.html')


def newsSource(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    search_query = request.GET.get('search', '').strip()
    
    sources = NewsSource.objects.all()
    
    if search_query:
        sources = sources.filter(
            name__icontains=search_query
        ) | sources.filter(
            url__icontains=search_query
        )
    
    sources = sources.order_by('name')
    
    # Pagination (12 per page – adjust as needed)
    paginator = Paginator(sources,15)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'newsSource.html', {
        'sources': page_obj,
        'search_query': search_query,
    })


@login_required
def keywordsAdd(request):
    if not check_access(request):
        return redirect('logout')  # Create this 
    
    """
    Add dangerous keywords with free-text category
    """
    
    # Get existing categories for suggestions (optional)
    existing_categories = DangerousKeyword.objects.values_list(
        'category', flat=True
    ).distinct().order_by('category')
    
    if request.method == 'POST':
        category = request.POST.get('category', '').strip()
        raw_words = request.POST.get('words', '').strip()

        # Validation: category is required
        if not category:
            messages.error(request, 'Please enter a category.')
            return render(request, 'keywords_add.html', {
                'form_data': request.POST,
                'existing_categories': existing_categories
            })

        # Validation: words are required
        if not raw_words:
            messages.error(request, 'Please enter at least one keyword.')
            return render(request, 'keywords_add.html', {
                'form_data': request.POST,
                'existing_categories': existing_categories
            })

        # Split and clean words
        word_list = re.split(r'[,;\n\s]+', raw_words)
        word_list = [w.strip() for w in word_list if w.strip()]

        if not word_list:
            messages.error(request, 'No valid keywords found.')
            return render(request, 'keywords_add.html', {
                'form_data': request.POST,
                'existing_categories': existing_categories
            })

        # Save keywords with current user
        added = 0
        duplicates = 0
        errors = 0
        
        for word in word_list:
            try:
                # Check if keyword already exists in this category
                exists = DangerousKeyword.objects.filter(
                    word=word.lower(),
                    category=category
                ).exists()
                
                if exists:
                    duplicates += 1
                    continue
                
                # Create new keyword
                DangerousKeyword.objects.create(
                    word=word,
                    category=category,
                    created_by=request.user,
                    is_active=True
                )
                added += 1
                
            except Exception as e:
                errors += 1
                print(f"Error saving keyword '{word}': {e}")

        # Success message
        if added > 0:
            success_msg = f"✅ Successfully added {added} keyword(s)"
            if duplicates:
                success_msg += f" (Skipped {duplicates} duplicate(s))"
            if errors:
                success_msg += f" ({errors} error(s) occurred)"
            messages.success(request, success_msg)
        else:
            if duplicates:
                messages.warning(request, f"All keywords already exist in this category. Skipped {duplicates} duplicate(s).")
            else:
                messages.error(request, "No keywords were added. Please check your input.")

        return redirect('keywords_add')  # Redirect to clear form

    # GET request
    return render(request, 'keywords_add.html', {
        'existing_categories': existing_categories
    })


@login_required
def CategoryAdd(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    View to add a new threat category
    """
    
    if request.method == 'POST':
        # Get form data
        category_name = request.POST.get('name', '').strip()
        is_active = request.POST.get('is_active', 'off') == 'on'
        
        # Basic validation
        if not category_name:
            messages.error(request, "Category name is required.")
            return render(request, 'keywords_add.html', {})
        
        # Check if category already exists for the current user
        # Allow duplicates only if created by different users
        existing_category = ThreatCategory.objects.filter(
            name__iexact=category_name,
            created_by=request.user  # Only check categories created by this user
        ).exists()
        
        if existing_category:
            messages.error(request, f"You already have a category named '{category_name}'.")
            return render(request, 'category_add.html', {})
        
        try:
            # Create new category
            new_category = ThreatCategory.objects.create(
                name=category_name,
                created_by=request.user,  # Set the logged-in user as creator
                is_active=is_active
            )
            
            messages.success(request, f"Category '{new_category.name}' added successfully!")
            return redirect('category_add')  # Redirect to appropriate page
            
        except Exception as e:
            messages.error(request, f"Error creating category: {str(e)}")
    
    # GET request - render empty form
    context = {
        'action': 'Add Category',
        'form_title': 'Create New Threat Category',
    }
    return render(request, 'category_add.html', context)


# This function loads dangerous keywords by category from the database once, caches them for 5 minutes, and reuses the cache to speed up comment scanning.
def get_keyword_map():
    keyword_map = cache.get('dangerous_keyword_map')
    if keyword_map is None:
        keyword_map = defaultdict(list)
        for kw in DangerousKeyword.objects.all():
            keyword_map[kw.category].append(kw.word)
        cache.set('dangerous_keyword_map', dict(keyword_map), timeout=300)
    return keyword_map

def is_dangerous_comment(text):
    if not text:
        return {
            'is_dangerous': False,
            'matches': [],
            'danger_score': 0,
            'highlighted_comment': mark_safe(escape('').replace('\n', '<br>'))
        }

    text_lower = text.lower()
    matches = []
    danger_score = 0
    keyword_map = get_keyword_map()
    matched_words = set()

    # Check core categories (violence, threats, dehumanizing)
    for category in ['violence', 'threats', 'dehumanizing']:
        for word in keyword_map.get(category, []):
            if word in text_lower:
                matches.append((category, word))
                matched_words.add(word)
                danger_score += 1

    # Check military category
    military_words = keyword_map.get('military', [])
    mobilization_words = keyword_map.get('mobilization', [])
    
    military_found = False
    for word in military_words:
        if word in text_lower:
            matches.append(('military', word))
            matched_words.add(word)
            danger_score += 1
            military_found = True

    # Check mobilization category
    mobilization_found = False
    for word in mobilization_words:
        if word in text_lower:
            matches.append(('mobilization', word))
            matched_words.add(word)
            danger_score += 1
            mobilization_found = True

    # HIGH RISK: Military + Violence/Threats combination
    # If military terms are used with violence/threats, increase danger significantly
    if military_found:
        for cat in ['violence', 'threats']:
            for word in keyword_map.get(cat, []):
                if word in text_lower:
                    matches.append(('contextual_threat', f'सेना + {word}'))
                    matched_words.add(word)
                    danger_score += 2  # Higher score for military + violence
                    break

    # HIGH RISK: Military + Mobilization combination
    # If military terms are used with mobilization/call to action
    if military_found and mobilization_found:
        matches.append(('contextual_threat', 'सेना + आह्वान'))
        danger_score += 3  # Highest risk - military mobilization
        # Add all military and mobilization words to matched words
        for word in military_words:
            if word in text_lower:
                matched_words.add(word)
        for word in mobilization_words:
            if word in text_lower:
                matched_words.add(word)

    # Highlight matched words
    safe_text = escape(text)
    for word in sorted(matched_words, key=len, reverse=True):
        escaped_word = re.escape(word)
        safe_text = re.sub(
            f'({escaped_word})',
            r'<mark class="highlight-danger">\1</mark>',
            safe_text,
            flags=re.IGNORECASE
        )

    highlighted = mark_safe(safe_text.replace('\n', '<br>'))
    
    # Adjusted danger threshold - military content is more sensitive
    is_dangerous = danger_score >= 1  # Lower threshold for military content
    
    return {
        'is_dangerous': is_dangerous,
        'matches': matches,
        'danger_score': danger_score,
        'highlighted_comment': highlighted
    }


import json
import re

def parse_comments_from_text(text):
    """
    Parse comments from either JSON format or the original text format.
    Supports both:
    1. JSON format with 'commenter' and 'text' fields
    2. Original text format with timestamps like '2d', '3h'
    """
    comments = []
    
    # First try to parse as JSON
    try:
        data = json.loads(text)
        
        # If it's a list of comments
        if isinstance(data, list):
            for item in data:
                comment_text = item.get('text', '')
                if comment_text:  # Only process if there's actual text
                    danger_info = is_dangerous_comment(comment_text)
                    sentiment = predict_sentiment(comment_text)
                    
                    comments.append({
                        'author': item.get('commenter', 'Unknown'),
                        'comment': comment_text,
                        'timestamp': item.get('timestamp', ''),
                        'sentiment': sentiment,
                        **danger_info
                    })
            return comments
    except (json.JSONDecodeError, TypeError):
        # If JSON parsing fails, fall back to original text format
        pass
    
    # Original text format parsing (unchanged)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    buffer = []
    time_pattern = re.compile(r'\d+[dhm]$')

    for line in lines:
        if time_pattern.match(line):
            if buffer:
                author = buffer[0]
                comment = '\n'.join(buffer[1:]).strip() if len(buffer) > 1 else ''

                danger_info = is_dangerous_comment(comment)
                sentiment = predict_sentiment(comment)

                comments.append({
                    'author': author,
                    'comment': comment,
                    'timestamp': line,
                    'sentiment': sentiment,
                    **danger_info
                })
            buffer = []
        elif line not in {'Reply', 'Edited'}:
            buffer.append(line)
    
    # Don't forget to process the last comment if buffer is not empty
    if buffer:
        author = buffer[0]
        comment = '\n'.join(buffer[1:]).strip() if len(buffer) > 1 else ''

        danger_info = is_dangerous_comment(comment)
        sentiment = predict_sentiment(comment)

        comments.append({
            'author': author,
            'comment': comment,
            'timestamp': '',  # No timestamp for last comment
            'sentiment': sentiment,
            **danger_info
        })

    return comments

def prepare_chart_data(comments):
    """Prepare simplified chart data that's guaranteed to work"""
    if not comments:
        return {
            'total_comments': 0,
            'suspicious_comments': 0,
            'clean_comments': 0,
            'suspicious_rate': 0,

            # Corrected sentiment counts (typo fixes)
            'sentiment_counts': {
                'positive': 0,
                'negative': 0,
                'neutral': 0  # Fixed spelling
            },

            'category_counts': {
                'violence': 0,
                'threats': 0,
                'dehumanizing': 0,
                'contextual_threat': 0
            },
            'score_distribution': {
                '0': 0, '1': 0, '2': 0, '3': 0, '4': 0, '5_plus': 0  # CHANGED: '5+' to '5_plus'
            },
            'comments_data': []
        }

    total_comments = len(comments)
    suspicious_comments = sum(1 for c in comments if c.get('is_dangerous', False))
    clean_comments = total_comments - suspicious_comments
    suspicious_rate = round((suspicious_comments / total_comments * 100), 1) if total_comments > 0 else 0

    # Category breakdown
    category_counts = {
        'violence': 0,
        'threats': 0,
        'dehumanizing': 0,
        'contextual_threat': 0
    }

    # Danger score distribution - CHANGED: '5+' to '5_plus'
    score_distribution = {'0': 0, '1': 0, '2': 0, '3': 0, '4': 0, '5_plus': 0}

    # Corrected sentiment counts
    sentiment_counts = {
        'positive': 0,
        'negative': 0,
        'neutral': 0  # Fixed spelling
    }

    for comment in comments:
        # Count categories
        for match in comment.get('matches', []):
            # Handle both list/tuple formats and string formats
            if isinstance(match, (list, tuple)) and len(match) > 0:
                category = str(match[0]).lower().strip()
            elif isinstance(match, str):
                category = match.lower().strip()
            else:
                continue
            
            # Map common variations to standard categories
            category_mapping = {
                'violent': 'violence',
                'violence': 'violence',
                'threat': 'threats',
                'threats': 'threats',
                'threatening': 'threats',
                'dehumanizing': 'dehumanizing',
                'dehumanize': 'dehumanizing',
                'contextual': 'contextual_threat',
                'contextual_threat': 'contextual_threat',
                'context': 'contextual_threat'
            }
            
            if category in category_mapping:
                mapped_category = category_mapping[category]
                category_counts[mapped_category] += 1

        # Count danger scores - CHANGED: '5+' to '5_plus'
        score = comment.get('danger_score', 0)
        if score >= 5:
            score_distribution['5_plus'] += 1  # CHANGED HERE
        else:
            score_distribution[str(score)] += 1

        # Count sentiment - handle different possible formats
        sentiment = comment.get('sentiment', 'neutral')
        
        # Convert to lowercase string for consistency
        if isinstance(sentiment, int):
            # Handle numeric sentiment: 0=negative, 1=positive, 2=neutral
            sentiment_map = {0: 'negative', 1: 'positive', 2: 'neutral'}
            sentiment = sentiment_map.get(sentiment, 'neutral')
        elif isinstance(sentiment, str):
            sentiment = sentiment.lower().strip()
            # Fix common typos
            if sentiment in ['positve', 'pos']:
                sentiment = 'positive'
            elif sentiment in ['neg', 'negatve']:
                sentiment = 'negative'
            elif sentiment in ['neut', 'neural', 'nutral']:  # Fix for "neural" typo
                sentiment = 'neutral'
        else:
            sentiment = 'neutral'
        
        # Count in sentiment_counts
        if sentiment in sentiment_counts:
            sentiment_counts[sentiment] += 1
        else:
            sentiment_counts['neutral'] += 1

    # Prepare comments data
    comments_data = []
    for comment in comments:
        # Normalize sentiment for display
        display_sentiment = comment.get('sentiment', 'neutral')
        if isinstance(display_sentiment, int):
            sentiment_map = {0: 'negative', 1: 'positive', 2: 'neutral'}
            display_sentiment = sentiment_map.get(display_sentiment, 'neutral')
        elif isinstance(display_sentiment, str):
            display_sentiment = display_sentiment.lower().strip()
            if display_sentiment in ['positve', 'pos']:
                display_sentiment = 'positive'
            elif display_sentiment in ['neg', 'negatve']:
                display_sentiment = 'negative'
            elif display_sentiment in ['neut', 'neural', 'nutral']:
                display_sentiment = 'neutral'
        
        # Process matches safely
        processed_matches = []
        for match in comment.get('matches', []):
            if isinstance(match, (list, tuple)):
                if len(match) >= 2:
                    processed_matches.append([str(match[0]), str(match[1])])
                elif len(match) == 1:
                    processed_matches.append([str(match[0]), ''])
            elif isinstance(match, str):
                processed_matches.append([match, ''])
        
        comments_data.append({
            'author': str(comment.get('author', 'Unknown')),
            'timestamp': str(comment.get('timestamp', 'Unknown')),
            'sentiment': display_sentiment,  # Use normalized sentiment
            'is_dangerous': bool(comment.get('is_dangerous', False)),
            'danger_score': int(comment.get('danger_score', 0)),
            'highlighted_comment': str(comment.get('highlighted_comment', '')),
            'matches': processed_matches
        })

    return {
        'total_comments': total_comments,
        'suspicious_comments': suspicious_comments,
        'clean_comments': clean_comments,
        'suspicious_rate': suspicious_rate,
        'sentiment_counts': sentiment_counts,
        'category_counts': category_counts,
        'score_distribution': score_distribution,
        'comments_data': comments_data
    }


# def commentAnalyze(request):
#     comments = []
#     uploaded = False
#     error = None
#     chart_data = {}

#     if request.method == 'POST' and request.FILES.get('txt_file'):
#         uploaded_file = request.FILES['txt_file']
#         if not uploaded_file.name.endswith('.txt'):
#             error = 'Please upload a .txt file.'
#             uploaded = True
#         else:
#             try:
#                 text = uploaded_file.read().decode('utf-8')
#                 comments = parse_comments_from_text(text) 
#                 print(111111111111, comments) # Your existing function
#                 uploaded = True
                
#                 # Prepare chart data
#                 chart_data = prepare_chart_data(comments)
                
#             except Exception as e:
#                 error = f'Error processing file: {str(e)}'
#                 uploaded = True

#     return render(request, 'comment_analyze.html', {
#         'comments': comments,
#         'uploaded': uploaded,
#         'error': error,
#         'chart_data': chart_data,
#     })
def commentAnalyze(request):
    if not check_access(request):
        return redirect('logout')  # Create this view

    comments = []
    uploaded = False
    error = None
    chart_data = {}

    if request.method == 'POST' and request.FILES.get('txt_file'):
        uploaded_file = request.FILES['txt_file']
        file_name = uploaded_file.name.lower()
        
        # Check file extension - accept both .txt and .json
        if not (file_name.endswith('.txt') or file_name.endswith('.json')):
            error = 'Please upload a .txt or .json file.'
            uploaded = True
        else:
            try:
                # Read file with UTF-8 encoding
                text = uploaded_file.read().decode('utf-8')
                
                # Parse comments using the universal function
                comments = parse_comments_from_text(text)
                uploaded = True
                
                # Prepare chart data
                chart_data = prepare_chart_data(comments)
                
            except UnicodeDecodeError:
                # Try other encodings if UTF-8 fails
                try:
                    uploaded_file.seek(0)  # Reset file pointer
                    text = uploaded_file.read().decode('latin-1')
                    comments = parse_comments_from_text(text)
                    uploaded = True
                    chart_data = prepare_chart_data(comments)
                except Exception as e:
                    error = f'File encoding error. Please use UTF-8 encoded files. Details: {str(e)}'
                    uploaded = True
            except json.JSONDecodeError as e:
                # Handle specific JSON errors for .json files
                if file_name.endswith('.json'):
                    error = f'Invalid JSON format in file. Error: {str(e)}'
                else:
                    error = f'Error parsing file: {str(e)}'
                uploaded = True
            except Exception as e:
                error = f'Error processing file: {str(e)}'
                uploaded = True

    return render(request, 'comment_analyze.html', {
        'comments': comments,
        'uploaded': uploaded,
        'error': error,
        'chart_data': chart_data,
    })
   
def newsAutofeeding(request=None):
    """
    Main function to fetch news from multiple sources and save to database
    """
    if not check_access(request):
        return redirect('logout')  # Create this view

    try:
        all_articles = []
        
        # Source 1: Kantipur News
        try:
            json_data = kantipur_to_json()
            data = json.loads(json_data)
            if data["metadata"]["status"] == "success":
                for article in data["articles"]:
                    article["source"] = "kantipur"  # Ensure source is set
                    all_articles.append(article)
                print(f"✅ Fetched {len(data['articles'])} articles from Kantipur")
            else:
                print(f"❌ Kantipur scraping failed: {data['metadata'].get('error', 'Unknown error')}")
        except Exception as e:
            print(f"❌ Error fetching Kantipur news: {e}")

        # Source 2: Techpana News
        try:
            # Call the Techpana function directly - it returns JSON string
            json_data = techpana_to_json()
            data = json.loads(json_data)
            
            if data["metadata"]["status"] == "success":
                for article in data["articles"]:
                    article["source"] = "techpana"  # Ensure source is set
                    all_articles.append(article)
                print(f"✅ Fetched {len(data['articles'])} articles from Techpana")
                
                # Print some statistics if available - UPDATED FIELD NAMES
                if "new_articles_added" in data["metadata"]:
                    print(f"📊 Techpana Stats: {data['metadata']['new_articles_added']} new articles")
                if "total_articles_found" in data["metadata"]:
                    print(f"📊 Techpana Total Found: {data['metadata']['total_articles_found']} articles on homepage")
                if "articles_skipped_no_keywords" in data["metadata"]:
                    print(f"📊 Techpana Skipped (no keywords): {data['metadata']['articles_skipped_no_keywords']} articles")
                if "articles_skipped_old" in data["metadata"]:
                    print(f"📊 Techpana Skipped (old): {data['metadata']['articles_skipped_old']} articles")
                
                # Content statistics
                if "content_statistics" in data["metadata"]:
                    content_stats = data["metadata"]["content_statistics"]
                    print(f"📝 Techpana Content Stats:")
                    print(f"   - With full content: {content_stats.get('articles_with_full_content', 0)}")
                    print(f"   - Preview only: {content_stats.get('articles_with_preview_only', 0)}")
                    print(f"   - Avg summary length: {content_stats.get('average_summary_length', 0)} chars")
                    print(f"   - Adequate content: {content_stats.get('articles_with_adequate_content', 0)}")
                
            else:
                print(f"❌ Techpana scraping failed: {data['metadata'].get('error', 'Unknown error')}")
        except Exception as e:
            print(f"❌ Error fetching Techpana news: {e}")
            import traceback
            print(f"❌ Techpana error details: {traceback.format_exc()}")

        # Source 3: Kathmandu Post News
        try:
            kathmandu_articles = kathmandu_post_extractor()
            for article in kathmandu_articles:
                # Convert Kathmandu Post format to match Kantipur format
                converted_article = {
                    "id": len(all_articles) + 1,
                    "title": article["title"],
                    "summary": article["summary"],
                    "url": article["news_link"],
                    "image_url": article["image_link"] or "",
                    "date": article["raw_date"],
                    "source": "kathmandu_post",
                    "threat_analysis": article.get("threat_analysis", {
                        "level": "low",
                        "keywords_found": [],
                        "categories": [],
                        "total_keywords_matched": 0
                    }),
                    "content_length": len(article["title"] + " " + article["summary"]),
                    "priority": article.get("threat_analysis", {}).get("priority", "normal")
                }
                all_articles.append(converted_article)
            print(f"✅ Fetched {len(kathmandu_articles)} articles from Kathmandu Post")
        except Exception as e:
            print(f"❌ Error fetching Kathmandu Post news: {e}")
        
        if not all_articles:
            error_msg = "❌ No articles fetched from any source"
            print(error_msg)
            return JsonResponse({
                "status": "error",
                "alert_type": "error",
                "alert_message": "No articles could be fetched from any news source.",
            })
        
        articles_saved = 0
        articles_updated = 0
        articles_skipped = 0
        
        # Process each article and save to database
        for article_data in all_articles:
            try:
                # Prepare data for saving
                threat_analysis = article_data.get("threat_analysis", {})
                
                # Convert lists to comma-separated strings
                keywords_found = threat_analysis.get("keywords_found", [])
                categories_list = threat_analysis.get("categories", [])
                
                keywords_str = ", ".join(keywords_found) if isinstance(keywords_found, list) else str(keywords_found)
                categories_str = ", ".join(categories_list) if isinstance(categories_list, list) else str(categories_list)
                
                # Handle date field - ensure it's a string
                article_date = article_data.get("date", "")
                if not article_date or article_date == "Unknown":
                    article_date = datetime.now().strftime("%Y-%m-%d")
                
                # Create or update article
                obj, created = AutoNewsArticle.objects.update_or_create(
                    url=article_data["url"],
                    defaults={
                        'title': article_data["title"][:500],  # Limit length for database
                        'summary': article_data.get("summary", "")[:1000],  # Limit length
                        'image_url': article_data.get("image_url", "")[:500],
                        'source': article_data.get("source", "unknown"),
                        'date': article_date,
                        'content_length': article_data.get("content_length", 0),
                        'priority': article_data.get("priority", "medium"),
                        'threat_level': threat_analysis.get("level", "low"),
                        'keywords': keywords_str[:500],  # Limit length
                        'categories': categories_str[:500],  # Limit length
                    }
                )
                
                if created:
                    articles_saved += 1
                    print(f"✅ Saved {article_data.get('source', 'unknown')}: {article_data['title'][:50]}...")
                else:
                    articles_updated += 1
                    print(f"🔄 Updated {article_data.get('source', 'unknown')}: {article_data['title'][:50]}...")
                    
            except Exception as e:
                articles_skipped += 1
                source = article_data.get('source', 'unknown')
                print(f"❌ Error saving {source} article '{article_data.get('title', 'Unknown')[:30]}...': {e}")
                continue
        
        # Create success message with all sources
        sources_used = []
        if any(article['source'] == 'kantipur' for article in all_articles):
            sources_used.append("Kantipur")
        if any(article['source'] == 'techpana' for article in all_articles):
            sources_used.append("Techpana")
        if any(article['source'] == 'kathmandu_post' for article in all_articles):
            sources_used.append("Kathmandu Post")
            
        sources_str = ", ".join(sources_used) if sources_used else "No sources"
        
        success_message = f"News autofeeding completed! Sources: {sources_str}. Saved: {articles_saved}, Updated: {articles_updated}, Skipped: {articles_skipped}"
        print(f"✅ {success_message}")
        
        # Return appropriate response
        if request and hasattr(request, 'is_ajax') and request.is_ajax():
            return JsonResponse({
                "status": "success",
                "alert_type": "success", 
                "alert_message": success_message,
                "stats": {
                    "saved": articles_saved,
                    "updated": articles_updated,
                    "skipped": articles_skipped,
                    "total_processed": len(all_articles),
                    "sources": sources_used
                }
            })
        else:
            return HttpResponseRedirect(
                reverse('autonews_view') + 
                f'?alert=success&message={success_message}&saved={articles_saved}&updated={articles_updated}'
            )
    
    except Exception as e:
        error_msg = f"❌ Autofeeding failed: {str(e)}"
        print(error_msg)
        import traceback
        print(f"❌ Full error: {traceback.format_exc()}")
        
        if request and hasattr(request, 'is_ajax') and request.is_ajax():
            return JsonResponse({
                "status": "error",
                "alert_type": "error",
                "alert_message": "News autofeeding failed. Please try again later.",
            })
        else:
            return HttpResponseRedirect(
                reverse('autonews_view') + 
                f'?alert=error&message=News autofeeding failed. Please try again later.'
            )


def keyboard_fetch(request):
    """Render the keyboard fetch page"""
    return render(request, 'newsportal_list.html')


from django.contrib.auth import get_user_model

User = get_user_model()

import json
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.shortcuts import HttpResponseRedirect
from django.urls import reverse
from datetime import datetime
from .models import AutoNewsArticle  # Your model

# Import your scraper functions
from collect.scrapers.keyboard_techpana import keyboard_techpana_to_json
from collect.scrapers.keyboard_nagarik import keyboard_nagariknews_to_json

def send_to_websocket(message):
    """Utility function to send messages to websocket"""
    # Implement your websocket sending logic here
    # For now, we'll just print
    print(f"📡 [WebSocket]: {message}")

@login_required
def keyboard_AutoFeed(request):
    """Auto-feed news from different sources"""
    if not check_access(request):
        return redirect('logout')  # Create this view

    try:
        current_user = request.user
        
        # Get requested source from AJAX or form
        source = None
        if request.method == 'POST':
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                try:
                    data = json.loads(request.body)
                    source = data.get('source')
                except:
                    source = request.POST.get('source')
            else:
                source = request.POST.get('source')
        
        send_to_websocket(f"Fetching started from { source}")
        
        # Define available scrapers
        SCRAPER_FUNCTIONS = {
            'nagariknews': keyboard_nagariknews_to_json,
            'techpana': keyboard_techpana_to_json,
            'kantipur': keyboard_kantipur_to_json,
            'kathmandupost': keyboard_kathmandu_post_to_json,
            'onlinekhabar': keyboard_onlinekhabar_to_json,
            'paschimnepal': keyboard_paschimnepal_to_json,
            'onlinetvnepal': keyboard_onlinetvnepal_to_json,
            'osnepal': keyboard_osnepal_to_json, 
            'eadarsha': keyboard_eadarsha_to_json,
            'arthasarokar': keyboard_arthasarokar_to_json,
            'newsofnepal': keyboard_newsofnepal_to_json,
            'rajdhanidaily': keyboard_rajdhanidaily_to_json,
            # Add more sources here as you create their scraper functions
        }
        
        # Determine which sources to scrape
        if source in SCRAPER_FUNCTIONS:
            sources_to_scrape = [source]
            print(333333333333333333333333, sources_to_scrape)
        elif source == 'all' or source is None:
            sources_to_scrape = list(SCRAPER_FUNCTIONS.keys())
        else:
            return JsonResponse({
                "status": "error",
                "message": f"Invalid source: {source}. Available: {', '.join(SCRAPER_FUNCTIONS.keys())}"
            })
        
        total_saved = 0
        source_stats = {}
        
        # Scrape each source
        for source_name in sources_to_scrape:
            try:
                send_to_websocket(f"📰 Fetching from {source_name}...")
                
                # Get the scraper function
                scraper_function = SCRAPER_FUNCTIONS[source_name]
                
                # Call the scraper
                json_result = scraper_function(request)
                
                # Parse result
                if isinstance(json_result, str):
                    data = json.loads(json_result)
                elif hasattr(json_result, 'content'):  # JsonResponse
                    data = json.loads(json_result.content.decode('utf-8'))
                else:
                    data = json_result
                
                # Check if scraping was successful
                if data.get("metadata", {}).get("status") != "success":
                    send_to_websocket(f"❌ {source_name}: Scraping failed")
                    source_stats[source_name] = {"error": "Scraping failed"}
                    continue
                
                articles = data.get("articles", [])
                saved_from_source = 0
                
                # Save each article
                for article in articles:
                    try:
                        # Check for duplicate (by URL for this user)
                        exists = AutoNewsArticle.objects.filter(
                            url=article['url'],
                            created_by=current_user
                        ).exists()
                        
                        if exists:
                            continue
                        
                        # Prepare keywords and categories
                        keywords_list = article.get('keywords', [])
                        if not keywords_list:
                            # Try to get from threat_analysis for backward compatibility
                            keywords_list = article.get('threat_analysis', {}).get('keywords_found', [])
                        
                        categories_list = article.get('categories', [])
                        if not categories_list:
                            # Try to get from threat_analysis for backward compatibility
                            categories_list = article.get('threat_analysis', {}).get('categories', [])
                        
                        # Convert lists to comma-separated strings
                        keywords_str = ','.join(keywords_list) if isinstance(keywords_list, list) else str(keywords_list)
                        categories_str = ','.join(categories_list) if isinstance(categories_list, list) else str(categories_list)
                        
                        # Determine threat level and priority
                        threat_level = article.get('threat_analysis', {}).get('level', 'low')
                        priority = 'medium'
                        if threat_level == 'high':
                            priority = 'high'
                        elif threat_level == 'low':
                            priority = 'low'
                        
                        # Create article
                        AutoNewsArticle.objects.create(
                            title=article['title'][:200],
                            summary=article.get('summary', article.get('content', '')[:300])[:500],
                            url=article['url'],
                            image_url=article.get('image_url', '')[:500],
                            source=source_name,  # Use actual source name
                            date=article.get('date', datetime.now().strftime('%Y-%m-%d')),
                            keywords=keywords_str[:200],
                            categories=categories_str[:200],
                            threat_level=threat_level,
                            priority=priority,
                            content_length=len(article.get('content', '')),
                            created_by=current_user
                        )
                        
                        saved_from_source += 1
                        total_saved += 1
                        
                        if saved_from_source % 5 == 0:
                            send_to_websocket(f"✅ {source_name}: Saved {saved_from_source} articles")
                            
                    except Exception as e:
                        # Skip article if error
                        continue
                
                source_stats[source_name] = {
                    "total_found": len(articles),
                    "saved": saved_from_source,
                    "status": "success"
                }
                
                send_to_websocket(f"✅ {source_name}: Saved {saved_from_source}/{len(articles)} articles")
                
            except Exception as e:
                send_to_websocket(f"❌ {source_name} error: {str(e)[:50]}")
                source_stats[source_name] = {
                    "error": str(e)[:100],
                    "status": "failed"
                }
                continue
        
        # Final summary
        send_to_websocket(f"💾 Total saved: {total_saved} articles")
        
        for source_name, stats in source_stats.items():
            if stats.get('status') == 'success':
                send_to_websocket(f"   📰 {source_name}: {stats.get('saved', 0)} saved")
            else:
                send_to_websocket(f"   ❌ {source_name}: Failed - {stats.get('error', 'Unknown error')}")
        
        send_to_websocket("=" * 40)
        
        # Return response
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                "status": "success",
                "message": f"Auto-feed completed. Saved {total_saved} articles.",
                "saved_count": total_saved,
                "source_stats": source_stats,
                "user": current_user.username
            })
        else:
            # For form submission
            messages.success(request, f"Auto-feed completed. Saved {total_saved} articles.")
            return redirect('autonews_view')  # Redirect to your news view
        
    except Exception as e:
        send_to_websocket(f"❌ Critical error in auto-feed: {str(e)}")
        
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                "status": "error",
                "message": f"Auto-feed failed: {str(e)[:100]}"
            })
        else:
            messages.error(request, f"Auto-feed failed: {str(e)[:100]}")
            return redirect('autonews_view')


@login_required
def autoNews(request):
    if not check_access(request):
        return redirect('logout')  # Create this view

    current_user = request.user
    
    # Get filter parameters
    threat_level_filter = request.GET.get('threat_level', '')
    priority_filter = request.GET.get('priority', '')
    date_filter = request.GET.get('date_filter', '')
    search_query = request.GET.get('search', '')
    

    # Start with articles created by current user only
    current_user_unit = request.user.unit
    articles = AutoNewsArticle.objects.filter(created_by__unit=current_user_unit)
    
    # Apply filters
    if threat_level_filter:
        articles = articles.filter(threat_level=threat_level_filter)
    
    if priority_filter:
        articles = articles.filter(priority=priority_filter)
    
    if date_filter:
        if date_filter == 'today':
            articles = articles.filter(created_at__date=timezone.now().date())
        elif date_filter == 'yesterday':
            yesterday = timezone.now().date() - timedelta(days=1)
            articles = articles.filter(created_at__date=yesterday)
        elif date_filter == 'week':
            week_ago = timezone.now().date() - timedelta(days=7)
            articles = articles.filter(created_at__date__gte=week_ago)
        elif date_filter == 'month':
            month_ago = timezone.now().date() - timedelta(days=30)
            articles = articles.filter(created_at__date__gte=month_ago)
    
    if search_query:
        articles = articles.filter(
            Q(title__icontains=search_query) |
            Q(summary__icontains=search_query) |
            Q(keywords__icontains=search_query) |
            Q(source__icontains=search_query)
        )
    
    # Order by creation date (newest first)
    articles = articles.order_by('-created_at')
    
    # Get all articles for statistics - ONLY FOR CURRENT USER
    all_articles_user = AutoNewsArticle.objects.filter(created_by__unit=current_user_unit)
    
    # Apply same filters to statistics if they exist
    stats_articles = all_articles_user
    
    # Apply the same threat level filter to stats if set
    if threat_level_filter:
        stats_articles = stats_articles.filter(threat_level=threat_level_filter)
    
    # Apply the same priority filter to stats if set
    if priority_filter:
        stats_articles = stats_articles.filter(priority=priority_filter)
    
    # Apply the same date filter to stats if set
    if date_filter:
        if date_filter == 'today':
            stats_articles = stats_articles.filter(created_at__date=timezone.now().date())
        elif date_filter == 'yesterday':
            yesterday = timezone.now().date() - timedelta(days=1)
            stats_articles = stats_articles.filter(created_at__date=yesterday)
        elif date_filter == 'week':
            week_ago = timezone.now().date() - timedelta(days=7)
            stats_articles = stats_articles.filter(created_at__date__gte=week_ago)
        elif date_filter == 'month':
            month_ago = timezone.now().date() - timedelta(days=30)
            stats_articles = stats_articles.filter(created_at__date__gte=month_ago)
    
    # Get articles by threat level for stats - ONLY FOR CURRENT USER WITH FILTERS
    critical_articles = stats_articles.filter(threat_level='critical')
    high_articles = stats_articles.filter(threat_level='high') 
    medium_articles = stats_articles.filter(threat_level='medium')
    low_articles = stats_articles.filter(threat_level='low')
    
    # Get articles by priority for stats - ONLY FOR CURRENT USER WITH FILTERS
    high_priority_articles = stats_articles.filter(priority='high')
    medium_priority_articles = stats_articles.filter(priority='medium')
    low_priority_articles = stats_articles.filter(priority='low')
    
    # Add pagination
    paginator = Paginator(articles, 15)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'auto_news.html', {
        'articles': page_obj,
        'total_articles': stats_articles.count(),  # Use filtered stats
        'critical_count': critical_articles.count(),
        'high_count': high_articles.count(),
        'medium_count': medium_articles.count(),
        'low_count': low_articles.count(),
        'high_priority_count': high_priority_articles.count(),
        'medium_priority_count': medium_priority_articles.count(),
        'low_priority_count': low_priority_articles.count(),
        'threat_level_filter': threat_level_filter,
        'priority_filter': priority_filter,
        'date_filter': date_filter,
        'search_query': search_query,
        'page_obj': page_obj,
        'current_user': current_user,
        'showing_user_only': True,  # Flag to indicate showing only user's data
    })

def visualizationMap(request):
    if not check_access(request):
        return redirect('logout')  # Create this view

    context = {
        'STATIC_URL': settings.STATIC_URL,
    }
    return render(request, 'map.html', context)

def create_marker(request):
    """Create a new marker"""
    try:
        # Parse JSON data
        data = json.loads(request.body)
        print("Received data:", data)  # Debug log
        
        # Validate required fields
        required_fields = ['title', 'latitude', 'longitude']
        for field in required_fields:
            if field not in data or data[field] is None or str(data[field]).strip() == '':
                return JsonResponse({
                    'success': False,
                    'error': f'{field} is required'
                }, status=400)
        
        # Validate coordinate ranges
        try:
            latitude = float(data['latitude'])
            longitude = float(data['longitude'])
            
            if not (-90 <= latitude <= 90):
                return JsonResponse({
                    'success': False,
                    'error': 'Latitude must be between -90 and 90'
                }, status=400)
                
            if not (-180 <= longitude <= 180):
                return JsonResponse({
                    'success': False,
                    'error': 'Longitude must be between -180 and 180'
                }, status=400)
        except ValueError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid coordinate format'
            }, status=400)
        
        # Create marker with transaction for safety
        with transaction.atomic():
            marker = MapMarker.objects.create(
                title=data['title'].strip(),
                description=data.get('description', '').strip(),
                category=data.get('category', 'general'),
                color=data.get('color', '#FF0000'),
                latitude=latitude,
                longitude=longitude,
                created_by=request.user if request.user.is_authenticated else None
            )
        
        # Return success response
        return JsonResponse({
            'success': True,
            'message': 'Marker created successfully',
            'marker': {
                'id': marker.id,
                'title': marker.title,
                'description': marker.description,
                'category': marker.category,
                'color': marker.color,
                'latitude': float(marker.latitude),
                'longitude': float(marker.longitude),
                'created_at': marker.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': marker.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            }
        }, status=201)
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
        
    except Exception as e:
        print(f"Error creating marker: {str(e)}")  # Debug log
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)



def get_markers(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Get all markers from database with filters"""
    try:
        # Get filters from request
        category = request.GET.get('category', '')
        date_filter = request.GET.get('date_filter', '')
        start_date_str = request.GET.get('start_date', '')
        end_date_str = request.GET.get('end_date', '')
        search = request.GET.get('search', '')
        
        # Start with all markers
        queryset = MapMarker.objects.all().order_by('-created_at')
        
        # Apply category filter
        if category and category != 'all':
            queryset = queryset.filter(category=category)
        
        # Apply date filter
        if date_filter and date_filter != 'all':
            today = timezone.now().date()
            
            if date_filter == 'today':
                queryset = queryset.filter(created_at__date=today)
            elif date_filter == 'week':
                week_ago = today - timedelta(days=7)
                queryset = queryset.filter(created_at__date__gte=week_ago)
            elif date_filter == 'month':
                month_ago = today - timedelta(days=30)
                queryset = queryset.filter(created_at__date__gte=month_ago)
            elif date_filter == 'custom' and start_date_str and end_date_str:
                # Parse custom date range
                try:
                    start_date = datetime.datetime.strptime(start_date_str, '%Y-%m-%d').date()
                    end_date = datetime.datetime.strptime(end_date_str, '%Y-%m-%d').date()
                    
                    # Filter between dates (inclusive)
                    queryset = queryset.filter(
                        created_at__date__gte=start_date,
                        created_at__date__lte=end_date
                    )
                except ValueError:
                    pass  # Invalid date format, ignore custom filter
        
        # Apply search filter
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(description__icontains=search)
            )
        
        # Convert to JSON format
        markers_list = []
        for marker in queryset:
            markers_list.append({
                'id': marker.id,
                'title': marker.title,
                'description': marker.description,
                'category': marker.category,
                'color': marker.color,
                'latitude': float(marker.latitude),
                'longitude': float(marker.longitude),
                'lat': float(marker.latitude),  # For compatibility
                'lng': float(marker.longitude), # For compatibility
                'created_at': marker.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': marker.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                'created_by': marker.created_by.username if marker.created_by else 'Anonymous',
            })
        
        return JsonResponse({
            'success': True,
            'count': len(markers_list),
            'markers': markers_list,
            'filters': {
                'category': category,
                'date_filter': date_filter,
                'start_date': start_date_str,
                'end_date': end_date_str
            }
        })
        
    except Exception as e:
        print(f"Error getting markers: {str(e)}")  # Debug log
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


def delete_marker(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Delete a single marker"""
    try:
        marker_id = request.GET.get('id')
        
        if not marker_id:
            return JsonResponse({
                'success': False,
                'error': 'Marker ID is required'
            }, status=400)
        
        try:
            marker = MapMarker.objects.get(id=marker_id)
        except MapMarker.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Marker not found'
            }, status=404)
        
        # Check permission (optional - add if you have user auth)
        # if request.user != marker.created_by and not request.user.is_staff:
        #     return JsonResponse({
        #         'success': False,
        #         'error': 'Permission denied'
        #     }, status=403)
        
        marker.delete()
        
        return JsonResponse({
            'success': True,
            'message': 'Marker deleted successfully'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


def delete_marker(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Delete a single marker"""
    try:
        marker_id = request.GET.get('id')
        
        if not marker_id:
            return JsonResponse({
                'success': False,
                'error': 'Marker ID is required'
            }, status=400)
        
        try:
            marker = MapMarker.objects.get(id=marker_id)
        except MapMarker.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Marker not found'
            }, status=404)
        
        # Check permission (optional - uncomment if using authentication)
        # if request.user.is_authenticated and request.user != marker.created_by and not request.user.is_staff:
        #     return JsonResponse({
        #         'success': False,
        #         'error': 'Permission denied'
        #     }, status=403)
        
        marker.delete()
        
        return JsonResponse({
            'success': True,
            'message': 'Marker deleted successfully'
        })
        
    except Exception as e:
        print(f"Error deleting marker: {str(e)}")  # Debug log
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)
    
def delete_all_markers(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Delete all markers"""
    try:
        # Optional: Add authentication check
        # if not request.user.is_staff:
        #     return JsonResponse({
        #         'success': False,
        #         'error': 'Permission denied. Admin only.'
        #     }, status=403)
        
        count = MapMarker.objects.all().count()
        MapMarker.objects.all().delete()
        
        return JsonResponse({
            'success': True,
            'message': f'All {count} markers deleted successfully'
        })
        
    except Exception as e:
        print(f"Error deleting all markers: {str(e)}")  # Debug log
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

def update_marker(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Update an existing marker"""
    try:
        # Parse JSON data
        data = json.loads(request.body)
        print("Update data received:", data)  # Debug log
        
        # Check if marker ID is provided
        marker_id = data.get('id')
        if not marker_id:
            return JsonResponse({
                'success': False,
                'error': 'Marker ID is required'
            }, status=400)
        
        try:
            # Get the marker to update
            marker = MapMarker.objects.get(id=marker_id)
        except MapMarker.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': f'Marker with ID {marker_id} not found'
            }, status=404)
        
        # Optional: Check permissions (uncomment if using authentication)
        # if request.user.is_authenticated and marker.created_by and request.user != marker.created_by:
        #     return JsonResponse({
        #         'success': False,
        #         'error': 'You do not have permission to edit this marker'
        #     }, status=403)
        
        # Track what fields are being updated
        updated_fields = []
        
        # Update title if provided
        if 'title' in data:
            new_title = data['title'].strip()
            if new_title and new_title != marker.title:
                marker.title = new_title
                updated_fields.append('title')
        
        # Update description if provided
        if 'description' in data:
            new_description = data['description'].strip()
            if new_description != marker.description:
                marker.description = new_description
                updated_fields.append('description')
        
        # Update category if provided
        if 'category' in data and data['category'] in dict(MapMarker.CATEGORY_CHOICES):
            if data['category'] != marker.category:
                marker.category = data['category']
                updated_fields.append('category')
        
        # Update color if provided
        if 'color' in data and data['color'] in dict(MapMarker.COLOR_CHOICES):
            if data['color'] != marker.color:
                marker.color = data['color']
                updated_fields.append('color')
        
        # Update coordinates if provided
        coordinates_updated = False
        new_lat = None
        new_lng = None
        
        if 'latitude' in data and data['latitude'] is not None:
            try:
                new_lat = float(data['latitude'])
                if not (-90 <= new_lat <= 90):
                    return JsonResponse({
                        'success': False,
                        'error': 'Latitude must be between -90 and 90'
                    }, status=400)
            except (ValueError, TypeError):
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid latitude format'
                }, status=400)
        
        if 'longitude' in data and data['longitude'] is not None:
            try:
                new_lng = float(data['longitude'])
                if not (-180 <= new_lng <= 180):
                    return JsonResponse({
                        'success': False,
                        'error': 'Longitude must be between -180 and 180'
                    }, status=400)
            except (ValueError, TypeError):
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid longitude format'
                }, status=400)
        
        # Update coordinates if both are valid
        if new_lat is not None and new_lng is not None:
            if new_lat != float(marker.latitude) or new_lng != float(marker.longitude):
                marker.latitude = new_lat
                marker.longitude = new_lng
                coordinates_updated = True
                updated_fields.extend(['latitude', 'longitude'])
        elif new_lat is not None or new_lng is not None:
            # If only one coordinate is provided, return error
            return JsonResponse({
                'success': False,
                'error': 'Both latitude and longitude must be provided together'
            }, status=400)
        
        # Save the marker if any fields were updated
        if updated_fields:
            marker.save()
            
            return JsonResponse({
                'success': True,
                'message': f'Marker updated successfully. Fields updated: {", ".join(updated_fields)}',
                'marker': {
                    'id': marker.id,
                    'title': marker.title,
                    'description': marker.description,
                    'category': marker.category,
                    'color': marker.color,
                    'latitude': float(marker.latitude),
                    'longitude': float(marker.longitude),
                    'created_at': marker.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'updated_at': marker.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'updated_fields': updated_fields,
                    'coordinates_changed': coordinates_updated
                }
            })
        else:
            return JsonResponse({
                'success': True,
                'message': 'No changes detected. Marker remains unchanged.',
                'marker': {
                    'id': marker.id,
                    'title': marker.title,
                    'description': marker.description,
                    'category': marker.category,
                    'color': marker.color,
                    'latitude': float(marker.latitude),
                    'longitude': float(marker.longitude),
                    'created_at': marker.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'updated_at': marker.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'updated_fields': []
                }
            })
            
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
        
    except Exception as e:
        print(f"Error updating marker: {str(e)}")  # Debug log
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500
        )
    
def manage_user(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    users = User.objects.all()
    context = {'users': users}
    return render(request, 'user_manage.html', context)

from django.shortcuts import render
from django.db.models import Q
from datetime import datetime
from django.shortcuts import render
from django.db.models import Q
from datetime import datetime

@login_required
def track_user(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return render(request, 'user_track.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'hierarchy_data': {'name': 'No Access', 'type': 'root', 'children': []},
            'selected_category': '',
            'start_date': '',
            'end_date': '',
            'all_categories': [],
            'total_threats': 0,
            'show_all_categories': False,
            'date_stats': {},
        })
    
    selected_category = request.GET.get('category')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    
    # Get threats based on user's unit
    if user.is_superuser:
        # Superusers can see all threats
        threats = ThreatAlert.objects.all()
    elif hasattr(user, 'role') and user.role == 2:  # Admin
        # Admins can see threats from their unit + their own + system threats
        threats = ThreatAlert.objects.filter(
            Q(created_by__unit=user.unit) |
            Q(created_by=user) |
            Q(created_by__isnull=True)
        )
    else:
        # Regular users can only see threats from their unit
        threats = ThreatAlert.objects.filter(
            created_by__unit=user.unit
        )
    
    # Apply category filter
    if selected_category:
        try:
            selected_category = int(selected_category)
            threats = threats.filter(category_id=selected_category)
        except ValueError:
            selected_category = None
    
    # Apply date range filter if provided
    if start_date:
        try:
            start_datetime = datetime.strptime(start_date, '%Y-%m-%d')
            threats = threats.filter(timestamp__date__gte=start_datetime.date())
        except ValueError:
            pass  # Invalid date format
    
    if end_date:
        try:
            end_datetime = datetime.strptime(end_date, '%Y-%m-%d')
            threats = threats.filter(timestamp__date__lte=end_datetime.date())
        except ValueError:
            pass  # Invalid date format
    
    threats = threats.order_by('timestamp')
    
    # Calculate date statistics
    date_stats = {}
    for threat in threats:
        date_str = threat.timestamp.strftime('%Y-%m-%d')
        date_stats[date_str] = date_stats.get(date_str, 0) + 1
    
    # Get accessible categories for filter dropdown
    if user.is_superuser:
        # Superusers see all categories
        accessible_categories = ThreatCategory.objects.all()
    else:
        # Get categories accessible to user's unit
        accessible_categories = ThreatCategory.objects.filter(
            Q(created_by__unit=user.unit) |
            Q(created_by=user) |
            Q(created_by__isnull=True)
        )
    
    # Get category choices from accessible categories
    category_choices = []
    for category in accessible_categories:
        category_choices.append((category.id, category.name))
    
    # Get selected category name
    selected_category_name = "All Categories"
    if selected_category:
        try:
            selected_category_obj = ThreatCategory.objects.get(id=selected_category)
            selected_category_name = selected_category_obj.name
        except ThreatCategory.DoesNotExist:
            selected_category_name = "Unknown Category"
    
    # Build hierarchy data
    hierarchy_data = {
        'name': selected_category_name,
        'type': 'root',
        'children': [],
        'date_stats': date_stats,
        'total_count': threats.count(),
        'date_range': {
            'start': start_date,
            'end': end_date
        },
        'user_unit': user.unit if hasattr(user, 'unit') else None,
        'unit_count': threats.filter(created_by__unit=user.unit).count() if hasattr(user, 'unit') else 0,
    }
    
    if selected_category:
        # Single category view
        date_groups = {}
        for threat in threats:
            date_str = threat.timestamp.strftime('%Y-%m-%d')
            if date_str not in date_groups:
                date_groups[date_str] = []
            
            # Convert threat to dictionary
            threat_dict = {
                'id': threat.id,
                'name': f"Threat-{threat.id}",
                'type': 'threat',
                'title': threat.title,
                'content': threat.content,
                'severity': threat.severity,
                'severity_display': threat.get_severity_display() if hasattr(threat, 'get_severity_display') else threat.severity.title(),
                'category': threat.category.name if threat.category else 'Uncategorized',
                'category_id': threat.category.id if threat.category else None,
                'source': threat.source,
                'url': threat.url,
                'province': threat.get_province_display() if hasattr(threat, 'get_province_display') and threat.province else threat.province or 'N/A',
                'image_url': threat.image.url if threat.image and hasattr(threat.image, 'url') else '',
                'has_image': bool(threat.image),
                'has_video': bool(threat.video),
                'date': date_str,
                'time': threat.timestamp.strftime('%H:%M'),
                'full_timestamp': threat.timestamp.isoformat(),
                'creator_name': threat.created_by.username if threat.created_by else 'System',
                'creator_unit': threat.created_by.unit if threat.created_by and hasattr(threat.created_by, 'unit') else 'System',
                'is_own_unit': threat.created_by.unit == user.unit if threat.created_by and hasattr(threat.created_by, 'unit') and hasattr(user, 'unit') else False,
            }
            date_groups[date_str].append(threat_dict)
        
        for date_str, date_threats in date_groups.items():
            hierarchy_data['children'].append({
                'name': date_str,
                'type': 'date',
                'count': len(date_threats),
                'children': date_threats  # Already converted to dictionaries
            })
    else:
        # All categories view - group by category
        category_groups = {}
        for threat in threats:
            category = threat.category
            category_id = category.id if category else 'uncategorized'
            category_name = category.name if category else 'Uncategorized'
            
            if category_id not in category_groups:
                category_groups[category_id] = {
                    'category_name': category_name,
                    'category_id': category_id,
                    'threats': []
                }
            
            # Convert threat to simple dictionary for preview
            threat_dict = {
                'id': threat.id,
                'title': threat.title[:50] + '...' if len(threat.title) > 50 else threat.title,
                'severity': threat.severity,
                'severity_display': threat.get_severity_display() if hasattr(threat, 'get_severity_display') else threat.severity.title(),
                'date': threat.timestamp.strftime('%Y-%m-%d'),
                'creator_unit': threat.created_by.unit if threat.created_by and hasattr(threat.created_by, 'unit') else 'System',
            }
            category_groups[category_id]['threats'].append(threat_dict)
        
        for category_id, group in category_groups.items():
            hierarchy_data['children'].append({
                'name': group['category_name'],
                'type': 'category',
                'category_id': category_id if category_id != 'uncategorized' else None,
                'count': len(group['threats']),
                'threats': group['threats'][:5]  # Show first 5 threats as preview
            })
    
    # Convert date_stats to list of dictionaries for easier JSON serialization
    date_stats_list = []
    for date_str, count in date_stats.items():
        date_stats_list.append({
            'date': date_str,
            'count': count
        })
    
    context = {
        'hierarchy_data_json': json.dumps(hierarchy_data, default=str),  # Convert to JSON for JavaScript
        'hierarchy_data': hierarchy_data,  # Keep Python object for template
        'selected_category': selected_category,
        'start_date': start_date,
        'end_date': end_date,
        'all_categories': category_choices,  # Use accessible categories only
        'total_threats': threats.count(),
        'show_all_categories': not selected_category,
        'date_stats': date_stats,
        'date_stats_json': json.dumps(date_stats_list, default=str),
        'user_unit': user.unit if hasattr(user, 'unit') else None,
        'is_superuser': user.is_superuser,
        'is_admin': getattr(user, 'role', 0) == 2,
        'unit_threat_count': threats.filter(created_by__unit=user.unit).count() if hasattr(user, 'unit') else 0,
    }
    
    return render(request, 'user_track.html', context)


def add_social_media_url(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    """
    View to add new social media URL with source department
    Using the same pattern as loginLogic
    """
    
    if request.method == 'POST':
        # Get form data
        url = request.POST.get('url', '').strip()
        department = request.POST.get('department', '').strip()
        
        # Validate required fields
        if not url:
            return render(request, 'add_url.html', {
                'alert_type': 'error',
                'alert_message': 'URL is required!',
                'url': url,
                'department': department,
            })
        
        if not department:
            return render(request, 'add_url.html', {
                'alert_type': 'error',
                'alert_message': 'Source department is required!',
                'url': url,
                'department': department,
            })
        
        # Validate URL format
        if not url.startswith(('http://', 'https://')):
            return render(request, 'add_url.html', {
                'alert_type': 'error',
                'alert_message': 'URL should start with http:// or https://',
                'url': url,
                'department': department,
            })
        
        # Check if URL already exists
        try:
            existing_url = SocialMediaURL.objects.get(url=url)
            return render(request, 'add_url.html', {
                'alert_type': 'error',
                'alert_message': f'This URL already exists',
                'url': url,
                'department': department,
            })
        except SocialMediaURL.DoesNotExist:
            pass  # URL doesn't exist, continue
        
        try:
            # Create new URL entry
            new_entry = SocialMediaURL.objects.create(
                url=url,
                source_department=department,
                status='pending',
                submitted_date=timezone.now()
            )
            
            # Redirect with success message in context
            return render(request, 'add_url.html', {
                'alert_type': 'success',
                'alert_message': f'✅ URL added successfully',
                'url': '',  # Clear form on success
                'department': '',
            })
            
        except Exception as e:
            return render(request, 'add_url.html', {
                'alert_type': 'error',
                'alert_message': f'Error saving URL: {str(e)}',
                'url': url,
                'department': department,
            })
    
    # GET request - show empty form
    return render(request, 'add_url.html', {
        'url': '',
        'department': '',
    })


def list_social_media_url(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """List all URLs with search functionality and pagination"""
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    
    # Get page number from request
    page = request.GET.get('page', 1)
    
    urls = SocialMediaURL.objects.all()
    
    # Search functionality - CORRECTED: using 'name' instead of 'full_name'
    if search_query:
        urls = urls.filter(
            Q(url__icontains=search_query) |
            Q(name__icontains=search_query) |  # Changed from full_name to name
            Q(personnel_no__icontains=search_query) |
            Q(user_id__icontains=search_query) |
            Q(source_department__icontains=search_query) |
            Q(rank__icontains=search_query) |  # Added rank to search
            Q(unit__icontains=search_query)    # Added unit to search
        )
    
    # Status filter
    if status_filter:
        urls = urls.filter(status=status_filter)
    
    # Date range filter (for submitted_date)
    if start_date:
        try:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            urls = urls.filter(submitted_date__date__gte=start_date_obj)
        except ValueError:
            pass
    
    if end_date:
        try:
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            # For end of day filtering
            urls = urls.filter(submitted_date__date__lte=end_date_obj)
        except ValueError:
            pass
    
    # Order by submitted_date (newest first)
    urls = urls.order_by('-submitted_date')
    
    # Pagination - 7 items per page
    paginator = Paginator(urls, 6)
    
    try:
        urls_paginated = paginator.page(page)
    except PageNotAnInteger:
        urls_paginated = paginator.page(1)
    except EmptyPage:
        urls_paginated = paginator.page(paginator.num_pages)
    
    context = {
        'urls': urls_paginated,
        'search_query': search_query,
        'status_filter': status_filter,
        'start_date': start_date,
        'end_date': end_date,
        'status_choices': SocialMediaURL.STATUS_CHOICES,
    }
    return render(request, 'url_list.html', context)


def update_social_media_url(request, url_id):
    if not check_access(request):
        return redirect('logout')  # Create this view

    """Update social media URL details via AJAX with image upload"""
    try:
        url_obj = SocialMediaURL.objects.get(id=url_id)
        
        # Update text fields
        url_obj.personnel_no = request.POST.get('personnel_no', '')
        url_obj.rank = request.POST.get('rank', '')
        url_obj.name = request.POST.get('name', '')
        url_obj.unit = request.POST.get('unit', '')
        url_obj.status = request.POST.get('status', 'pending')
        url_obj.platform = request.POST.get('platform', '')
        url_obj.user_id = request.POST.get('user_id', '')
        url_obj.description = request.POST.get('description', '')
        url_obj.remarks = request.POST.get('remarks', '')
        
        # Function to handle single photo upload
        def handle_single_photo(field_name, remove_field_name):
            """Handle upload for a single photo field"""
            if field_name in request.FILES:
                photo_files = request.FILES.getlist(field_name)
                if photo_files:
                    photo_file = photo_files[0]  # Take the first file if multiple
                    
                    # Validate file type
                    allowed_types = ['image/jpeg', 'image/png', 'image/gif', 'image/jpg', 'image/webp']
                    if photo_file.content_type not in allowed_types:
                        raise ValueError(f'Invalid file type for {field_name}. Only JPG, PNG, GIF, and WebP are allowed.')
                    
                    # Validate file size (5MB limit)
                    if photo_file.size > 5 * 1024 * 1024:  # 5MB
                        raise ValueError(f'File size too large for {field_name} ({photo_file.size/1024/1024:.1f}MB). Maximum size is 5MB.')
                    
                    # Get current photo field
                    current_photo = getattr(url_obj, field_name)
                    
                    # Delete old photo file if exists
                    if current_photo:
                        current_photo.delete(save=False)
                    
                    # Save new photo
                    setattr(url_obj, field_name, photo_file)
                    print(f"✅ {field_name} uploaded: {photo_file.name}")
        
        # Handle all three photo fields
        try:
            # Handle main photo
            handle_single_photo('photo', 'remove_photo')
            
            # Handle photo_one
            handle_single_photo('photo_one', 'remove_photo_one')
            
            # Handle photo_two
            handle_single_photo('photo_two', 'remove_photo_two')
            
        except ValueError as ve:
            return JsonResponse({
                'success': False, 
                'error': str(ve),
                'alert_type': 'error',
                'alert_message': f'❌ {str(ve)}'
            }, status=400)
        
        # Handle remove photo checkboxes
        if request.POST.get('remove_photo') == 'on':
            if url_obj.photo:
                url_obj.photo.delete(save=False)
                url_obj.photo = None
        
        if request.POST.get('remove_photo_one') == 'on':
            if url_obj.photo_one:
                url_obj.photo_one.delete(save=False)
                url_obj.photo_one = None
        
        if request.POST.get('remove_photo_two') == 'on':
            if url_obj.photo_two:
                url_obj.photo_two.delete(save=False)
                url_obj.photo_two = None
        
        # Save the object
        url_obj.completed_date = timezone.now()
        url_obj.save()
        
        # Prepare response data with alert info
        response_data = {
            'success': True, 
            'message': 'URL updated successfully',
            'alert_type': 'success',
            'alert_message': '✅ URL updated successfully',
            'data': {
                'id': url_obj.id,
                'name': url_obj.name,
                'personnel_no': url_obj.personnel_no,
                'status': url_obj.status,
                'platform': url_obj.platform,
                'has_photo': bool(url_obj.photo),
                'has_photo_one': bool(url_obj.photo_one),
                'has_photo_two': bool(url_obj.photo_two)
            },
            'redirect_url': '/social-media/list/'
        }
        
        # Function to get photo info
        def get_photo_info(field_name):
            photo = getattr(url_obj, field_name)
            if photo:
                return {
                    'type': 'uploaded',
                    'url': photo.url if hasattr(photo, 'url') else '',
                    'name': photo.name if hasattr(photo, 'name') else ''
                }
            return None
        
        # Add photo info to response for all three fields
        response_data['photo'] = get_photo_info('photo')
        response_data['photo_one'] = get_photo_info('photo_one')
        response_data['photo_two'] = get_photo_info('photo_two')
        
        return JsonResponse(response_data)
        
    except SocialMediaURL.DoesNotExist:
        return JsonResponse({
            'success': False, 
            'error': 'URL not found',
            'alert_type': 'error',
            'alert_message': '❌ URL not found'
        }, status=404)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'success': False, 
            'error': str(e),
            'alert_type': 'error',
            'alert_message': f'❌ Error: {str(e)}'
        }, status=500)
         

def dashboard_social_media(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Social Media tracking dashboard with custom date range"""
    
    # Get parameters from request
    interval = request.GET.get('interval', 'month')
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    custom_range = request.GET.get('custom_range', 'false')
    
    now = timezone.now()
    start_date = None
    end_date = now
    date_range_label = ""
    
    # Handle custom date range
    if custom_range == 'true' and start_date_str and end_date_str:
        try:
            # Parse dates
            start_date_parsed = parse_date(start_date_str)
            end_date_parsed = parse_date(end_date_str)
            
            if start_date_parsed and end_date_parsed:
                # Create timezone-aware datetimes
                start_date = timezone.make_aware(
                    datetime.combine(start_date_parsed, datetime.min.time())
                )
                end_date = timezone.make_aware(
                    datetime.combine(end_date_parsed, datetime.max.time())
                )
                
                date_range_label = f"{start_date_str} to {end_date_str}"
                interval = 'custom'
        except (ValueError, TypeError):
            # Fallback to default interval if date parsing fails
            interval = 'month'
            custom_range = 'false'
    
    # Handle predefined intervals if not using custom range
    if custom_range != 'true':
        if interval == 'today':
            start_date = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
            date_range_label = "Today"
        elif interval == 'week':
            start_date = timezone.now() - timedelta(days=7)
            date_range_label = "Last 7 Days"
        elif interval == 'month':
            start_date = timezone.now() - timedelta(days=30)
            date_range_label = "Last 30 Days"
        elif interval == 'quarter':
            start_date = timezone.now() - timedelta(days=90)
            date_range_label = "Last 90 Days"
        elif interval == 'year':
            start_date = timezone.now() - timedelta(days=365)
            date_range_label = "Last 365 Days"
        else:  # all
            start_date = None
            date_range_label = "All Time"
    
    # Base queryset
    urls = SocialMediaURL.objects.all()
    
    # Apply date filters
    if start_date:
        urls = urls.filter(submitted_date__gte=start_date)
    if end_date and interval == 'custom':
        urls = urls.filter(submitted_date__lte=end_date)
    
    total_urls = urls.count()
    
    # Status data
    status_data = []
    for status_code, status_name in SocialMediaURL.STATUS_CHOICES:
        count = urls.filter(status=status_code).count()
        percentage = round((count / total_urls * 100), 1) if total_urls > 0 else 0
        
        # Get icon and color based on status
        if status_code == 'pending':
            icon = 'fa-clock'
            color = '#f6c23e'
        elif status_code == 'searching':
            icon = 'fa-search'
            color = '#36b9cc'
        elif status_code == 'found':
            icon = 'fa-check-circle'
            color = '#1cc88a'
        elif status_code == 'not_found':
            icon = 'fa-times-circle'
            color = '#e74a3b'
        else:  # error
            icon = 'fa-exclamation-triangle'
            color = '#858796'
        
        status_data.append({
            'code': status_code,
            'label': status_name,
            'count': count,
            'percentage': percentage,
            'icon': icon,
            'color': color
        })
    
    # Recent activity counts (always relative to today)
    today_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
    last_7_days = SocialMediaURL.objects.filter(
        submitted_date__gte=today_start - timedelta(days=7)
    ).count()
    
    last_30_days = SocialMediaURL.objects.filter(
        submitted_date__gte=today_start - timedelta(days=30)
    ).count()
    
    # For custom range, calculate days difference
    days_in_range = None
    if interval == 'custom' and start_date and end_date:
        # Make both naive for days calculation
        start_naive = start_date.replace(tzinfo=None)
        end_naive = end_date.replace(tzinfo=None)
        days_in_range = (end_naive - start_naive).days + 1
    
    context = {
        'selected_interval': interval,
        'start_date': start_date_str if start_date_str else '',
        'end_date': end_date_str if end_date_str else '',
        'date_range_label': date_range_label,
        'days_in_range': days_in_range,
        'status_data': status_data,
        'total_urls': total_urls,
        'last_7_days': last_7_days,
        'last_30_days': last_30_days,
    }
    
    return render(request, 'dashboard_social_media.html', context)

from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from datetime import datetime
from .models import SocialMediaURL

def generate_social_media_report(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Generate Word document report for selected social media URLs"""
    if request.method == 'POST':
        try:
            # Parse JSON data
            data = json.loads(request.body)
            url_ids = data.get('url_ids', [])
            nepali_format = data.get('nepali_format', True)
            
            if not url_ids:
                return JsonResponse({
                    'success': False,
                    'error': 'No URLs selected for report'
                })
            
            # Create Word document
            document = Document()
            
            # Use PORTRAIT orientation
            section = document.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            
            # Set margins
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
            
            # Add title
            title_paragraph = document.add_paragraph()
            if nepali_format:
                title_text = "सामाजिक संजाल सम्बन्धि सैनिक व्यक्तिको विवरण"
            else:
                title_text = "Social Media Report"
                
            title_run = title_paragraph.add_run(title_text)
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.name = 'Calibri'
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            date_paragraph = document.add_paragraph()
            current_date = datetime.now().strftime("%Y-%m-%d %H:%M")
            date_text = f"मिति: {current_date}" if nepali_format else f"Date: {current_date}"
            date_run = date_paragraph.add_run(date_text)
            date_run.font.size = Pt(10)
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            document.add_paragraph()  # Add spacing
            
            # Create table - Now with 9 columns including remarks
            table = document.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            columns = table.columns
            widths = [
                0.5,   # Serial (सि.नं.)
                0.8,   # Personnel No (व्य.नं.)
                1.0,   # Rank (दर्जा)
                2.0,   # Name (नामथर)
                1.5,   # Unit (युनिट)
                1.2,   # Platform (सामाजिक संजाल)
                1.2,   # User ID (User ID)
                2.0,   # Description (गतिविधि विवरण)
                1.5    # Remarks (कै./Remarks)
            ]
            
            for i, width in enumerate(widths):
                columns[i].width = Inches(width)
            
            # Add table headers
            if nepali_format:
                headers = [
                    "सि.नं.",           # Serial
                    "व्य.नं.",          # Personnel No
                    "दर्जा",            # Rank
                    "नामथर",           # Name
                    "युनिट",            # Unit
                    "सामाजिक संजाल",   # Platform
                    "User ID",          # User ID
                    "गतिविधि विवरण",   # Description
                    "कै."              # Remarks
                ]
            else:
                headers = [
                    "SN",               # Serial
                    "Personnel No",     # Personnel No
                    "Rank",             # Rank
                    "Full Name",        # Name
                    "Unit",             # Unit
                    "Platform",         # Platform
                    "User ID",          # User ID
                    "Activity Description", # Description
                    "Remarks"          # Remarks
                ]
            
            # Style header cells
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = header
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
            
            # Add data rows
            serial_number = 1
            for url_id in url_ids:
                try:
                    url_obj = SocialMediaURL.objects.filter(id=url_id).first()
                    if not url_obj:
                        continue
                    
                    # Add new row
                    row_cells = table.add_row().cells
                    
                    # Fill row data
                    # Serial Number
                    row_cells[0].text = str(serial_number)
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Personnel No
                    row_cells[1].text = str(url_obj.personnel_no or "")
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Rank
                    row_cells[2].text = url_obj.rank or ""
                    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Name
                    row_cells[3].text = url_obj.name or ""
                    row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Unit
                    row_cells[4].text = url_obj.unit or ""
                    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Platform
                    row_cells[5].text = url_obj.platform or ""
                    row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # User ID
                    row_cells[6].text = url_obj.user_id or ""
                    row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Description (truncate if too long)
                    description = url_obj.description or ""
                    if len(description) > 150:  # Reduced from 200 to accommodate more columns
                        description = description[:147] + "..."
                    if url_obj.url:
                        url_display = url_obj.url
                        if len(url_display) > 50:
                            url_display = url_display[:47] + "..."
                        description += f"\nURL: {url_display}"
                    
                    row_cells[7].text = description
                    row_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Remarks
                    remarks = url_obj.remarks or ""
                    if len(remarks) > 100:  # Truncate if too long
                        remarks = remarks[:97] + "..."
                    row_cells[8].text = remarks
                    row_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Style all cells
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                run.font.name = 'Calibri'
                    
                    serial_number += 1
                    
                except Exception as e:
                    print(f"Error processing URL ID {url_id}: {str(e)}")
                    continue
            
            # Add summary
            if serial_number > 1:
                document.add_paragraph()
                
                summary_para = document.add_paragraph()
                if nepali_format:
                    summary_text = f"जम्मा रेकर्डहरू: {serial_number - 1}"
                else:
                    summary_text = f"Total Records: {serial_number - 1}"
                    
                summary_run = summary_para.add_run(summary_text)
                summary_run.font.size = Pt(10)
                summary_run.bold = True
                summary_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add a footer note if needed
                document.add_paragraph()
                footer_para = document.add_paragraph()
                if nepali_format:
                    footer_text = "टिप्पणी: यो विवरण सामाजिक संजाल व्यवस्थापन प्रणालीबाट स्वचालित रूपमा तयार गरिएको हो।"
                else:
                    footer_text = "Note: This report is automatically generated from the Social Media Management System."
                
                footer_run = footer_para.add_run(footer_text)
                footer_run.font.size = Pt(9)
                footer_run.italic = True
                footer_run.font.name = 'Calibri'
                footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Save document
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"social_media_report_{timestamp}.docx"
            
            # Return response
            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Error generating report: {str(e)}'
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Invalid request method'
    }, status=400)


def photo_social_media(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Display all photos from SocialMediaURL model in table format"""
    # Get filter parameters
    personnel_no_filter = request.GET.get('personnel_no', '')
    name_filter = request.GET.get('name', '')
    rank_filter = request.GET.get('rank', '')
    unit_filter = request.GET.get('unit', '')
    
    # Start with all objects that have at least one photo
    urls = SocialMediaURL.objects.filter(
        Q(photo__isnull=False) | 
        Q(photo_one__isnull=False) | 
        Q(photo_two__isnull=False)
    ).exclude(
        Q(photo='') & Q(photo_one='') & Q(photo_two='')
    ).distinct()
    
    # Apply individual filters
    if personnel_no_filter:
        urls = urls.filter(personnel_no__icontains=personnel_no_filter)
    
    if name_filter:
        urls = urls.filter(name__icontains=name_filter)
    
    if rank_filter:
        urls = urls.filter(rank__icontains=rank_filter)
    
    if unit_filter:
        urls = urls.filter(unit__icontains=unit_filter)
    
    # Add pagination
    paginator = Paginator(urls, 10)  # Show 10 per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Prepare context for template
    context = {
        'urls': page_obj,
        'personnel_no_filter': personnel_no_filter,
        'name_filter': name_filter,
        'rank_filter': rank_filter,
        'unit_filter': unit_filter,
        'total_urls': urls.count(),
    }
    
    return render(request, 'photo_social_media.html', context)

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.http import JsonResponse
from django.views.decorators.http import require_POST, require_http_methods
from django.db.models import Q
from .models import ThreatCategory
import json

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q

@login_required
def category_list(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    """
    List categories created by the logged-in user with pagination and search
    """
    if not check_access(request):
        messages.error(request, 'Access denied!')
        return redirect('dashboard')  # Create this view
    
    user = request.user
    
    # Check if user account is active
    if not user.is_active:
        return render(request, 'category_list.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.',
            'categories': [],
            'search_query': '',
            'active_count': 0,
            'inactive_count': 0,
        })
    
    # Get search query
    search_query = request.GET.get('search', '').strip()
    
    # Get categories created ONLY by the current user
    categories_list = ThreatCategory.objects.filter(created_by=user)
    
    # Apply search filter if query exists
    if search_query:
        categories_list = categories_list.filter(
            Q(name__icontains=search_query)
        )
    
    # Get counts for stats
    total_count = categories_list.count()
    active_count = categories_list.filter(is_active=True).count()
    inactive_count = categories_list.filter(is_active=False).count()
    
    # Order by name (or by creation date if preferred)
    categories_list = categories_list.order_by('name')
    
    # Pagination - 5 items per page
    paginator = Paginator(categories_list, 5)
    page = request.GET.get('page')
    
    try:
        categories = paginator.page(page)
    except PageNotAnInteger:
        categories = paginator.page(1)
    except EmptyPage:
        categories = paginator.page(paginator.num_pages)
    
    context = {
        'categories': categories,
        'search_query': search_query,
        'total_count': total_count,
        'active_count': active_count,
        'inactive_count': inactive_count,
    }
    return render(request, 'category_list.html', context)

@login_required
def category_add(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    """
    Add a new category
    """
    if request.method == 'POST':
        category_name = request.POST.get('name', '').strip()
        is_active = request.POST.get('is_active', 'off') == 'on'
        
        # Validation
        if not category_name:
            messages.error(request, "Category name is required.")
            return render(request, 'category_add.html', {
                'form_data': request.POST
            })
        
        # Clean and format category name
        category_name = category_name.strip()
        
        try:
            # Try to create the category
            category = ThreatCategory.objects.create(
                name=category_name,
                created_by=request.user,
                is_active=is_active
            )
            
            messages.success(request, f"Category '{category.name}' added successfully!")
            return redirect('category_list')
            
        except IntegrityError:
            # Handle duplicate name error
            messages.error(request, f"A category named '{category_name}' already exists.")
            return render(request, 'category_add.html', {
                'form_data': request.POST
            })
        except Exception as e:
            # Handle other errors
            messages.error(request, f"Error creating category: {str(e)}")
            return render(request, 'category_add.html', {
                'form_data': request.POST
            })
    
    return render(request, 'category_add.html', {})

@login_required
def category_edit(request, category_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Edit an existing category
    """
    category = get_object_or_404(ThreatCategory, id=category_id, created_by=request.user)
    
    if request.method == 'POST':
        category_name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        is_active = request.POST.get('is_active', 'off') == 'on'
        
        # Validation
        if not category_name:
            messages.error(request, "Category name is required.")
            return render(request, 'category_edit.html', {'category': category})
        
        # Check for duplicate name (excluding current category)
        duplicate = ThreatCategory.objects.filter(
            name__iexact=category_name,
            created_by=request.user
        ).exclude(id=category_id).exists()
        
        if duplicate:
            messages.error(request, f"You already have a category named '{category_name}'.")
            return render(request, 'category_edit.html', {'category': category})
        
        # Update category
        category.name = category_name
        category.description = description
        category.is_active = is_active
        category.save()
        
        messages.success(request, f"Category '{category.name}' updated successfully!")
        return redirect('category_list')
    
    return render(request, 'category_edit.html', {'category': category})

@login_required
@require_http_methods(["POST"])
def category_delete(request, category_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Delete a category (POST only)
    """
    try:
        category = get_object_or_404(ThreatCategory, id=category_id, created_by=request.user)
        category_name = category.name
        
        # Check if category has alerts
        alert_count = category.threat_alerts.count()
        
        # Delete the category
        category.delete()
        
        if alert_count > 0:
            messages.success(request, f"Category '{category_name}' deleted. {alert_count} alert(s) updated.")
        else:
            messages.success(request, f"Category '{category_name}' deleted successfully.")
            
    except ThreatCategory.DoesNotExist:
        messages.error(request, "Category not found or you don't have permission to delete it.")
    except Exception as e:
        messages.error(request, f"Error deleting category: {str(e)}")
    
    return redirect('category_list')

@login_required
@require_POST
def toggle_category_status(request, category_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Toggle category active/inactive status (AJAX - POST only)
    """
    try:
        # Check if it's an AJAX request
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            # Get category with user permission check
            category = get_object_or_404(ThreatCategory, id=category_id, created_by=request.user)
            
            # Get current status
            current_status = category.is_active
            
            # Toggle the status
            category.is_active = not current_status
            category.save()
            
            # Prepare response data
            response_data = {
                'success': True,
                'is_active': category.is_active,
                'message': f"Category '{category.name}' is now {'Active' if category.is_active else 'Inactive'}",
                'category_id': category.id,
                'category_name': category.name,
                'previous_status': current_status,
                'new_status': category.is_active
            }
            
            return JsonResponse(response_data)
        
        else:
            # If not AJAX request, handle as regular POST
            category = get_object_or_404(ThreatCategory, id=category_id, created_by=request.user)
            category.is_active = not category.is_active
            category.save()
            
            messages.success(
                request, 
                f"Category '{category.name}' is now {'Active' if category.is_active else 'Inactive'}"
            )
            return redirect('category_list')
            
    except ThreatCategory.DoesNotExist:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False,
                'message': 'Category not found or you do not have permission to modify it.'
            }, status=404)
        else:
            messages.error(request, "Category not found or you don't have permission to modify it.")
            return redirect('category_list')
            
    except Exception as e:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False,
                'message': f'An error occurred: {str(e)}'
            }, status=400)
        else:
            messages.error(request, f"An error occurred: {str(e)}")
            return redirect('category_list')

@login_required
def bulk_toggle_status(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Bulk toggle status for multiple categories
    """
    if request.method == 'POST' and request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        try:
            data = json.loads(request.body)
            category_ids = data.get('category_ids', [])
            new_status = data.get('status')
            
            if not category_ids or new_status is None:
                return JsonResponse({
                    'success': False,
                    'message': 'Invalid request data'
                }, status=400)
            
            # Get categories for current user
            categories = ThreatCategory.objects.filter(
                id__in=category_ids,
                created_by=request.user
            )
            
            updated_count = 0
            for category in categories:
                category.is_active = new_status
                category.save()
                updated_count += 1
            
            return JsonResponse({
                'success': True,
                'message': f'Successfully updated {updated_count} category(s)',
                'updated_count': updated_count
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Error: {str(e)}'
            }, status=400)
    
    return JsonResponse({'success': False, 'message': 'Invalid request method'}, status=400)


@login_required
def keyword_listing(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Dangerous Keywords Listing with search, filtering, and pagination
    - Shows user's own keywords + system keywords
    """
    user = request.user
    
    # Get search parameters
    search_query = request.GET.get('search', '')
    category_filter = request.GET.get('category', '')
    page_number = request.GET.get('page', 1)
    
    # Start with user's keywords + system keywords (created_by=None)
    keywords = DangerousKeyword.objects.filter(
        Q(created_by=user) | Q(created_by__isnull=True)
    ).order_by('-created_at')  # Order by newest first
    
    # Apply search filter
    if search_query:
        keywords = keywords.filter(
            Q(word__icontains=search_query) |
            Q(category__icontains=search_query)
        )
    
    # Apply category filter
    if category_filter:
        keywords = keywords.filter(category=category_filter)
    
    # Get unique categories from user's + system keywords
    categories = keywords.values_list(
        'category', flat=True
    ).distinct().order_by('category')
    
    # Get counts BEFORE pagination
    total_count = keywords.count()
    active_count = keywords.filter(is_active=True).count()
    inactive_count = keywords.filter(is_active=False).count()
    
    # Get user's personal count vs system count
    user_count = DangerousKeyword.objects.filter(created_by=user).count()
    system_count = DangerousKeyword.objects.filter(created_by__isnull=True).count()
    
    # Pagination - 15 items per page
    paginator = Paginator(keywords, 5)
    
    try:
        keywords_page = paginator.page(page_number)
    except PageNotAnInteger:
        keywords_page = paginator.page(1)
    except EmptyPage:
        keywords_page = paginator.page(paginator.num_pages)
    
    # Calculate start index for numbering
    start_index = (keywords_page.number - 1) * paginator.per_page
    
    context = {
        'keywords': keywords_page,  # Changed from queryset to paginated page
        'paginator': paginator,
        'page_obj': keywords_page,
        'total_count': total_count,
        'active_count': active_count,
        'inactive_count': inactive_count,
        'user_count': user_count,
        'system_count': system_count,
        'search_query': search_query,
        'category_filter': category_filter,
        'categories': categories,
        'current_user': user,
        'start_index': start_index + 1,
    }
    
    return render(request, 'dangerouskeyboard.html', context)

@login_required
def keyword_edit(request, id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Edit an existing dangerous keyword
    """
    # Get the keyword - if you want user-specific editing, add filter
    keyword = get_object_or_404(DangerousKeyword, id=id)
    # If you want only user's own keywords: get_object_or_404(DangerousKeyword, id=id, created_by=request.user)
    
    if request.method == 'POST':
        word = request.POST.get('word', '').strip().lower()
        category = request.POST.get('category', '').strip()
        is_active = request.POST.get('is_active', 'off') == 'on'
        
        # Validation
        if not word:
            messages.error(request, "Keyword is required.")
            return render(request, 'dangerousedit.html', {'keyword': keyword})
        
        if not category:
            messages.error(request, "Category is required.")
            return render(request, 'dangerousedit.html', {'keyword': keyword})
        
        # Check for duplicate word-category combination (excluding current keyword)
        duplicate = DangerousKeyword.objects.filter(
            Q(word__iexact=word) & Q(category__iexact=category)
        ).exclude(id=id).exists()
        
        if duplicate:
            messages.error(request, f"The keyword '{word}' already exists in category '{category}'.")
            return render(request, 'dangerousedit.html', {'keyword': keyword})
        
        # Update keyword
        keyword.word = word
        keyword.category = category
        keyword.is_active = is_active
        keyword.save()
        
        messages.success(request, f"Keyword '{keyword.word}' updated successfully!")
        return redirect('keywords_listing')
    
    return render(request, 'dangerousedit.html', {'keyword': keyword})

@login_required
def keyword_delete(request, id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Delete dangerous keyword
    """
    keyword = get_object_or_404(DangerousKeyword, id=id)
    
    if request.method == 'POST':
        word = keyword.word
        keyword.delete()
        messages.success(request, f'Keyword "{word}" deleted successfully!')
        return redirect('keywords_listing')
    
    return render(request, 'dangerousdelete.html', {'keyword': keyword})

from django.shortcuts import render
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException, TimeoutException, NoSuchElementException, StaleElementReferenceException, InvalidSessionIdException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
import time
import traceback
import random
import os
import re
from datetime import datetime
from urllib.parse import urlparse, parse_qs, urlencode
import json
import hashlib
import mimetypes
from django.http import HttpResponse, FileResponse, JsonResponse
import io
from django.views.decorators.csrf import csrf_exempt
import threading

# Global variables
MAIN_COMMENTS_FILE = None
EXTRACTED_COMMENTS_HASHES = set()  # Track comment hashes to avoid duplicates
SESSION_START_TIME = datetime.now().strftime("%Y%m%d_%H%M%S")
ALL_COMMENTS_DATA = []  # Store all comments for JSON export

# Progress tracking (simple single-user approach)
current_progress = {
    'status': 'idle',
    'step': '',
    'percentage': 0,
    'message': '',
    'comments_count': 0,
    'is_running': False,
    'results': None,
    'error': None
}

def update_progress(step, percentage, message, comments_count=None):
    """Update progress for all users (simple single-user approach)"""
    global current_progress
    current_progress.update({
        'status': 'running',
        'step': step,
        'percentage': min(percentage, 95),  # Never reach 100% until complete
        'message': message,
        'is_running': True,
        'error': None
    })
    if comments_count is not None:
        current_progress['comments_count'] = comments_count
    # Minimal print - only show important updates
    if step in ['Setting up browser', 'Logging into Facebook', 'Navigating to post', 
                'Extracting data', 'Creating files', 'Taking screenshot', 'Finalizing']:
        print(f"[{step}] {percentage}% - {message}")

def reset_progress():
    """Reset progress"""
    global current_progress
    if current_progress['results']:
        current_progress['results'] = None
    current_progress = {
        'status': 'idle',
        'step': '',
        'percentage': 0,
        'message': '',
        'comments_count': 0,
        'is_running': False,
        'results': None,
        'error': None
    }

def set_progress_results(results):
    """Set final results"""
    global current_progress
    current_progress.update({
        'results': results,
        'status': 'complete',
        'percentage': 100,
        'step': 'Complete',
        'message': 'Process completed successfully!',
        'comments_count': results.get('extracted_comments', 0),
        'is_running': False
    })

def set_progress_error(error_msg):
    """Set error state"""
    global current_progress
    current_progress.update({
        'error': error_msg,
        'status': 'error',
        'message': f'Error: {error_msg}',
        'is_running': False
    })

@csrf_exempt
def check_progress(request):
    """API endpoint to check progress"""
    return JsonResponse(current_progress)

def get_desktop_path():
    """Get the desktop path based on OS"""
    if os.name == 'nt':  # Windows
        return os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    else:  # macOS or Linux
        return os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop')

DESKTOP_PATH = get_desktop_path()

def init_main_file(session_id=None):
    """Initialize the main comments file on Desktop"""
    global MAIN_COMMENTS_FILE
    if session_id is None:
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create filename for desktop
    filename = f"facebook_comments_{session_id}.txt"
    MAIN_COMMENTS_FILE = os.path.join(DESKTOP_PATH, filename)
    
    # Create file with header if it doesn't exist
    if not os.path.exists(MAIN_COMMENTS_FILE):
        try:
            with open(MAIN_COMMENTS_FILE, 'w', encoding='utf-8') as f:
                f.write("="*60 + "\n")
                f.write("FACEBOOK COMMENTS EXTRACTION\n")
                f.write(f"Session ID: {session_id}\n")
                f.write(f"Started at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("="*60 + "\n\n")
        except Exception as e:
            # Fallback to current directory
            filename = f"facebook_comments_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            MAIN_COMMENTS_FILE = os.path.join(os.getcwd(), filename)
    
    return MAIN_COMMENTS_FILE

def get_comment_hash(comment):
    """Generate hash for a comment to detect duplicates"""
    text = comment.get('text', '').strip().lower()
    commenter = comment.get('commenter', '').strip().lower()
    return hashlib.md5(f"{text}|{commenter}".encode()).hexdigest()

def save_comments_to_main_file(comments, checkpoint=False):
    """Save comments to the main file (append new ones only)"""
    global EXTRACTED_COMMENTS_HASHES, MAIN_COMMENTS_FILE, ALL_COMMENTS_DATA
    
    if not MAIN_COMMENTS_FILE:
        MAIN_COMMENTS_FILE = init_main_file()
    
    new_comments = []
    for comment in comments:
        comment_hash = get_comment_hash(comment)
        if comment_hash not in EXTRACTED_COMMENTS_HASHES:
            EXTRACTED_COMMENTS_HASHES.add(comment_hash)
            new_comments.append(comment)
            # Also add to ALL_COMMENTS_DATA for JSON export
            ALL_COMMENTS_DATA.append({
                'commenter': comment.get('commenter', 'Unknown'),
                'text': comment.get('text', ''),
                'timestamp': comment.get('timestamp', ''),
                'extraction_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'hash': comment_hash
            })
    
    if not new_comments:
        return len(EXTRACTED_COMMENTS_HASHES)
    
    try:
        mode = 'a' if os.path.exists(MAIN_COMMENTS_FILE) else 'w'
        
        with open(MAIN_COMMENTS_FILE, mode, encoding='utf-8') as f:
            if checkpoint and mode == 'a':
                f.write("\n" + "-"*60 + "\n")
                f.write(f"[Checkpoint: {datetime.now().strftime('%H:%M:%S')}]\n")
                f.write(f"Added {len(new_comments)} new comments\n")
                f.write("-"*60 + "\n\n")
            
            # Count existing comments to get starting number
            start_number = 1
            if os.path.exists(MAIN_COMMENTS_FILE) and mode == 'a':
                try:
                    with open(MAIN_COMMENTS_FILE, 'r', encoding='utf-8') as rf:
                        content = rf.read()
                        # Count "COMMENT #" occurrences
                        start_number = content.count('COMMENT #') + 1
                except:
                    pass
            
            # Write new comments
            for i, comment in enumerate(new_comments, start_number):
                f.write(f"COMMENT #{i}:\n")
                f.write("-"*30 + "\n")
                f.write(f"User: {comment.get('commenter', 'Unknown')}\n")
                if comment.get('timestamp'):
                    f.write(f"Time: {comment['timestamp']}\n")
                f.write("-"*30 + "\n")
                f.write(comment.get('text', '') + "\n")
                f.write("="*40 + "\n\n")
            
            # Update footer
            f.write(f"\n[Updated at: {datetime.now().strftime('%H:%M:%S')}] ")
            f.write(f"Total unique comments: {len(EXTRACTED_COMMENTS_HASHES)}\n")
        
        return len(EXTRACTED_COMMENTS_HASHES)
        
    except Exception as e:
        # Try emergency save
        try:
            emergency_file = os.path.join(DESKTOP_PATH, f"emergency_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            with open(emergency_file, 'w', encoding='utf-8') as ef:
                for comment in new_comments:
                    ef.write(f"{comment.get('text', '')}\n")
        except:
            pass
        return len(EXTRACTED_COMMENTS_HASHES)

def save_comments_to_json(all_comments, session_id=None):
    """Save all comments to JSON file on Desktop"""
    if session_id is None:
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    json_filename = os.path.join(DESKTOP_PATH, f"facebook_comments_{session_id}.json")
    
    try:
        with open(json_filename, 'w', encoding='utf-8') as f:
            json.dump(all_comments, f, ensure_ascii=False, indent=2)
        return json_filename
    except Exception as e:
        return None

def create_summary_file(data, all_comments, session_id):
    """Create summary file on Desktop"""
    summary_filename = os.path.join(DESKTOP_PATH, f"facebook_extraction_summary_{session_id}.txt")
    
    try:
        with open(summary_filename, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("FACEBOOK DATA EXTRACTION - COMPLETE SUMMARY\n")
            f.write("="*80 + "\n\n")
            
            f.write(f"EXTRACTION DETAILS:\n")
            f.write("-"*40 + "\n")
            f.write(f"URL: {data.get('url', 'N/A')}\n")
            f.write(f"Extraction Date: {data.get('timestamp', 'N/A')}\n")
            f.write(f"Session ID: {session_id}\n")
            f.write(f"Total Comments Extracted: {len(all_comments)}\n")
            f.write("-"*40 + "\n\n")
            
            f.write("ENGAGEMENT METRICS:\n")
            f.write("-"*40 + "\n")
            f.write(f"Likes: {data.get('likes', 'N/A')}\n")
            f.write(f"Comments: {data.get('comments', 'N/A')}\n")
            f.write(f"Shares: {data.get('shares', 'N/A')}\n")
            f.write(f"Reads/Views: {data.get('reads', 'N/A')}\n")
            f.write("-"*40 + "\n\n")
            
            f.write("POST CONTENT:\n")
            f.write("-"*40 + "\n")
            post_text = data.get('post_text', '')
            if post_text:
                f.write(post_text + "\n")
            else:
                f.write("No post text extracted\n")
            f.write("-"*40 + "\n\n")
            
            f.write("DATA FILES:\n")
            f.write("-"*40 + "\n")
            f.write(f"1. JSON File: facebook_comments_{session_id}.json\n")
            f.write(f"   - Contains all {len(all_comments)} comments in structured format\n")
            f.write(f"   - Location: {DESKTOP_PATH}\n")
            f.write(f"2. Text File: facebook_comments_{session_id}.txt\n")
            f.write(f"   - Human-readable format of all comments\n")
            f.write(f"3. This Summary File: facebook_extraction_summary_{session_id}.txt\n")
            f.write("-"*40 + "\n\n")
            
            f.write("COMMENTS ANALYSIS READY:\n")
            f.write("-"*40 + "\n")
            f.write("The JSON file is formatted for easy analysis with:\n")
            f.write("• Python (pandas, json libraries)\n")
            f.write("• Excel (import JSON data)\n")
            f.write("• Data visualization tools\n")
            f.write("• Natural Language Processing (NLP)\n")
            f.write("-"*40 + "\n")
        
        return summary_filename
        
    except Exception as e:
        return None

def human_type(element, text, field_name=""):
    """Type like a human with random delays"""
    try:
        element.click()
        time.sleep(random.uniform(0.1, 0.3))
        element.send_keys(Keys.CONTROL, 'a')
        time.sleep(random.uniform(0.1, 0.2))
        element.send_keys(Keys.BACKSPACE)
        time.sleep(random.uniform(0.2, 0.4))
        
        for char in text:
            element.send_keys(char)
            time.sleep(random.uniform(0.05, 0.15))
        
    except Exception as e:
        try:
            element.parent.execute_script(f"arguments[0].value = '{text}';", element)
            element.parent.execute_script("arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", element)
        except:
            pass

def human_mouse_move(driver, element):
    """Move mouse like a human"""
    try:
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", element)
        time.sleep(random.uniform(0.3, 0.7))
        
        actions = ActionChains(driver)
        actions.move_to_element(element).perform()
        time.sleep(random.uniform(0.2, 0.5))
        
    except Exception as e:
        try:
            element.click()
        except:
            pass

def random_human_delay(min_seconds=0.5, max_seconds=2.0):
    """Random delay"""
    delay = random.uniform(min_seconds, max_seconds)
    time.sleep(delay)
    return delay

def handle_cookie_popups(driver):
    """Handle cookie consent popups"""
    overlay_selectors = [
        (By.XPATH, "//button[contains(text(), 'Allow') or contains(text(), 'Accept') or contains(text(), 'OK')]"),
        (By.XPATH, "//button[contains(text(), 'Only allow essential cookies')]"),
        (By.XPATH, "//button[contains(text(), 'Decline optional cookies')]"),
        (By.XPATH, "//button[@aria-label='Close' or contains(@aria-label, 'close')]"),
    ]
    
    for by_method, selector in overlay_selectors:
        try:
            elements = driver.find_elements(by_method, selector)
            for elem in elements:
                try:
                    if elem.is_displayed():
                        try:
                            elem.click()
                        except:
                            driver.execute_script("arguments[0].click();", elem)
                        time.sleep(1)
                        return True
                except:
                    continue
        except:
            continue
    
    return False

def handle_save_password_popup(driver):
    """Handle Chrome save password popup"""
    try:
        time.sleep(1)
        actions = ActionChains(driver)
        actions.send_keys(Keys.ESCAPE).perform()
        time.sleep(0.5)
        return True
    except:
        return False

def click_element_safely(driver, element, element_name=""):
    """Safely click an element"""
    try:
        try:
            element.click()
            return True
        except:
            pass
        
        try:
            driver.execute_script("arguments[0].click();", element)
            return True
        except:
            pass
        
        try:
            actions = ActionChains(driver)
            actions.move_to_element(element).click().perform()
            return True
        except:
            pass
        
        return False
        
    except Exception as e:
        return False

def click_most_relevant_then_all_comments(driver):
    """Click 'Most relevant' dropdown then 'All comments' option"""
    max_attempts = 3
    for attempt in range(max_attempts):
        try:
            most_relevant_selectors = [
                (By.XPATH, "//span[contains(text(), 'Most relevant')]"),
                (By.XPATH, "//span[contains(text(), 'most relevant')]"),
                (By.XPATH, "//div[contains(text(), 'Most relevant')]"),
                (By.XPATH, "//div[contains(text(), 'most relevant')]"),
                (By.XPATH, "//*[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'most relevant')]"),
            ]
            
            most_relevant_element = None
            for by_method, selector in most_relevant_selectors:
                try:
                    elements = driver.find_elements(by_method, selector)
                    for elem in elements:
                        if elem.is_displayed():
                            most_relevant_element = elem
                            break
                    if most_relevant_element:
                        break
                except:
                    continue
            
            if most_relevant_element:
                if click_element_safely(driver, most_relevant_element, "Most relevant"):
                    time.sleep(2)
                    
                    all_comments_selectors = [
                        (By.XPATH, "//span[contains(text(), 'All comments')]"),
                        (By.XPATH, "//span[contains(text(), 'all comments')]"),
                        (By.XPATH, "//div[contains(text(), 'All comments')]"),
                        (By.XPATH, "//div[contains(text(), 'all comments')]"),
                        (By.XPATH, "//*[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'all comments')]"),
                    ]
                    
                    for by_method, selector in all_comments_selectors:
                        try:
                            elements = driver.find_elements(by_method, selector)
                            for elem in elements:
                                if elem.is_displayed():
                                    if click_element_safely(driver, elem, "All comments"):
                                        time.sleep(3)
                                        return True
                        except:
                            continue
            
            time.sleep(1)
            
        except Exception as e:
            time.sleep(1)
    
    return False

def is_session_active(driver):
    """Check if the browser session is still active"""
    try:
        driver.current_url
        return True
    except (InvalidSessionIdException, WebDriverException):
        return False
    except:
        return False

def extract_current_comments(driver):
    """Extract comments currently visible on page"""
    current_comments = []
    
    if not is_session_active(driver):
        return current_comments
    
    try:
        comment_selectors = [
            (By.XPATH, "//div[contains(@class, 'x1y1aw1k') or contains(@class, 'x1lliihq')]//div[@dir='auto']"),
            (By.XPATH, "//div[@role='article']//div[@dir='auto' and string-length(text()) > 5]"),
            (By.XPATH, "//div[@data-commentid]//div[@dir='auto']"),
            (By.XPATH, "//div[contains(@class, 'comment')]//div[@dir='auto']"),
            (By.XPATH, "//div[@dir='auto' and string-length(text()) > 10]"),
        ]
        
        seen_texts_in_current = set()
        
        for by_method, selector in comment_selectors:
            try:
                if not is_session_active(driver):
                    break
                    
                comment_elements = driver.find_elements(by_method, selector)
                
                for elem in comment_elements:
                    try:
                        if not is_session_active(driver):
                            break
                            
                        comment_text = elem.text.strip()
                        
                        if (comment_text and 
                            len(comment_text) > 10 and
                            comment_text not in seen_texts_in_current and
                            not any(x in comment_text.lower() for x in [
                                'like', 'reply', 'share', 'comment', 
                                'most relevant', 'all comments',
                                'show more comments', 'view more comments',
                                'write a comment', 'post a comment'
                            ])):
                            
                            # Get commenter name
                            commenter = "Unknown"
                            try:
                                ancestor = elem.find_element(By.XPATH, "./ancestor::div[contains(@class, 'x1y1aw1k') or contains(@class, 'x1lliihq') or @data-commentid]")
                                links = ancestor.find_elements(By.TAG_NAME, "a")
                                for link in links:
                                    link_text = link.text.strip()
                                    if link_text and link_text not in comment_text:
                                        commenter = link_text
                                        break
                            except:
                                pass
                            
                            # Get timestamp
                            timestamp = ""
                            try:
                                time_elem = elem.find_element(By.XPATH, 
                                    "./following::a[contains(@aria-label, 'Posted') or contains(@aria-label, 'Comment')]"
                                )
                                timestamp = time_elem.get_attribute('aria-label') or time_elem.text
                            except:
                                pass
                            
                            comment_data = {
                                'commenter': commenter,
                                'text': comment_text,
                                'timestamp': timestamp
                            }
                            
                            current_comments.append(comment_data)
                            seen_texts_in_current.add(comment_text)
                            
                    except StaleElementReferenceException:
                        continue
                    except Exception as e:
                        if "invalid session" in str(e).lower():
                            break
                        continue
                        
                if not is_session_active(driver):
                    break
                    
            except Exception as e:
                if "invalid session" in str(e).lower():
                    break
                continue
        
        return current_comments
        
    except Exception as e:
        return current_comments

def click_all_show_more_comments(driver):
    """Click ALL 'Show more comments' buttons and save to single file"""
    total_clicks = 0
    consecutive_no_buttons = 0
    max_consecutive_no_buttons = 3
    
    while True:
        if not is_session_active(driver):
            break
        
        try:
            show_more_selectors = [
                (By.XPATH, "//span[contains(text(), 'Show more comments')]"),
                (By.XPATH, "//span[contains(text(), 'View more comments')]"),
                (By.XPATH, "//span[contains(text(), 'See more comments')]"),
                (By.XPATH, "//span[contains(text(), 'Load more comments')]"),
            ]
            
            found_buttons = []
            for by_method, selector in show_more_selectors:
                try:
                    elements = driver.find_elements(by_method, selector)
                    for elem in elements:
                        try:
                            if elem.is_displayed() and elem.is_enabled():
                                found_buttons.append(elem)
                        except:
                            continue
                except:
                    continue
            
            # Remove duplicates
            unique_buttons = []
            seen_ids = set()
            for btn in found_buttons:
                try:
                    btn_id = btn.id
                    if btn_id not in seen_ids:
                        seen_ids.add(btn_id)
                        unique_buttons.append(btn)
                except:
                    unique_buttons.append(btn)
            
            if len(unique_buttons) == 0:
                consecutive_no_buttons += 1
                
                if consecutive_no_buttons >= max_consecutive_no_buttons:
                    break
                
                try:
                    current_comments = extract_current_comments(driver)
                    if current_comments:
                        comments_count = save_comments_to_main_file(current_comments, checkpoint=False)
                        update_progress('Extracting data', 60, f'Extracting comments... Found {comments_count} comments so far', comments_count)
                except:
                    pass
                
                try:
                    driver.execute_script("window.scrollBy(0, 500);")
                    time.sleep(2)
                except:
                    break
                continue
            
            consecutive_no_buttons = 0
            
            clicked_any = False
            for i, btn in enumerate(unique_buttons):
                try:
                    if not is_session_active(driver):
                        break
                    
                    driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", btn)
                    time.sleep(0.5)
                    
                    if btn.is_displayed() and btn.is_enabled():
                        if click_element_safely(driver, btn, f"Show more comments {i+1}"):
                            total_clicks += 1
                            clicked_any = True
                            time.sleep(2)
                            
                            try:
                                new_comments = extract_current_comments(driver)
                                if new_comments:
                                    comments_count = save_comments_to_main_file(new_comments, checkpoint=False)
                                    update_progress('Extracting data', 60, f'Extracting comments... Found {comments_count} comments so far', comments_count)
                            except:
                                pass
                            
                            try:
                                driver.execute_script("window.scrollBy(0, 300);")
                                time.sleep(0.5)
                            except:
                                pass
                except:
                    continue
            
            if not clicked_any:
                consecutive_no_buttons += 1
                
                try:
                    new_comments = extract_current_comments(driver)
                    if new_comments:
                        comments_count = save_comments_to_main_file(new_comments, checkpoint=False)
                        update_progress('Extracting data', 60, f'Extracting comments... Found {comments_count} comments so far', comments_count)
                except:
                    pass
                
                if consecutive_no_buttons >= max_consecutive_no_buttons:
                    break
                
                try:
                    driver.execute_script("window.scrollBy(0, 500);")
                    time.sleep(1)
                except:
                    break
                continue
            
        except Exception as e:
            try:
                current_comments = extract_current_comments(driver)
                if current_comments:
                    save_comments_to_main_file(current_comments, checkpoint=True)
            except:
                pass
            
            if "invalid session" in str(e).lower() or "session id" in str(e).lower():
                break
            time.sleep(1)
    
    return total_clicks

def scroll_and_extract_comments(driver):
    """Scroll through page and extract all comments to single file"""
    try:
        last_height = driver.execute_script("return document.body.scrollHeight")
    except:
        return
    
    consecutive_same_height = 0
    max_consecutive_same_height = 3
    scroll_attempts = 0
    max_scroll_attempts = 50
    
    while scroll_attempts < max_scroll_attempts:
        scroll_attempts += 1
        
        if not is_session_active(driver):
            break
        
        try:
            current_comments = extract_current_comments(driver)
            if current_comments:
                comments_count = save_comments_to_main_file(current_comments, checkpoint=False)
                if scroll_attempts % 5 == 0:
                    update_progress('Scrolling', 65, f'Scrolling page... Found {comments_count} comments so far', comments_count)
            
            driver.execute_script("window.scrollBy(0, 800);")
            time.sleep(1.5)
            
            new_height = driver.execute_script("return document.body.scrollHeight")
            
            if new_height == last_height:
                consecutive_same_height += 1
                if consecutive_same_height >= max_consecutive_same_height:
                    break
            else:
                consecutive_same_height = 0
            
            last_height = new_height
            
            try:
                current_pos = driver.execute_script("return window.pageYOffset + window.innerHeight")
                if current_pos >= new_height:
                    break
            except:
                pass
                
        except Exception as e:
            if "invalid session" in str(e).lower():
                break
            time.sleep(1)
            continue
    
    # Final extraction after scrolling
    try:
        if is_session_active(driver):
            final_comments = extract_current_comments(driver)
            if final_comments:
                comments_count = save_comments_to_main_file(final_comments, checkpoint=True)
                update_progress('Scrolling', 70, f'Finalizing extraction... Total: {comments_count} comments', comments_count)
    except:
        pass

def extract_engagement_metrics(driver):
    """Extract Read, Like, Comment, Share counts"""
    metrics = {
        'likes': '0',
        'comments': '0',
        'shares': '0',
        'reads': '0'
    }
    
    try:
        page_source = driver.page_source
        
        likes_match = re.search(r'(\d+(?:\.\d+)?[KMB]?)\s*Likes?', page_source, re.IGNORECASE)
        if likes_match:
            metrics['likes'] = likes_match.group(1)
        
        comments_match = re.search(r'(\d+(?:\.\d+)?[KMB]?)\s*Comments?', page_source, re.IGNORECASE)
        if comments_match:
            metrics['comments'] = comments_match.group(1)
        
        shares_match = re.search(r'(\d+(?:\.\d+)?[KMB]?)\s*Shares?', page_source, re.IGNORECASE)
        if shares_match:
            metrics['shares'] = shares_match.group(1)
        
        views_match = re.search(r'(\d+(?:\.\d+)?[KMB]?)\s*(Views|Seen|Reads?)', page_source, re.IGNORECASE)
        if views_match:
            metrics['reads'] = views_match.group(1)
        
    except:
        pass
    
    return metrics

def extract_post_data_enhanced(driver, target_url):
    """Enhanced extraction - saves to Desktop and creates JSON"""
    global ALL_COMMENTS_DATA
    
    # Initialize variables
    session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    init_main_file(session_id)
    
    # Reset ALL_COMMENTS_DATA for this session
    ALL_COMMENTS_DATA = []
    
    data = {
        'likes': '0',
        'comments': '0',
        'shares': '0',
        'reads': '0',
        'post_text': '',
        'comments_count': 0,
        'comments_file': MAIN_COMMENTS_FILE,
        'json_file': '',
        'summary_file': '',
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'url': target_url,
        'session_id': session_id
    }
    
    try:
        update_progress('Extracting data', 55, 'Loading post content...')
        time.sleep(5)
        
        if not is_session_active(driver):
            return data, []
        
        # ========== EXTRACT POST CONTENT ==========
        post_text = ""
        text_selectors = [
            (By.XPATH, "//div[@dir='auto' and string-length(text()) > 20]"),
            (By.XPATH, "//div[contains(@class, 'x1lliihq') and @dir='auto']"),
            (By.XPATH, "//div[@data-testid='post_message']//div[@dir='auto']"),
            (By.XPATH, "//div[contains(@class, 'userContent')]//div[@dir='auto']"),
        ]
        
        for by_method, selector in text_selectors:
            try:
                elements = driver.find_elements(by_method, selector)
                for elem in elements:
                    text = elem.text.strip()
                    if text and len(text) > 60:
                        post_text = text
                        break
                if post_text:
                    break
            except:
                continue
        
        data['post_text'] = post_text
        
        # ========== EXTRACT ENGAGEMENT METRICS ==========
        metrics = extract_engagement_metrics(driver)
        data.update(metrics)
        
        # ========== CLICK MOST RELEVANT → ALL COMMENTS ==========
        if is_session_active(driver):
            switched = click_most_relevant_then_all_comments(driver)
            data['switched_to_all'] = switched
        
        # ========== CLICK ALL SHOW MORE COMMENTS ==========
        update_progress('Extracting data', 60, 'Loading comments... Found 0 comments so far', 0)
        if is_session_active(driver):
            total_clicks = click_all_show_more_comments(driver)
            data['total_clicks'] = total_clicks
        else:
            data['total_clicks'] = 0
        
        # ========== SCROLL AND EXTRACT ALL COMMENTS ==========
        if is_session_active(driver):
            scroll_and_extract_comments(driver)
            data['comments_count'] = len(EXTRACTED_COMMENTS_HASHES)
        else:
            data['comments_count'] = len(EXTRACTED_COMMENTS_HASHES)
        
        # ========== CREATE JSON AND SUMMARY FILES ==========
        update_progress('Creating files', 75, 'Saving data to files...', data['comments_count'])
        if ALL_COMMENTS_DATA:
            json_filename = save_comments_to_json(ALL_COMMENTS_DATA, session_id)
            data['json_file'] = json_filename
            
            summary_file = create_summary_file(data, ALL_COMMENTS_DATA, session_id)
            data['summary_file'] = summary_file
        
        return data, ALL_COMMENTS_DATA
        
    except Exception as e:
        return data, []

def take_screenshot_safe(driver):
    """Take screenshot with session checking"""
    try:
        if not is_session_active(driver):
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        screenshot_name = os.path.join(DESKTOP_PATH, f"facebook_screenshot_{timestamp}.png")
        driver.save_screenshot(screenshot_name)
        return screenshot_name
    except:
        return None

def catch_url(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    return render(request, 'catchinglink.html')

def download_shared_file(request, file_id):
    if not check_access(request):
        return redirect('logout')
    
    """Serve files for download using file_id"""
    try:
        # Get file from database using file_id
        shared_file = SharedFile.objects.get(id=file_id)
        
        # Check if user has permission to download
        if not (shared_file.uploaded_by == request.user or 
                request.user in shared_file.shared_with.all()):
            return HttpResponse("Access denied", status=403)
        
        # Get the actual file path from FileField
        file_path = shared_file.file.path  # Assuming 'file' is a FileField
        file_name = shared_file.file.name  # Get filename from database
        
        print(f"Downloading: {file_name} (ID: {file_id})")
        
        if not os.path.exists(file_path):
            return HttpResponse(f"File not found: {file_name}", status=404)
        
        # Determine content type from file extension
        import mimetypes
        content_type, encoding = mimetypes.guess_type(file_path)
        if not content_type:
            content_type = 'application/octet-stream'
        
        # Increment download count if you have that field
        if hasattr(shared_file, 'download_count'):
            shared_file.download_count += 1
            shared_file.save()
        
        # Serve the file
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        response = HttpResponse(file_content, content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_name)}"'
        response['Content-Length'] = len(file_content)
        
        return response
        
    except SharedFile.DoesNotExist:
        return HttpResponse(f"File with ID {file_id} not found", status=404)
    except Exception as e:
        return HttpResponse(f"Error downloading file: {str(e)}", status=500)
    
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.template.loader import render_to_string
from datetime import datetime
import os
import tempfile
import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

@login_required
def sentiment_report(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Create and download sentiment reports directly from form data"""
    form_data = {}
    
    if request.method == 'POST':
        # Get form data
        form_data = {
            'title': request.POST.get('title', '').strip(),
            'description': request.POST.get('description', '').strip(),
            'positive_percentage': request.POST.get('positive_percentage', '89.2'),
            'negative_percentage': request.POST.get('negative_percentage', '6.9'),
            'neutral_percentage': request.POST.get('neutral_percentage', '3.9'),
            'jay_nepal_likes': request.POST.get('jay_nepal_likes', '70000'),
            'jay_nepal_shares': request.POST.get('jay_nepal_shares', '13800'),
            'jay_nepal_comments': request.POST.get('jay_nepal_comments', '2000'),
        }
        
        # Validate required fields
        if not form_data['title']:
            messages.error(request, 'Report title is required.')
            return render(request, 'sentiment_reports.html', {
                'form_data': form_data
            })
        
        try:
            # Convert to proper types
            positive = float(form_data['positive_percentage'])
            negative = float(form_data['negative_percentage'])
            neutral = float(form_data['neutral_percentage'])
            likes = int(form_data['jay_nepal_likes'])
            shares = int(form_data['jay_nepal_shares'])
            comments = int(form_data['jay_nepal_comments'])
            
            # Ensure percentages sum to 100
            total_percent = positive + negative + neutral
            if abs(total_percent - 100) > 0.1:  # Allow small rounding errors
                # Normalize to 100%
                factor = 100 / total_percent
                positive = round(positive * factor, 1)
                negative = round(negative * factor, 1)
                neutral = round(neutral * factor, 1)
            
        except (ValueError, TypeError) as e:
            messages.error(request, f'Invalid number format: {str(e)}')
            return render(request, 'sentiment_reports.html', {
                'form_data': form_data
            })
        
        # Handle image upload
        image_base64 = None
        image_mime_type = None
        
        if 'image' in request.FILES:
            uploaded_file = request.FILES['image']
            
            # Validate file type
            allowed_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
            file_ext = os.path.splitext(uploaded_file.name)[1].lower()
            
            if file_ext not in allowed_extensions:
                messages.error(request, 'Invalid image format. Please upload JPG, PNG, GIF, BMP, or WebP.')
                return render(request, 'sentiment_reports.html', {
                    'form_data': form_data
                })
            
            # Validate file size (max 5MB)
            if uploaded_file.size > 5 * 1024 * 1024:
                messages.error(request, 'Image size must be less than 5MB.')
                return render(request, 'sentiment_reports.html', {
                    'form_data': form_data
                })
            
            # Read image data and convert to base64
            image_data = uploaded_file.read()
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # Determine mime type based on file extension
            if file_ext in ['.jpg', '.jpeg']:
                image_mime_type = 'image/jpeg'
            elif file_ext == '.png':
                image_mime_type = 'image/png'
            elif file_ext == '.gif':
                image_mime_type = 'image/gif'
            elif file_ext == '.bmp':
                image_mime_type = 'image/bmp'
            elif file_ext == '.webp':
                image_mime_type = 'image/webp'
            else:
                image_mime_type = 'image/jpeg'  # default
        
        # Get current datetime once
    
        current_time = datetime.now()
        
        # Calculate gradient positions for pie chart
        positive_end = positive
        negative_end = positive + negative
        
        # Prepare data for report
        report_data = {
            'title': form_data['title'],
            'description': form_data['description'],
            'positive_percentage': round(positive, 1),
            'negative_percentage': round(negative, 1),
            'neutral_percentage': round(neutral, 1),
            'jay_nepal_likes': likes,
            'jay_nepal_shares': shares,
            'jay_nepal_comments': comments,
            'total_engagement': likes + shares + comments,
            'supporters_count': likes,
            'critics_count': comments // 10,
            'created_at': current_time,
            'created_by': request.user.get_full_name() or request.user.username,
            'rank': request.user.rank if hasattr(request.user, 'rank') else 'N/A',
            # Pie chart data for CSS conic gradient
            'positive_end': positive_end,
            'negative_end': negative_end,
        }
        
        # Prepare context for template
        context = {
            'report': report_data,
            'now': current_time,
            'user': request.user,
            'image_base64': image_base64,
            'image_mime_type': image_mime_type,
        }
        
        # Generate HTML report
        return generate_html_report(request, context)
    
    # GET request - show empty form
    return render(request, 'sentiment_reports.html', {
        'form_data': form_data
    })

def generate_html_report(request, context):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Generate and return HTML report"""
    # Render HTML template
    html_string = render_to_string('report_template.html', context)
    
    # Create HTTP response
    response = HttpResponse(html_string, content_type='text/html')
    filename = f"sentiment_report_{context['report']['title'].replace(' ', '_')}_{context['now'].strftime('%Y%m%d_%H%M%S')}.html"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

from utils.websocket_helper import send_to_websocket

def websocket_test(request):
    print("socket test page rendering")
    """Simple view to test WebSocket connection"""
    # Send your exact message
    send_to_websocket("🚀 TECHPANA SCRAPER STARTED (USER-SPECIFIC, NO LIMITS)")
    
    return render(request, 'websocket_test.html')

# views.py - COMPLETE CORRECTED VERSION
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from datetime import datetime, timedelta
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Count, Q
from django.http import JsonResponse
from django.db.models.functions import TruncDate
from .models import ThreatAlert, ThreatCategory
import json
from django.utils.timesince import timesince

def news_central(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # Get all distinct categories for filter
    all_categories = ThreatCategory.objects.all().distinct()
    
    # Get selected categories from GET parameters
    selected_categories = request.GET.getlist('category')
    
    # Date range filtering
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    
    # Parse dates correctly
    start_date = None
    end_date = None
    
    if start_date_str:
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        except ValueError:
            start_date = None
    
    if end_date_str:
        try:
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except ValueError:
            end_date = None
    
    # Severity filter
    severity_filter = request.GET.get('severity')
    
    # Sort ordering
    sort_by = request.GET.get('sort', '-created_at')
    valid_sort_fields = ['-created_at', 'created_at', '-severity', 'severity', '-likes_count', 'likes_count']
    if sort_by not in valid_sort_fields:
        sort_by = '-created_at'
    
    # Search query
    search_query = request.GET.get('search', '').strip()
    
    # Build base query
    threats = ThreatAlert.objects.all().select_related('category', 'created_by')
    
    # Apply search filter
    if search_query:
        threats = threats.filter(
            Q(title__icontains=search_query) |
            Q(content__icontains=search_query) |
            Q(source__icontains=search_query) |
            Q(category__name__icontains=search_query)
        )
    
    # Apply category filter
    if selected_categories:
        threats = threats.filter(category_id__in=selected_categories)
    
    # Apply severity filter
    if severity_filter and severity_filter in ['low', 'medium', 'high', 'critical']:
        threats = threats.filter(severity=severity_filter)
    
    # Apply date range filter - FIXED VERSION
    if start_date:
        # Create datetime at start of the day (00:00:00)
        start_datetime = datetime.combine(start_date, datetime.min.time())
        # Make it timezone aware
        start_datetime = timezone.make_aware(start_datetime)
        threats = threats.filter(created_at__gte=start_datetime)
    
    if end_date:
        # Create datetime at end of the day (23:59:59.999999)
        end_datetime = datetime.combine(end_date, datetime.max.time())
        # Make it timezone aware
        end_datetime = timezone.make_aware(end_datetime)
        threats = threats.filter(created_at__lte=end_datetime)  # Use lte, not lt
    
    # Apply sorting
    threats = threats.order_by(sort_by)
    
    # Get counts for stats
    total_count = threats.count()
    critical_count = threats.filter(severity='critical').count()
    high_count = threats.filter(severity='high').count()
    
    # Default date range (last 30 days)
    default_end_date = timezone.now().date()
    default_start_date = default_end_date - timedelta(days=30)
    
    # Format dates for template
    start_date_for_template = start_date_str if start_date_str else default_start_date.strftime('%Y-%m-%d')
    end_date_for_template = end_date_str if end_date_str else default_end_date.strftime('%Y-%m-%d')
    
    # Regular page request - get paginated data
    page = request.GET.get('page', 1)
    rows_per_page = int(request.GET.get('rows', 25))
    
    paginator = Paginator(threats, rows_per_page)
    
    try:
        threats_page = paginator.page(page)
    except PageNotAnInteger:
        threats_page = paginator.page(1)
    except EmptyPage:
        threats_page = paginator.page(paginator.num_pages)
    
    context = {
        'threats': threats_page,
        'all_categories': all_categories,
        'selected_categories': selected_categories,
        'start_date': start_date_for_template,
        'end_date': end_date_for_template,
        'default_start_date': default_start_date.strftime('%Y-%m-%d'),
        'default_end_date': default_end_date.strftime('%Y-%m-%d'),
        'severity_filter': severity_filter or '',
        'search_query': search_query,
        'sort_by': sort_by,
        'total_count': total_count,
        'critical_count': critical_count,
        'high_count': high_count,
        'rows_per_page': rows_per_page,
    }
    
    return render(request, 'news_central.html', context)

@login_required
def like_alert(request, alert_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Facebook-style like system with mutual exclusivity
    - Like removes any existing unlike
    - User can only either like or unlike, not both
    """
    if request.method == 'POST':
        try:
            alert = get_object_or_404(ThreatAlert, id=alert_id)
            user = request.user
            
            # Check current state
            is_liked = user in alert.users_liked.all()
            is_unliked = user in alert.users_unliked.all()
            
            response_data = {
                'success': True,
                'liked': False,
                'unliked': False,
                'likes_count': alert.likes_count,
                'unlikes_count': alert.unlikes_count,
                'user_liked': False,
                'user_unliked': False
            }
            
            if is_liked:
                # User already liked - remove like
                alert.users_liked.remove(user)
                alert.likes_count = max(0, alert.likes_count - 1)
                response_data['liked'] = False
                response_data['user_liked'] = False
            else:
                # User hasn't liked - add like
                alert.users_liked.add(user)
                alert.likes_count += 1
                response_data['liked'] = True
                response_data['user_liked'] = True
                
                # Remove unlike if it exists (mutual exclusivity)
                if is_unliked:
                    alert.users_unliked.remove(user)
                    alert.unlikes_count = max(0, alert.unlikes_count - 1)
                    response_data['user_unliked'] = False
            
            alert.save()
            
            # Get fresh counts and state
            alert.refresh_from_db()
            response_data['likes_count'] = alert.likes_count
            response_data['unlikes_count'] = alert.unlikes_count
            
            # Check actual state after save
            response_data['user_liked'] = user in alert.users_liked.all()
            response_data['user_unliked'] = user in alert.users_unliked.all()
            
            return JsonResponse(response_data)
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Invalid request method. Use POST.'
    }, status=405)

@login_required
def unlike_alert(request, alert_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    Facebook-style unlike system with mutual exclusivity
    - Unlike removes any existing like
    - User can only either like or unlike, not both
    """
    if request.method == 'POST':
        try:
            alert = get_object_or_404(ThreatAlert, id=alert_id)
            user = request.user
            
            # Check current state
            is_unliked = user in alert.users_unliked.all()
            is_liked = user in alert.users_liked.all()
            
            response_data = {
                'success': True,
                'liked': False,
                'unliked': False,
                'likes_count': alert.likes_count,
                'unlikes_count': alert.unlikes_count,
                'user_liked': False,
                'user_unliked': False
            }
            
            if is_unliked:
                # User already unliked - remove unlike
                alert.users_unliked.remove(user)
                alert.unlikes_count = max(0, alert.unlikes_count - 1)
                response_data['unliked'] = False
                response_data['user_unliked'] = False
            else:
                # User hasn't unliked - add unlike
                alert.users_unliked.add(user)
                alert.unlikes_count += 1
                response_data['unliked'] = True
                response_data['user_unliked'] = True
                
                # Remove like if it exists (mutual exclusivity)
                if is_liked:
                    alert.users_liked.remove(user)
                    alert.likes_count = max(0, alert.likes_count - 1)
                    response_data['user_liked'] = False
            
            alert.save()
            
            # Get fresh counts and state
            alert.refresh_from_db()
            response_data['likes_count'] = alert.likes_count
            response_data['unlikes_count'] = alert.unlikes_count
            
            # Check actual state after save
            response_data['user_liked'] = user in alert.users_liked.all()
            response_data['user_unliked'] = user in alert.users_unliked.all()
            
            return JsonResponse(response_data)
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    
    return JsonResponse({
        'success': False,
        'error': 'Invalid request method. Use POST.'
    }, status=405)

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from .models import User  # Import your custom User model

@login_required
def user_list(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """
    List all users with search and pagination
    Only accessible to Admin and SuperAdmin
    """
    # Check user permission
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # Get all users
    users = User.objects.filter(is_void=False).order_by('-created_at')
    
    # Search functionality
    search_query = request.GET.get('search', '').strip()
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(name__icontains=search_query) |
            Q(phone__icontains=search_query) |
            Q(unit__icontains=search_query)
        )
    
    # Role filter
    role_filter = request.GET.get('role', '')
    if role_filter:
        users = users.filter(role=role_filter)
    
    # Pagination
    page = request.GET.get('page', 1)
    paginator = Paginator(users, 10)
    
    try:
        users_page = paginator.page(page)
    except PageNotAnInteger:
        users_page = paginator.page(1)
    except EmptyPage:
        users_page = paginator.page(paginator.num_pages)
    
    # Get counts
    total_users = users.count()
    user_count_by_role = {
        'User': users.filter(role='User').count(),
        'CyberUser': users.filter(role='CyberUser').count(),
        'Admin': users.filter(role='Admin').count(),
        'SuperAdmin': users.filter(role='SuperAdmin').count(),
    }
    
    context = {
        'users': users_page,
        'search_query': search_query,
        'role_filter': role_filter,
        'total_users': total_users,
        'user_count_by_role': user_count_by_role,
        'role_choices': User.ROLE_CHOICES,
        'paginator': paginator,
    }
    
    return render(request, 'user_list.html', context)

from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, redirect
from django.views.decorators.http import require_http_methods

@login_required
@require_http_methods(["GET", "POST"])
def add_user(request):
    if not check_access(request):
        return redirect('logout')  # Create this view

    """Add new user"""
    if request.user.role not in ['Admin', 'SuperAdmin']:
        messages.error(request, "Permission denied.")
        return redirect('user_list')
    
    # Get role choices from model
    role_choices = User.ROLE_CHOICES
    
    if request.method == 'POST':
        # Get form data
        username = request.POST.get('username')
        email = request.POST.get('email')
        first_name = request.POST.get('first_name', '')
        last_name = request.POST.get('last_name', '')
        phone = request.POST.get('phone', '')
        unit = request.POST.get('unit', '')
        rank = request.POST.get('rank', '')
        role = request.POST.get('role', 'User')
        password = request.POST.get('password')
        password2 = request.POST.get('password2')
        social_media = request.POST.get('social_media') == 'on'
        is_active = request.POST.get('is_active') == 'on'
        is_void = request.POST.get('is_void') == 'on'
        
        # Validate required fields
        if not all([username, email, password, password2]):
            messages.error(request, "Please fill all required fields.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # Check if username already exists
        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # Check if email already exists
        if User.objects.filter(email=email).exists():
            messages.error(request, "Email already exists.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # Validate passwords match
        if password != password2:
            messages.error(request, "Passwords do not match.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # Validate password strength
        if len(password) < 8:
            messages.error(request, "Password must be at least 8 characters long.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # Restrict Admin/SuperAdmin creation based on current user's role
        current_user_role = request.user.role
        
        # Regular Admin can only create User or CyberUser
        if current_user_role == 'Admin' and role in ['Admin', 'SuperAdmin']:
            messages.error(request, "Admin can only create User or CyberUser roles.")
            return render(request, 'add_user.html', {
                'role_choices': role_choices
            })
        
        # SuperAdmin can create any role, others default to User
        if current_user_role != 'SuperAdmin':
            if role in ['Admin', 'SuperAdmin']:
                role = 'User'
        
        try:
            # Create new user
            user = User.objects.create(
                username=username,
                email=email,
                first_name=first_name,
                last_name=last_name,
                phone=phone,
                unit=unit,
                rank=rank,
                role=role,
                social_media=social_media,
                is_active=is_active,
                is_void=is_void,
            )
            
            # Set password properly using Django's password hashing
            user.set_password(password)
            user.save()
            
            messages.success(request, f"User '{username}' has been created successfully.")
            return redirect('user_list')
            
        except Exception as e:
            messages.error(request, f"Error creating user: {str(e)}")
            return render(request, 'add_user.html', {
                'role_choices': role_choices,
                'form_data': request.POST
            })
    
    # GET request - show empty form
    return render(request, 'add_user.html', {
        'role_choices': role_choices
    })

@login_required
def edit_user(request, user_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Edit user information"""
    # Check if user has permission to edit (only SuperAdmin or Admin can edit)
    if request.user.role not in ['SuperAdmin', 'Admin']:
        messages.error(request, "You don't have permission to edit users.")
        return redirect('user_list')
    
    user = get_object_or_404(User, id=user_id)
    
    # Prevent editing self's role/status if not SuperAdmin
    if user.id == request.user.id and request.user.role != 'SuperAdmin':
        messages.warning(request, "You cannot edit your own role/status. Contact SuperAdmin.")
        return redirect('user_list')
    
    # Get role choices from model
    role_choices = User.ROLE_CHOICES
    
    if request.method == 'POST':
        # Get form data
        username = request.POST.get('username')
        email = request.POST.get('email')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        phone = request.POST.get('phone')
        unit = request.POST.get('unit')
        rank = request.POST.get('rank')
        role = request.POST.get('role')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        social_media = request.POST.get('social_media') == 'on'
        is_active = request.POST.get('is_active') == 'on'
        is_void = request.POST.get('is_void') == 'on'  # Get is_void checkbox value
        
        # Only SuperAdmin can change is_void status
        if request.user.role != 'SuperAdmin' and is_void != user.is_void:
            messages.error(request, "Only SuperAdmin can change deletion status.")
            return redirect('edit_user', user_id=user_id)
        
        # Prevent self-deletion
        if user.id == request.user.id and is_void:
            messages.error(request, "You cannot delete yourself!")
            return redirect('edit_user', user_id=user_id)
        
        # Validate username uniqueness (if changed)
        if username != user.username:
            if User.objects.filter(username=username).exclude(id=user.id).exists():
                messages.error(request, "Username already exists.")
                return redirect('edit_user', user_id=user_id)
        
        # Validate email uniqueness (if changed)
        if email != user.email:
            if User.objects.filter(email=email).exclude(id=user.id).exists():
                messages.error(request, "Email already exists.")
                return redirect('edit_user', user_id=user_id)
        
        # Validate passwords match if provided
        if password and password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return redirect('edit_user', user_id=user_id)
        
        # Update user fields
        user.username = username
        user.email = email
        user.first_name = first_name
        user.last_name = last_name
        user.phone = phone
        user.unit = unit
        user.rank = rank
        user.role = role
        user.social_media = social_media
        user.is_active = is_active
        
        # Only update is_void if user is SuperAdmin
        if request.user.role == 'SuperAdmin':
            user.is_void = is_void
        
        # Update password if provided
        if password:
            user.set_password(password)
        
        try:
            user.save()
            messages.success(request, f"User '{username}' has been updated successfully.")
            return redirect('user_list')
        except Exception as e:
            messages.error(request, f"Error updating user: {str(e)}")
            return redirect('edit_user', user_id=user_id)
    
    # For GET request, render the edit form
    context = {
        'user': user,
        'role_choices': role_choices,
    }
    
    return render(request, 'edit_user.html', context)

@login_required
def delete_user(request, user_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    print("Deleted page rendering")
    """Delete user"""
    if request.user.role != 'SuperAdmin':
        messages.error(request, "Only SuperAdmin can delete users.")
        return redirect('user_list')
    
    user = get_object_or_404(User, id=user_id)
    
    # Prevent self-deletion
    if user.id == request.user.id:
        messages.error(request, "You cannot delete yourself!")
        return redirect('user_list')
    
    username = user.username
    user.is_void = True
    user.save()  # Save the change
    messages.success(request, f"User '{username}' has been deleted successfully.")
    return redirect('user_list')

@login_required
def sharing_files(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # File type filter options
    file_types = [
        ('pdf', 'PDF'),
        ('doc', 'Word(doc)'),
        ('docx', 'Word(docx)'),
        ('txt', 'Text(txt)'),
        ('jpg', 'Image(jpg)'),
        ('jpeg', 'Image(jpeg)'),
        ('png', 'Image(png)'),
        ('zip', 'Archive(zip)'),
        ('rar', 'Archive(rar)'),
        ('xls', 'Excel(xls)'),
        ('xlsx', 'Excel(xlsr)'),
        ('csv', 'CSV'),
    ]
    
    # Handle file upload (POST request)
    if request.method == 'POST' and request.FILES.get('file'):
        try:
            uploaded_file = request.FILES['file']
            
            # Validate file size (max 10MB = 10 * 1024 * 1024 bytes)
            if uploaded_file.size > 10 * 1024 * 1024:
                messages.error(request, 'File size exceeds 10MB limit.')
                return redirect('sharing_files')
            
            # Get file extension
            file_name = uploaded_file.name
            file_extension = os.path.splitext(file_name)[1].lower().replace('.', '')
            
            # Create SharedFile object
            shared_file = SharedFile.objects.create(
                name=file_name,
                file=uploaded_file,
                uploaded_by=request.user,
                extension=file_extension,
                size=uploaded_file.size
            )
            
            # Handle sharing with users
            share_with = request.POST.get('share_with', '').strip()
            if share_with:
                users_to_share = []
                share_list = [s.strip() for s in share_with.split(',') if s.strip()]
                
                for identifier in share_list:
                    try:
                        # Try to find user by username or email
                        if '@' in identifier:
                            user = User.objects.get(email=identifier)
                        else:
                            user = User.objects.get(username=identifier)
                        
                        if user != request.user:  # Don't share with self
                            users_to_share.append(user)
                    except User.DoesNotExist:
                        messages.warning(request, f'User "{identifier}" not found.')
                        continue
                
                # Add users to shared_with
                if users_to_share:
                    shared_file.shared_with.add(*users_to_share)
                    messages.info(request, f'File shared with {len(users_to_share)} user(s).')
            
            messages.success(request, f'File "{file_name}" uploaded successfully!')
            return redirect('sharing_files')
            
        except Exception as e:
            messages.error(request, f'Error uploading file: {str(e)}')
            return redirect('sharing_files')
    
    # Handle GET request - display files with filtering
    selected_types = request.GET.getlist('file_type')
    
    # Start with all files visible to user
    # User can see files they uploaded OR files shared with them
    files = SharedFile.objects.filter(
        Q(uploaded_by=request.user) | Q(shared_with=request.user)
    ).distinct().order_by('-uploaded_at')
    
    # Apply file type filter if selected
    if selected_types:
        files = files.filter(extension__in=selected_types)
    
    context = {
        'files': files,
        'file_types': file_types,
        'selected_types': selected_types,
    }
    
    return render(request, 'sharing_files.html', context)

@login_required
def delete_file(request, file_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Delete a file (only if user owns it)"""
    try:
        file = SharedFile.objects.get(id=file_id, uploaded_by=request.user)
        file_name = file.name
        
        # Delete the physical file
        if file.file and os.path.exists(file.file.path):
            os.remove(file.file.path)
        
        file.delete()
        messages.success(request, f'File "{file_name}" deleted successfully!')
    except SharedFile.DoesNotExist:
        messages.error(request, 'File not found or you do not have permission to delete it.')
    except Exception as e:
        messages.error(request, f'Error deleting file: {str(e)}')
    
    return redirect('sharing_files')

@login_required
def download_file(request, file_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Increment download count and serve file"""
    try:
        file = SharedFile.objects.get(
            Q(id=file_id) & 
            (Q(uploaded_by=request.user) | Q(shared_with=request.user))
        )
        file.increment_download_count()
        
        # Return file for download
        file_path = file.file.path
        if os.path.exists(file_path):
            response = FileResponse(open(file_path, 'rb'))
            response['Content-Disposition'] = f'attachment; filename="{file.name}"'
            return response
        else:
            messages.error(request, 'File not found on server.')
            
    except SharedFile.DoesNotExist:
        messages.error(request, 'File not found or access denied.')
    
    return redirect('sharing_files')

@login_required
def share_file(request, file_id):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Share a file with other users"""
    if request.method == 'POST':
        try:
            file = SharedFile.objects.get(id=file_id, uploaded_by=request.user)
            share_with = request.POST.get('share_with', '').strip()
            
            if share_with:
                users_to_share = []
                share_list = [s.strip() for s in share_with.split(',') if s.strip()]
                
                for identifier in share_list:
                    try:
                        if '@' in identifier:
                            user = User.objects.get(email=identifier)
                        else:
                            user = User.objects.get(username=identifier)
                        
                        if user != request.user:
                            users_to_share.append(user)
                    except User.DoesNotExist:
                        messages.warning(request, f'User "{identifier}" not found.')
                        continue
                
                if users_to_share:
                    file.shared_with.add(*users_to_share)
                    messages.success(request, f'File shared with {len(users_to_share)} user(s)')
                else:
                    messages.warning(request, 'No valid users found to share with.')
            else:
                messages.warning(request, 'Please enter usernames or emails to share with.')
                
        except SharedFile.DoesNotExist:
            messages.error(request, 'File not found or you do not have permission to share it.')
        except Exception as e:
            messages.error(request, f'Error sharing file: {str(e)}')
    
    return redirect('sharing_files')

@login_required
def bulk_delete_files(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    """Delete multiple files at once"""
    if request.method == 'POST':
        file_ids = request.POST.getlist('file_ids')
        deleted_count = 0
        
        for file_id in file_ids:
            try:
                file = SharedFile.objects.get(id=int(file_id), uploaded_by=request.user)
                
                # Delete the physical file
                if file.file and os.path.exists(file.file.path):
                    os.remove(file.file.path)
                
                file.delete()
                deleted_count += 1
            except (SharedFile.DoesNotExist, ValueError):
                continue
        
        if deleted_count > 0:
            messages.success(request, f'{deleted_count} file(s) deleted successfully!')
        else:
            messages.warning(request, 'No files were deleted.')
    
    return redirect('sharing_files')

def password_reset(request, user_id):
    if not check_access(request):
        return redirect('logout')  # Create this view

    user = get_object_or_404(User, id=user_id)
    
    
    # Check if user is disabled
    if not user.is_active:
        return render(request, 'login.html', {
            'alert_type': 'error',
            'alert_message': 'Your account is disabled. Contact administrator.'
        })
    
    # Set the new password
    DEFAULT_PASSWORD = f"{user.username}@123"
    user.set_password(DEFAULT_PASSWORD)
    user.save()
    
    # Show success on login page (or wherever you want)
    return render(request, 'dashboard.html', {
        'alert_type': 'success',
        'alert_message': f'Password reset for {user.username}. New password: {DEFAULT_PASSWORD}'
    })

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.hashers import check_password

@login_required
def password_change(request):
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # If POST request, process the form
    if request.method == 'POST':
        old_password = request.POST.get('old_password')
        new_password1 = request.POST.get('new_password1')
        new_password2 = request.POST.get('new_password2')
        
        print("Password change attempt:", old_password, new_password1, new_password2)
        
        # Validate inputs
        if not old_password or not new_password1 or not new_password2:
            return render(request, 'change_password.html', {
                'alert_type': 'error',
                'alert_message': 'All fields are required.'
            })
        
        if not check_password(old_password, request.user.password):
            return render(request, 'change_password.html', {
                'alert_type': 'error',
                'alert_message': 'Your current password is incorrect.'
            })
        
        if new_password1 != new_password2:
            return render(request, 'change_password.html', {
                'alert_type': 'error', 
                'alert_message': 'New passwords do not match.'
            })
        
        if len(new_password1) < 8:
            return render(request, 'change_password.html', {
                'alert_type': 'error',
                'alert_message': 'Password must be at least 8 characters long.'
            })
        
        if new_password1.isdigit():
            return render(request, 'change_password.html', {
                'alert_type': 'error',
                'alert_message': 'Password cannot be entirely numeric.'
            })
        
        common_passwords = ['password', '12345678', 'qwerty123', 'admin123', 'password123', '123456789']
        if new_password1.lower() in common_passwords:
            return render(request, 'change_password.html', {
                'alert_type': 'error',
                'alert_message': 'Password is too common. Choose a stronger password.'
            })
        
        # Update password if all validations pass
        request.user.set_password(new_password1)
        request.user.save()
        update_session_auth_hash(request, request.user)  # Keep user logged in
        
        return render(request, 'change_password.html', {
            'alert_type': 'success',
            'alert_message': 'Your password has been changed successfully!'
        })
    
    # GET request - show the form
    return render(request, 'change_password.html')


@session_auth_required
def event_list(request):
    """
    List threat alerts - SuperAdmin has full access to all alerts
    SuperAdmin from Cyber unit gets full edit/delete access
    Created users can edit/delete their own threats
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_username = request.session.get('user_username', '')
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')  # Get user_id from session
    
    query = request.GET.get('q', '').strip()
    
    # Get all threats
    threats = ThreatAlert.objects.all()
    threats = threats.order_by('-timestamp')
    
    # Search functionality
    if query:
        threats = threats.filter(
            Q(title__icontains=query) |
            Q(content__icontains=query) |
            Q(source__icontains=query) |
            Q(category__name__icontains=query) |
            Q(province__icontains=query)
        )
    
    # Get counts for dashboard
    total_threats = threats.count()
    severity_counts = {
        'critical': threats.filter(severity='critical').count(),
        'high': threats.filter(severity='high').count(),
        'medium': threats.filter(severity='medium').count(),
        'low': threats.filter(severity='low').count(),
    }
    
    # Recent threats (last 7 days)
    import datetime
    one_week_ago = datetime.datetime.now() - datetime.timedelta(days=7)  # ✅ Correct
    # one_week_ago = datetime.now() - timedelta(days=7)
    recent_threats = threats.filter(timestamp__gte=one_week_ago).count()
    
    # Top categories
    top_categories = threats.values('category__name').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # IMPORTANT: Check if user is SuperAdmin from Cyber unit
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin_cyber = (role_str == 'superadmin' or role_str == '1') and unit_str == 'cyber'
    
    # Check for each threat if current user can edit/delete
    threats_with_permissions = []
    for threat in threats:
        # Check if user created this threat
        user_can_edit = False
        if threat.created_by and user_id and str(threat.created_by.id) == str(user_id):
            user_can_edit = True
        
        # Check if user is SuperAdmin Cyber
        if is_superadmin_cyber:
            user_can_edit = True
        
        # Add permission flag to threat object
        threat.user_can_edit = user_can_edit
        threats_with_permissions.append(threat)
    
    # Paginate the queryset
    paginator = Paginator(threats_with_permissions, 10)
    page_number = request.GET.get('page')
    threats_page = paginator.get_page(page_number)
    
    context = {
        'threats': threats_page,
        'total_threats': total_threats,
        'critical_severity': severity_counts['critical'],
        'high_severity': severity_counts['high'],
        'medium_severity': severity_counts['medium'],
        'low_severity': severity_counts['low'],
        'recent_threats': recent_threats,
        'session_username': user_username,
        'user_unit': user_unit,
        'user_role': user_role,
        'search_query': query,
        'top_categories': top_categories,
        'is_superadmin_cyber': is_superadmin_cyber,
        'user_id': user_id,  # Pass user_id to template
    }
    
    return render(request, 'event_list.html', context)

@session_auth_required
def threat_edit(request, pk):
    """
    Edit a threat alert - Only SuperAdmin from Cyber unit or the creator can edit
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    threat = get_object_or_404(ThreatAlert, pk=pk)
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')
    
    # Check if user has permission to edit
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin_cyber = (role_str == 'superadmin' or role_str == '1') and unit_str == 'cyber'
    
    # Check if user is the creator
    is_creator = threat.created_by and user_id and str(threat.created_by.id) == str(user_id)
    
    if not (is_superadmin_cyber or is_creator):
        messages.error(request, "You don't have permission to edit this threat alert.")
        return redirect('event_list')
    
    # Get all categories for the dropdown
    categories = ThreatCategory.objects.all().order_by('name')
    
    # Use model's choices
    SEVERITY_CHOICES = ThreatAlert.SEVERITY_CHOICES
    PROVINCE_CHOICES = ThreatAlert.PROVINCE_CHOICES
    
    # For GET request, prepare form data
    if request.method == 'GET':
        context = {
            'threat': threat,
            'categories': categories,
            'severity_choices': SEVERITY_CHOICES,
            'province_choices': PROVINCE_CHOICES,
            'current_title': threat.title,
            'current_content': threat.content,
            'current_severity': threat.severity,
            'current_province': threat.province,
            'current_source': threat.source,
            'current_url': threat.url,
            'current_category_id': threat.category.id if threat.category else '',
            'next': request.GET.get('next', 'event_list'),
            'is_creator': is_creator,
            'is_superadmin_cyber': is_superadmin_cyber,
        }
        return render(request, 'threat_edit.html', context)
    
    # Handle POST request (form submission)
    elif request.method == 'POST':
        try:
            with transaction.atomic():
                # Get form data
                title = request.POST.get('title', '').strip()
                content = request.POST.get('content', '').strip()
                category_id = request.POST.get('category', '')
                severity = request.POST.get('severity', 'medium')
                province = request.POST.get('province', '')
                source = request.POST.get('source', 'unknown')
                url = request.POST.get('url', '').strip()
                
                # Initialize error tracking
                has_errors = False
                
                # Validate required fields
                if not title:
                    messages.error(request, "Title is required.")
                    has_errors = True
                elif len(title) > 300:
                    messages.error(request, "Title cannot exceed 300 characters.")
                    has_errors = True
                
                if not content:
                    messages.error(request, "Content is required.")
                    has_errors = True
                
                if not url:
                    messages.error(request, "URL is required.")
                    has_errors = True
                
                # Validate severity choice
                valid_severities = [choice[0] for choice in SEVERITY_CHOICES]
                if severity not in valid_severities:
                    messages.error(request, "Invalid severity level selected.")
                    has_errors = True
                
                # Validate province choice
                if province:  # Province is optional
                    valid_provinces = [choice[0] for choice in PROVINCE_CHOICES]
                    if province not in valid_provinces:
                        messages.error(request, "Invalid province selected.")
                        has_errors = True
                
                # If validation errors, re-render form with data
                if has_errors:
                    context = {
                        'threat': threat,
                        'categories': categories,
                        'severity_choices': SEVERITY_CHOICES,
                        'province_choices': PROVINCE_CHOICES,
                        'current_title': title,
                        'current_content': content,
                        'current_severity': severity,
                        'current_province': province,
                        'current_source': source,
                        'current_url': url,
                        'current_category_id': category_id,
                        'next': request.POST.get('next', 'event_list'),
                        'is_creator': is_creator,
                        'is_superadmin_cyber': is_superadmin_cyber,
                    }
                    return render(request, 'threat_edit.html', context)
                
                # Check if URL is unique (excluding current threat)
                if ThreatAlert.objects.filter(url=url).exclude(pk=threat.pk).exists():
                    messages.error(request, "A threat alert with this URL already exists.")
                    context = {
                        'threat': threat,
                        'categories': categories,
                        'severity_choices': SEVERITY_CHOICES,
                        'province_choices': PROVINCE_CHOICES,
                        'current_title': title,
                        'current_content': content,
                        'current_severity': severity,
                        'current_province': province,
                        'current_source': source,
                        'current_url': url,
                        'current_category_id': category_id,
                        'next': request.POST.get('next', 'event_list'),
                        'is_creator': is_creator,
                        'is_superadmin_cyber': is_superadmin_cyber,
                    }
                    return render(request, 'threat_edit.html', context)
                
                # Update the threat object
                threat.title = title
                threat.content = content
                threat.severity = severity
                threat.source = source
                threat.url = url
                threat.updated_at = timezone.now()
                
                # Handle province (optional)
                threat.province = province if province else None
                
                # Handle category
                if category_id:
                    try:
                        category = ThreatCategory.objects.get(id=category_id)
                        threat.category = category
                    except ThreatCategory.DoesNotExist:
                        threat.category = None
                        messages.warning(request, "Selected category was not found. Category has been removed.")
                else:
                    threat.category = None
                
                # Handle image upload
                if 'image' in request.FILES:
                    image_file = request.FILES['image']
                    # Validate file type
                    allowed_image_types = ['image/jpeg', 'image/jpg', 'image/png', 'image/gif', 'image/webp']
                    if image_file.content_type not in allowed_image_types:
                        messages.error(request, "Invalid image format. Please upload JPEG, PNG, GIF, or WebP images.")
                        context = {
                            'threat': threat,
                            'categories': categories,
                            'severity_choices': SEVERITY_CHOICES,
                            'province_choices': PROVINCE_CHOICES,
                            'current_title': title,
                            'current_content': content,
                            'current_severity': severity,
                            'current_province': province,
                            'current_source': source,
                            'current_url': url,
                            'current_category_id': category_id,
                            'next': request.POST.get('next', 'event_list'),
                            'is_creator': is_creator,
                            'is_superadmin_cyber': is_superadmin_cyber,
                        }
                        return render(request, 'threat_edit.html', context)
                    
                    # Validate file size (5MB limit)
                    if image_file.size > 5 * 1024 * 1024:
                        messages.error(request, "Image size exceeds 5MB limit.")
                        context = {
                            'threat': threat,
                            'categories': categories,
                            'severity_choices': SEVERITY_CHOICES,
                            'province_choices': PROVINCE_CHOICES,
                            'current_title': title,
                            'current_content': content,
                            'current_severity': severity,
                            'current_province': province,
                            'current_source': source,
                            'current_url': url,
                            'current_category_id': category_id,
                            'next': request.POST.get('next', 'event_list'),
                            'is_creator': is_creator,
                            'is_superadmin_cyber': is_superadmin_cyber,
                        }
                        return render(request, 'threat_edit.html', context)
                    
                    # Delete old image if exists
                    if threat.image:
                        # Get the file path and delete it
                        if os.path.isfile(threat.image.path):
                            os.remove(threat.image.path)
                    
                    threat.image = image_file
                
                # Handle video upload
                if 'video' in request.FILES:
                    video_file = request.FILES['video']
                    # Validate file type
                    allowed_video_types = ['video/mp4', 'video/webm', 'video/ogg']
                    if video_file.content_type not in allowed_video_types:
                        messages.error(request, "Invalid video format. Please upload MP4, WebM, or OGG videos.")
                        context = {
                            'threat': threat,
                            'categories': categories,
                            'severity_choices': SEVERITY_CHOICES,
                            'province_choices': PROVINCE_CHOICES,
                            'current_title': title,
                            'current_content': content,
                            'current_severity': severity,
                            'current_province': province,
                            'current_source': source,
                            'current_url': url,
                            'current_category_id': category_id,
                            'next': request.POST.get('next', 'event_list'),
                            'is_creator': is_creator,
                            'is_superadmin_cyber': is_superadmin_cyber,
                        }
                        return render(request, 'threat_edit.html', context)
                    
                    # Validate file size (50MB limit)
                    if video_file.size > 50 * 1024 * 1024:
                        messages.error(request, "Video size exceeds 50MB limit.")
                        context = {
                            'threat': threat,
                            'categories': categories,
                            'severity_choices': SEVERITY_CHOICES,
                            'province_choices': PROVINCE_CHOICES,
                            'current_title': title,
                            'current_content': content,
                            'current_severity': severity,
                            'current_province': province,
                            'current_source': source,
                            'current_url': url,
                            'current_category_id': category_id,
                            'next': request.POST.get('next', 'event_list'),
                            'is_creator': is_creator,
                            'is_superadmin_cyber': is_superadmin_cyber,
                        }
                        return render(request, 'threat_edit.html', context)
                    
                    # Delete old video if exists
                    if threat.video:
                        if os.path.isfile(threat.video.path):
                            os.remove(threat.video.path)
                    
                    threat.video = video_file
                
                # Handle image removal checkbox
                if request.POST.get('remove_image') == 'true':
                    if threat.image:
                        if os.path.isfile(threat.image.path):
                            os.remove(threat.image.path)
                        threat.image = None
                
                # Handle video removal checkbox
                if request.POST.get('remove_video') == 'true':
                    if threat.video:
                        if os.path.isfile(threat.video.path):
                            os.remove(threat.video.path)
                        threat.video = None
                
                # Save the updated threat
                threat.save()
                
                messages.success(request, "Threat alert has been updated successfully.")
                
                # Redirect based on next parameter or default to event list
                redirect_to = request.POST.get('next', 'event_list')
                if redirect_to:
                    return redirect(redirect_to)
                
                return redirect('event_list')
                
        except Exception as e:
            messages.error(request, f"An error occurred while updating the threat alert: {str(e)}")
            
            # Re-render form with submitted data
            context = {
                'threat': threat,
                'categories': categories,
                'severity_choices': SEVERITY_CHOICES,
                'province_choices': PROVINCE_CHOICES,
                'current_title': request.POST.get('title', ''),
                'current_content': request.POST.get('content', ''),
                'current_severity': request.POST.get('severity', 'medium'),
                'current_province': request.POST.get('province', ''),
                'current_source': request.POST.get('source', 'unknown'),
                'current_url': request.POST.get('url', ''),
                'current_category_id': request.POST.get('category', ''),
                'next': request.POST.get('next', 'event_list'),
                'is_creator': is_creator,
                'is_superadmin_cyber': is_superadmin_cyber,
            }
            return render(request, 'threat_edit.html', context)

@session_auth_required
def threat_delete(request, pk):
    """
    Delete a threat alert - Only SuperAdmin from Cyber unit or the creator can delete
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    threat = get_object_or_404(ThreatAlert, pk=pk)
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')
    
    # Check if user has permission to delete
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin_cyber = (role_str == 'superadmin' or role_str == '1') and unit_str == 'cyber'
    
    # Check if user is the creator
    is_creator = threat.created_by and user_id and str(threat.created_by.id) == str(user_id)
    
    if not (is_superadmin_cyber or is_creator):
        messages.error(request, "You don't have permission to delete this threat alert.")
        return redirect('event_list')
    
    if request.method == 'POST':
        try:
            # Store threat info for message
            threat_title = threat.title
            
            # Delete associated media files
            if threat.image:
                if os.path.isfile(threat.image.path):
                    os.remove(threat.image.path)
            
            if threat.video:
                if os.path.isfile(threat.video.path):
                    os.remove(threat.video.path)
            
            # Delete the threat object
            threat.delete()
            
            messages.success(request, f'Threat alert "{threat_title}" has been deleted.')
            return redirect('event_list')
            
        except Exception as e:
            messages.error(request, f"An error occurred while deleting the threat alert: {str(e)}")
            return redirect('event_list')
    
    # If GET request, redirect to event list
    return redirect('event_list')

@session_auth_required
def list_autonews(request):
    """
    List AutoNews articles - with permission logic
    SuperAdmin has full access, users can edit/delete their own articles
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_username = request.session.get('user_username', '')
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')
    
    query = request.GET.get('q', '').strip()
    
    # Get all auto news articles
    articles = AutoNewsArticle.objects.all()
    articles = articles.order_by('-created_at')
    
    # Search functionality
    if query:
        articles = articles.filter(
            Q(title__icontains=query) |
            Q(summary__icontains=query) |
            Q(source__icontains=query) |
            Q(keywords__icontains=query) |
            Q(categories__icontains=query)
        )
    
    # Get counts for dashboard
    total_articles = articles.count()
    
    # Priority counts
    priority_counts = {
        'high': articles.filter(priority='high').count(),
        'medium': articles.filter(priority='medium').count(),
        'low': articles.filter(priority='low').count(),
    }
    
    # Threat level counts
    threat_level_counts = {
        'critical': articles.filter(threat_level='critical').count(),
        'high': articles.filter(threat_level='high').count(),
        'medium': articles.filter(threat_level='medium').count(),
        'low': articles.filter(threat_level='low').count(),
    }
    
    # Recent articles (last 7 days)
    #  import datetime
    one_week_ago = datetime.now() - timedelta(days=7)  # ✅ Correct
    # one_week_ago = datetime.now() - timedelta(days=7)
    recent_articles = articles.filter(created_at__gte=one_week_ago).count()
    
    # Top sources
    top_sources = articles.values('source').annotate(
        count=Count('id')
    ).order_by('-count')[:5]
    
    # IMPORTANT: Check if user is SuperAdmin
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin = (role_str == 'superadmin' or role_str == '1')
    
    # Check for each article if current user can edit/delete
    articles_with_permissions = []
    for article in articles:
        # Check if user created this article
        user_can_edit = False
        if article.created_by and user_id and str(article.created_by.id) == str(user_id):
            user_can_edit = True
        
        # Check if user is SuperAdmin
        if is_superadmin:
            user_can_edit = True
        
        # Add permission flag to article object
        article.user_can_edit = user_can_edit
        articles_with_permissions.append(article)
    
    # Paginate the queryset
    paginator = Paginator(articles_with_permissions, 10)
    page_number = request.GET.get('page')
    articles_page = paginator.get_page(page_number)
    
    context = {
        'articles': articles_page,
        'total_articles': total_articles,
        'priority_counts': priority_counts,
        'threat_level_counts': threat_level_counts,
        'recent_articles': recent_articles,
        'session_username': user_username,
        'user_unit': user_unit,
        'user_role': user_role,
        'search_query': query,
        'top_sources': top_sources,
        'is_superadmin': is_superadmin,
        'user_id': user_id,
    }
    
    return render(request, 'autonews_list.html', context)

@session_auth_required
def edit_autonews(request, pk):
    """
    Edit an AutoNews article - Only SuperAdmin or the creator can edit
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    article = get_object_or_404(AutoNewsArticle, pk=pk)
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')
    
    # Check if user has permission to edit
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin = (role_str == 'superadmin' or role_str == '1')
    
    # Check if user is the creator
    is_creator = article.created_by and user_id and str(article.created_by.id) == str(user_id)
    
    if not (is_superadmin or is_creator):
        messages.error(request, "You don't have permission to edit this news article.")
        return redirect('list_autonews')
    
    # Define choices for dropdowns (add these to your model if not already)
    PRIORITY_CHOICES = [
        ('low', 'Low'),
        ('medium', 'Medium'),
        ('high', 'High'),
    ]
    
    THREAT_LEVEL_CHOICES = [
        ('low', 'Low'),
        ('medium', 'Medium'),
        ('high', 'High'),
        ('critical', 'Critical'),
    ]
    
    # For GET request, prepare form data
    if request.method == 'GET':
        context = {
            'article': article,
            'priority_choices': PRIORITY_CHOICES,
            'threat_level_choices': THREAT_LEVEL_CHOICES,
            'current_title': article.title,
            'current_summary': article.summary,
            'current_url': article.url,
            'current_image_url': article.image_url or '',
            'current_source': article.source,
            'current_date': article.date,
            'current_content_length': article.content_length,
            'current_priority': article.priority,
            'current_threat_level': article.threat_level,
            'current_keywords': article.keywords,
            'current_categories': article.categories,
            'is_creator': is_creator,
            'is_superadmin': is_superadmin,
        }
        return render(request, 'autonews_edit.html', context)
    
    # Handle POST request (form submission)
    elif request.method == 'POST':
        try:
            with transaction.atomic():
                # Get form data
                title = request.POST.get('title', '').strip()
                summary = request.POST.get('summary', '').strip()
                url = request.POST.get('url', '').strip()
                image_url = request.POST.get('image_url', '').strip()
                source = request.POST.get('source', '').strip()
                date = request.POST.get('date', '').strip()
                content_length = request.POST.get('content_length', '0')
                priority = request.POST.get('priority', 'medium')
                threat_level = request.POST.get('threat_level', 'low')
                keywords = request.POST.get('keywords', '').strip()
                categories = request.POST.get('categories', '').strip()
                
                # Initialize error tracking
                has_errors = False
                
                # Validate required fields
                if not title:
                    messages.error(request, "Title is required.")
                    has_errors = True
                elif len(title) > 500:
                    messages.error(request, "Title cannot exceed 500 characters.")
                    has_errors = True
                
                if not summary:
                    messages.error(request, "Summary is required.")
                    has_errors = True
                
                if not url:
                    messages.error(request, "URL is required.")
                    has_errors = True
                elif len(url) > 1000:
                    messages.error(request, "URL cannot exceed 1000 characters.")
                    has_errors = True
                
                if not source:
                    messages.error(request, "Source is required.")
                    has_errors = True
                
                if not date:
                    messages.error(request, "Date is required.")
                    has_errors = True
                
                # Validate priority choice
                valid_priorities = [choice[0] for choice in PRIORITY_CHOICES]
                if priority not in valid_priorities:
                    messages.error(request, "Invalid priority selected.")
                    has_errors = True
                
                # Validate threat level choice
                valid_threat_levels = [choice[0] for choice in THREAT_LEVEL_CHOICES]
                if threat_level not in valid_threat_levels:
                    messages.error(request, "Invalid threat level selected.")
                    has_errors = True
                
                # Validate content_length is integer
                try:
                    content_length_int = int(content_length)
                    if content_length_int < 0:
                        messages.error(request, "Content length cannot be negative.")
                        has_errors = True
                except ValueError:
                    messages.error(request, "Content length must be a number.")
                    has_errors = True
                
                # If validation errors, re-render form with data
                if has_errors:
                    context = {
                        'article': article,
                        'priority_choices': PRIORITY_CHOICES,
                        'threat_level_choices': THREAT_LEVEL_CHOICES,
                        'current_title': title,
                        'current_summary': summary,
                        'current_url': url,
                        'current_image_url': image_url,
                        'current_source': source,
                        'current_date': date,
                        'current_content_length': content_length,
                        'current_priority': priority,
                        'current_threat_level': threat_level,
                        'current_keywords': keywords,
                        'current_categories': categories,
                        'is_creator': is_creator,
                        'is_superadmin': is_superadmin,
                    }
                    return render(request, 'autonews_edit.html', context)
                
                # Check if URL is unique for this user (excluding current article)
                # Since unique_together is ['url', 'created_by']
                if AutoNewsArticle.objects.filter(
                    url=url, 
                    created_by=article.created_by
                ).exclude(pk=article.pk).exists():
                    messages.error(request, "You already have an article with this URL.")
                    context = {
                        'article': article,
                        'priority_choices': PRIORITY_CHOICES,
                        'threat_level_choices': THREAT_LEVEL_CHOICES,
                        'current_title': title,
                        'current_summary': summary,
                        'current_url': url,
                        'current_image_url': image_url,
                        'current_source': source,
                        'current_date': date,
                        'current_content_length': content_length,
                        'current_priority': priority,
                        'current_threat_level': threat_level,
                        'current_keywords': keywords,
                        'current_categories': categories,
                        'is_creator': is_creator,
                        'is_superadmin': is_superadmin,
                    }
                    return render(request, 'autonews_edit.html', context)
                
                # Update the article object
                article.title = title
                article.summary = summary
                article.url = url
                article.image_url = image_url if image_url else None
                article.source = source
                article.date = date
                article.content_length = content_length_int
                article.priority = priority
                article.threat_level = threat_level
                article.keywords = keywords
                article.categories = categories
                
                # Save the updated article
                article.save()
                
                messages.success(request, "News article has been updated successfully.")
                return redirect('list_autonews')
                
        except Exception as e:
            messages.error(request, f"An error occurred while updating the article: {str(e)}")
            
            # Re-render form with submitted data
            context = {
                'article': article,
                'priority_choices': PRIORITY_CHOICES,
                'threat_level_choices': THREAT_LEVEL_CHOICES,
                'current_title': request.POST.get('title', ''),
                'current_summary': request.POST.get('summary', ''),
                'current_url': request.POST.get('url', ''),
                'current_image_url': request.POST.get('image_url', ''),
                'current_source': request.POST.get('source', ''),
                'current_date': request.POST.get('date', ''),
                'current_content_length': request.POST.get('content_length', '0'),
                'current_priority': request.POST.get('priority', 'medium'),
                'current_threat_level': request.POST.get('threat_level', 'low'),
                'current_keywords': request.POST.get('keywords', ''),
                'current_categories': request.POST.get('categories', ''),
                'is_creator': is_creator,
                'is_superadmin': is_superadmin,
            }
            return render(request, 'autonews_edit.html', context)

@session_auth_required
def delete_autonews(request, pk):
    """
    Delete an AutoNews article - Only SuperAdmin or the creator can delete
    """
    if not check_access(request):
        return redirect('logout')  # Create this view
    
    article = get_object_or_404(AutoNewsArticle, pk=pk)
    
    # Get user info from session
    user_role = request.session.get('user_role', 0)
    user_unit = request.session.get('user_unit', '')
    user_id = request.session.get('user_id', '')
    
    # Check if user has permission to delete
    role_str = str(user_role).lower() if user_role else ''
    unit_str = str(user_unit).lower() if user_unit else ''
    
    is_superadmin = (role_str == 'superadmin' or role_str == '1')
    
    # Check if user is the creator
    is_creator = article.created_by and user_id and str(article.created_by.id) == str(user_id)
    
    if not (is_superadmin or is_creator):
        messages.error(request, "You don't have permission to delete this news article.")
        return redirect('list_autonews')
    
    if request.method == 'POST':
        try:
            # Store article info for message
            article_title = article.title
            
            # Delete the article object
            article.delete()
            
            messages.success(request, f'News article "{article_title}" has been deleted.')
            return redirect('list_autonews')
            
        except Exception as e:
            messages.error(request, f"An error occurred while deleting the article: {str(e)}")
            return redirect('list_autonews')
    
    # If GET request, redirect to list
    return redirect('list_autonews')


# views.py
import os
from django.http import FileResponse
from django.conf import settings
from django.shortcuts import render

def download_exe_file(request):
    # Get filename from query parameter
    filename = request.GET.get('file', 'FacebookScraper.exe')
    
    # Define the path to your .exe file
    # Make sure FacebookScraper.exe is in your project's base directory
    file_path = os.path.join(settings.BASE_DIR, filename)
    
    # Check if file exists
    if not os.path.exists(file_path):
        # If file doesn't exist, show an error page
        context = {
            'error': f"File '{filename}' not found on server.",
            'file_available': False
        }
        return render(request, 'your_template.html', context)
    
    try:
        # Open the file in binary mode
        file = open(file_path, 'rb')
        
        # Create FileResponse with appropriate headers
        response = FileResponse(file)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Add file size header (optional)
        response['Content-Length'] = os.path.getsize(file_path)
        
        return response
    
    except Exception as e:
        context = {
            'error': f"Error downloading file: {str(e)}",
            'file_available': False
        }
        return render(request, 'your_template.html', context)

def download_page(request):
    # This is the view that renders your download page template
    filename = "FacebookScraper.exe"
    file_path = os.path.join(settings.BASE_DIR, filename)
    
    # Check if file exists and get its size
    file_available = os.path.exists(file_path)
    file_size = None
    
    if file_available:
        # Convert bytes to MB
        size_bytes = os.path.getsize(file_path)
        file_size = round(size_bytes / (1024 * 1024), 2)
    
    context = {
        'file_available': file_available,
        'file_size': file_size,
        'file_version': "2.1.0",  # You might want to get this dynamically
        'error': None if file_available else "File not available for download"
    }
    
    return render(request, 'catchinglink.html', context)

import sqlite3
import os
import datetime
from django.http import HttpResponse
from django.conf import settings
from django.contrib import messages
from django.shortcuts import redirect

# Removed @staff_member_required - now ANY user can download
def download_database_backup(request):
    """
    Direct database download when clicked in sidebar - accessible to all users
    """
    try:
        # Generate filename with timestamp
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'database_backup_{timestamp}.sql'
        
        # Connect to SQLite database
        db_path = os.path.join(settings.BASE_DIR, 'db.sqlite3')
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Create backup content as a list
        backup_lines = []
        
        # Header
        backup_lines.append(f"""-- Database Backup created on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
SET FOREIGN_KEY_CHECKS = 0;
START TRANSACTION;

""")
        
        # Get all tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = [t[0] for t in cursor.fetchall() if not t[0].startswith('sqlite_')]
        
        # Drop tables
        for table in reversed(tables):
            backup_lines.append(f"DROP TABLE IF EXISTS `{table}`;\n")
        backup_lines.append("\n")
        
        # Process each table
        for table in tables:
            # Get column info
            cursor.execute(f"PRAGMA table_info('{table}');")
            columns = cursor.fetchall()
            
            if not columns:
                continue
            
            # Build CREATE TABLE
            col_defs = []
            for col in columns:
                col_name = col[1]
                col_type = col[2].upper()
                not_null = col[3]
                is_pk = col[5]
                
                # Map types
                if 'INT' in col_type:
                    mysql_type = 'int(11)'
                elif 'BIGINT' in col_type:
                    mysql_type = 'bigint(20)'
                elif 'TEXT' in col_type:
                    mysql_type = 'text'
                elif 'BOOL' in col_type:
                    mysql_type = 'tinyint(1)'
                elif 'DATETIME' in col_type:
                    mysql_type = 'datetime'
                elif 'VARCHAR' in col_type:
                    mysql_type = col_type.lower()
                else:
                    mysql_type = 'text'
                
                col_def = f"`{col_name}` {mysql_type}"
                if not_null:
                    col_def += " NOT NULL"
                if is_pk and col_name == 'id':
                    col_def += " PRIMARY KEY AUTO_INCREMENT"
                elif is_pk:
                    col_def += " PRIMARY KEY"
                
                col_defs.append(col_def)
            
            backup_lines.append(f"CREATE TABLE `{table}` (\n  " + ",\n  ".join(col_defs) + "\n);\n\n")
            
            # Get data
            cursor.execute(f"SELECT * FROM '{table}';")
            rows = cursor.fetchall()
            
            if rows:
                col_names = [f"`{c[1]}`" for c in columns]
                backup_lines.append(f"INSERT INTO `{table}` ({', '.join(col_names)}) VALUES\n")
                
                values = []
                for row in rows:
                    vals = []
                    for v in row:
                        if v is None:
                            vals.append("NULL")
                        elif isinstance(v, (int, float)):
                            vals.append(str(v))
                        elif isinstance(v, bool):
                            vals.append("1" if v else "0")
                        else:
                            # Fix: Don't use f-string with backslash, use regular string concatenation
                            escaped = str(v).replace("'", "''")
                            vals.append("'" + escaped + "'")
                    values.append("(" + ", ".join(vals) + ")")
                
                backup_lines.append(",\n".join(values) + ";\n\n")
        
        # Footer
        backup_lines.append("COMMIT;\nSET FOREIGN_KEY_CHECKS = 1;\n")
        
        conn.close()
        
        # Join all lines and create response
        backup_content = ''.join(backup_lines)
        response = HttpResponse(backup_content, content_type='application/sql')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        messages.error(request, f'Backup failed: {str(e)}')
        return redirect(request.META.get('HTTP_REFERER', 'admin:index'))

import requests
import time
import urllib3
import warnings
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from django.core.cache import cache
from django.utils import timezone
from django.db import models
from .models import Website
import pytz
from datetime import datetime

# Set up logging
logger = logging.getLogger(__name__)

# Add file handler for down sites log
down_logger = logging.getLogger('down_sites')
down_logger.setLevel(logging.INFO)

# Create file handler
file_handler = logging.FileHandler('down_sites.log')
file_handler.setLevel(logging.INFO)

# Create formatter
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s', 
                              datefmt='%Y-%m-%d %H:%M:%S')
file_handler.setFormatter(formatter)

# Add handler to logger
down_logger.addHandler(file_handler)

# Also log to console
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(formatter)
down_logger.addHandler(console_handler)

# Suppress SSL warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
warnings.filterwarnings('ignore', message='Unverified HTTPS request')

# Timeout settings
TIMEOUT = 5
MAX_RETRIES = 3
RETRY_DELAY = 3
CACHE_TIMEOUT = 60

# Set Kathmandu timezone
KATHMANDU_TZ = pytz.timezone('Asia/Kathmandu')

def get_kathmandu_time():
    """Get current time in Kathmandu timezone in HH:MM:SS format"""
    return timezone.now().astimezone(KATHMANDU_TZ).strftime("%H:%M:%S")

def get_kathmandu_datetime():
    """Get current datetime in Kathmandu timezone for logging"""
    return timezone.now().astimezone(KATHMANDU_TZ).strftime("%Y-%m-%d %H:%M:%S")

def log_down_site(website, error, attempts, response_time=None, status_code=None):
    """Log details when a site is down"""
    log_entry = (
        f"DOWN - {website.name} | "
        f"URL: {website.url} | "
        f"Error: {error} | "
        f"Attempts: {attempts} | "
        f"Status Code: {status_code if status_code else 'N/A'} | "
        f"Response Time: {response_time if response_time else 'N/A'}ms"
    )
    down_logger.info(log_entry)
    
    # Also log to console with color (if supported)
    print(f"\033[91m[DOWN] {get_kathmandu_datetime()} - {website.name} - {error}\033[0m")

def log_site_recovery(website, response_time, status_code, attempts):
    """Log when a previously down site recovers"""
    log_entry = (
        f"RECOVERED - {website.name} | "
        f"URL: {website.url} | "
        f"Response Time: {response_time}ms | "
        f"Status Code: {status_code} | "
        f"Attempts: {attempts}"
    )
    down_logger.info(log_entry)
    print(f"\033[92m[RECOVERED] {get_kathmandu_datetime()} - {website.name} - Back online\033[0m")

def log_site_check(website, status_info):
    """Log all site checks (optional - for debugging)"""
    if status_info['status'] == 'up':
        logger.debug(f"UP - {website.name} - {status_info['response_time']}ms")
    else:
        logger.warning(f"DOWN - {website.name} - {status_info['error']}")

def check_website_with_retry(website, max_retries=MAX_RETRIES, timeout=TIMEOUT):
    """Check website status with multiple retry attempts and logging"""
    url = website.url
    attempts = 0
    last_error = None
    
    while attempts < max_retries:
        attempts += 1
        try:
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            start = time.time()
            
            response = requests.get(
                url, 
                timeout=timeout,
                headers={'User-Agent': 'Mozilla/5.0'},
                allow_redirects=True,
                verify=False
            )
            
            response_time = round((time.time() - start) * 1000)
            
            # If this was a recovery (previous attempts failed)
            if attempts > 1:
                log_site_recovery(website, response_time, response.status_code, attempts)
            
            return {
                'status': 'up',
                'response_time': response_time,
                'status_code': response.status_code,
                'error': None,
                'attempts': attempts
            }
            
        except requests.Timeout:
            last_error = 'Timeout'
            logger.debug(f"Attempt {attempts}/{max_retries} timeout for {website.name}")
            if attempts < max_retries:
                time.sleep(RETRY_DELAY)
                continue
                
        except requests.ConnectionError:
            last_error = 'Connection failed'
            logger.debug(f"Attempt {attempts}/{max_retries} connection failed for {website.name}")
            if attempts < max_retries:
                time.sleep(RETRY_DELAY)
                continue
                
        except Exception as e:
            last_error = str(e)[:30]
            logger.debug(f"Attempt {attempts}/{max_retries} error for {website.name}: {last_error}")
            if attempts < max_retries:
                time.sleep(RETRY_DELAY)
                continue
    
    # All attempts failed - log the down site
    error_msg = f'{last_error} after {max_retries} attempts'
    log_down_site(website, error_msg, max_retries)
    
    return {
        'status': 'down',
        'response_time': None,
        'status_code': None,
        'error': error_msg,
        'attempts': max_retries
    }

def check_website_batch(websites):
    """Check multiple websites in parallel with logging"""
    results = []
    up_count = 0
    down_count = 0
    fast_sites = 0
    slow_sites = 0
    very_slow_sites = 0
    
    # Log batch start
    logger.info(f"Starting batch check of {len(websites)} websites at {get_kathmandu_datetime()}")
    
    with ThreadPoolExecutor(max_workers=10) as executor:
        future_to_website = {
            executor.submit(check_website_with_retry, website): website 
            for website in websites
        }
        
        for future in as_completed(future_to_website):
            website = future_to_website[future]
            try:
                status_info = future.result(timeout=30)
            except Exception as e:
                status_info = {
                    'status': 'down',
                    'response_time': None,
                    'status_code': None,
                    'error': f'Check failed',
                    'attempts': MAX_RETRIES
                }
                # Log unexpected errors
                log_down_site(website, f'Unexpected error: {str(e)[:50]}', MAX_RETRIES)
            
            if status_info['status'] == 'up':
                up_count += 1
                response_time = status_info.get('response_time')
                
                if response_time and response_time < 1000:
                    fast_sites += 1
                elif response_time and response_time < 3000:
                    slow_sites += 1
                elif response_time:
                    very_slow_sites += 1
            else:
                down_count += 1
            
            results.append({
                'id': website.id,
                'name': website.name,
                'url': website.url,
                'status': status_info['status'],
                'response_time': status_info.get('response_time'),
                'status_code': status_info.get('status_code'),
                'error': status_info.get('error', 'Unknown error'),
                'attempts': status_info.get('attempts', 1)
            })
    
    # Log batch summary
    logger.info(f"Batch check completed - Up: {up_count}, Down: {down_count}, Total: {len(websites)}")
    
    return results, up_count, down_count, fast_sites, slow_sites, very_slow_sites

def monitor_sites(request):
    """Main dashboard view with logging"""
    start_total = time.time()
    
    search_query = request.GET.get('search', '').strip()
    cache_key = f"website_monitor_{search_query}"
    cached_data = cache.get(cache_key)
    force_refresh = request.GET.get('refresh') == '1'
    
    if cached_data and not force_refresh:
        context = cached_data
        context['cached'] = True
        context['load_time'] = round((time.time() - start_total) * 1000)
        return render(request, 'website_monitor.html', context)
    
    websites = Website.objects.filter(is_active=True)
    
    if search_query:
        websites = websites.filter(
            models.Q(name__icontains=search_query) | 
            models.Q(url__icontains=search_query)
        )
    
    website_list = list(websites)
    total = len(website_list)
    
    # Log the check initiation
    logger.info(f"Website check initiated by user at {get_kathmandu_datetime()}")
    
    if total == 0:
        context = {
            'results': [],
            'total': 0,
            'up_count': 0,
            'down_count': 0,
            'fast_sites': 0,
            'slow_sites': 0,
            'very_slow_sites': 0,
            'uptime_percentage': 0,
            'search_query': search_query,
            'last_checked': get_kathmandu_time(),
            'load_time': round((time.time() - start_total) * 1000)
        }
        return render(request, 'website_monitor.html', context)
    
    results, up_count, down_count, fast_sites, slow_sites, very_slow_sites = check_website_batch(website_list)
    uptime_percentage = round((up_count / total) * 100) if total > 0 else 0
    
    results.sort(key=lambda x: (0 if x['status'] == 'up' else 1, x.get('response_time') or 9999))
    
    context = {
        'results': results,
        'total': total,
        'up_count': up_count,
        'down_count': down_count,
        'fast_sites': fast_sites,
        'slow_sites': slow_sites,
        'very_slow_sites': very_slow_sites,
        'uptime_percentage': uptime_percentage,
        'search_query': search_query,
        'last_checked': get_kathmandu_time(),
        'load_time': round((time.time() - start_total) * 1000),
        'retry_count': MAX_RETRIES,
        'retry_delay': RETRY_DELAY
    }
    
    cache.set(cache_key, context, CACHE_TIMEOUT)
    
    return render(request, 'website_monitor.html', context)